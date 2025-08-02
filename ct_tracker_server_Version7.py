import os
import re
import datetime
import json
import time
import threading
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import sqlite3
import zipfile
import uuid
from flask import Flask, Response, render_template_string, request, send_file, abort, session, redirect, url_for, flash, jsonify, after_this_request
import shutil
from functools import wraps
from urllib.parse import unquote
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from pydicom import dcmread
from pydicom.dataset import Dataset, FileMetaDataset
from pydicom.uid import generate_uid, ExplicitVRLittleEndian
from pynetdicom import AE, debug_logger, evt
from pynetdicom.sop_class import (
    StudyRootQueryRetrieveInformationModelFind,
    EncapsulatedPDFStorage,
    SecondaryCaptureImageStorage,
    ModalityWorklistInformationFind,
    ModalityPerformedProcedureStep,
    Verification
)
from PIL import Image
from docx import Document
import random

# --- Enable pynetdicom logging for debugging ---
# debug_logger()

# --- List of all possible permissions for the new system ---
# --- Enable pynetdicom logging for debugging ---
# debug_logger()

# --- List of all possible permissions for the new system ---
# --- List of all possible permissions for the new system ---
ALL_PERMISSIONS = [
    'view_dashboard', 'view_approval_tracker', 'view_us_approval_tracker', 'view_mwl_server', 'view_admin_page', 'view_settings_page',
    'upload_reports', 'delete_reports', # New report permissions
    'manage_approvals', 'add_approval_notes', 'manage_all_subscriptions',
    'manage_own_subscriptions', 'manage_users', 'manage_mwl_entries', 'test_smtp',
    'preregister_patient',
    'download_images',
    'view_resources_page', 'manage_resources'
]

# --- DEFAULT CONFIGURATION ---
DEFAULT_CONFIG = {
    "INSTITUTION_NAME": "CRH Radiology",
    "DICOM_ROOT": "D:\\CT OUTSIDE",
    "PACS_IP": "10.48.25.35",
    "PACS_PORT": 11121,
    "PACS_AE_TITLE": "BOXOUTSIDE",
    "LOCAL_AE_TITLE": "CT_TRACKER",
    "JPG_WATCH_FOLDERS": [
        "C:\\Users\\konicaminolta\\Box\\Dr. Jones\\CT\\DAILY WORKLIST\\TODAY",
        "C:\\Users\\konicaminolta\\Box\\Dr. Palmer\\CT\\DAILY WORKLIST\\TODAY",
        "C:\\Users\\konicaminolta\\Box\\Dr. Aleks\\CT\\DAILY WORKLIST\\TODAY"
    ],
    "US_WATCH_FOLDERS": [
        "C:\\Users\\konicaminolta\\Box\\Ultrasound\\DAILY WORKLIST\\TODAY"
    ],
    "PACS_POLL_INTERVAL": 300,
    "APPROVAL_POLL_INTERVAL": 60,
    "EMAIL_ENABLED": False,
    "RADIOLOGIST_EMAILS": [],
    "ARCHIVING_EMAILS": [], # NEW: For archiving report actions
    "EMAIL_SUBJECT": "New CT Studies Detected",
    "SMTP_HOST": "smtp.example.com",
    "SMTP_PORT": 587,
    "SMTP_USER": "user@example.com",
    "DOCX_TEMPLATE_PATH": "C:\\Program Files\\CT_Tracker\\RADTEMPLATE.docx",
    "SMTP_PASSWORD": "password",
    "SMTP_SENDER_EMAIL": "tracker@example.com",
    "EMAIL_TEMPLATES": {
        "base_template": """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; }}
    .main-container {{ width: 100%; max-width: 600px; margin: 0 auto; border: 1px solid #dee2e6; border-radius: 5px; }}
    .header {{ background-color: #003366; color: white; padding: 20px; text-align: center; font-size: 24px; border-top-left-radius: 5px; border-top-right-radius: 5px;}}
    .content-body {{ padding: 25px; }}
    .content-body h3 {{ color: #003366; margin-top: 0; }}
    .note {{ background-color: #e9ecef; padding: 15px; border-left: 4px solid #0056b3; margin: 15px 0; }}
    .footer {{ background-color: #f1f1f1; color: #666; padding: 15px; font-size: 12px; text-align: center; border-bottom-left-radius: 5px; border-bottom-right-radius: 5px;}}
    .activity-table {{ border-collapse: collapse; width: 100%; margin-top: 15px; }}
    .activity-table th, .activity-table td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
    .activity-table thead {{ background-color: #e9ecef; }}
</style>
</head>
<body>
<div class="main-container">
    <div class="header">RadTrac®</div>
    <div class="content-body">{email_content}</div>
    <div class="footer">This is an automated notification from RadTrac®. Please do not reply to this email.</div>
</div>
</body>
</html>""",
        "new_subscriber": """
<h3>Subscription Confirmed</h3>
<p>Hello {user_full_name},</p>
<p>This email confirms your subscription to updates for patient <strong>{patient_name}</strong> (ID: {patient_id}).</p>
<p>The current status of this request is: <strong>{status}</strong>.</p>
<p>You will be notified of any further changes to this request.</p>
<hr>
<h4>Recent Activity</h4>
{activity_history_table}
""",
        "status_update": """
<h3>Status Update for {patient_name}</h3>
<p>The status of the imaging request for patient <strong>{patient_name}</strong> has been updated to <strong>{status}</strong> by {user_full_name}.</p>
""",
        "new_note": """
<h3>New Note for {patient_name}</h3>
<p>A new note has been added by {user_full_name} regarding the request for <strong>{patient_name}</strong>:</p>
<p class="note"><em>"{note_text}"</em></p>
""",
        "request_completed": """
<h3>Request Completed for {patient_name}</h3>
<p>The imaging study for patient <strong>{patient_name}</strong> (Accession: {accession_number}) has been successfully received by PACS.</p>
<p>The associated approval request has been automatically marked as 'Completed' in the tracker.</p>
""",
        "report_uploaded_subscriber": """
<h3>Report Uploaded for {patient_name}</h3>
<p>A new report has been uploaded for patient <strong>{patient_name}</strong> (Accession: {accession_number}) by {user_full_name}.</p>
<p>You can view the report by logging into the RadTrac® dashboard.</p>
""",
        "report_archive": """
<h3>{action} Report: {patient_name}</h3>
<p>This is an automated notification for archival purposes.</p>
<p>A report for patient <strong>{patient_name}</strong> (Accession: {accession_number}) was <strong>{action_past_tense}</strong> by user <strong>{user_full_name}</strong> at {timestamp}.</p>
<p>The report file is attached to this email for your records.</p>
""",
    },
    # --- MWL Server Settings ---
    "MWL_ENABLED": False,
    "MWL_AE_TITLE": "CT_TRACKER_MWL",
    "MWL_PORT": 11104,
    "MWL_MPPS_ENABLED": True,
    "DEFAULT_ACCESSION_PREFIX": "CRH",
    "DEFAULT_SCHEDULED_STATION_AE": "ANY_MODALITY",
    # --- DOCX Output Folders by Modality ---
    "DOCX_OUTPUT_FOLDERS": {
        "General": "C:\\CT_Tracker_Reports\\_General",
        "CT": "C:\\CT_Tracker_Reports\\CT",
        "DX": "C:\\CT_Tracker_Reports\\DX",
        "US": "C:\\CT_Tracker_Reports\\US",
        "MG": "C:\\CT_Tracker_Reports\\MG",
        "MR": "C:\\CT_Tracker_Reports\\MR"
    },
    # --- User Role Permissions ---
    "USER_ROLE_PERMISSIONS": {
        "admin": ALL_PERMISSIONS,
        "radiology_staff": [
            'view_dashboard', 'view_approval_tracker', 'view_us_approval_tracker', 'view_mwl_server',
            'upload_reports', 'delete_reports', 'manage_approvals', 'add_approval_notes',
            'manage_all_subscriptions', 'manage_own_subscriptions', 'manage_mwl_entries',
            'preregister_patient', 'download_images',
            'view_resources_page', 'manage_resources'
        ],
        "doctor": [
            'view_dashboard', 'view_approval_tracker', 'view_us_approval_tracker', 'add_approval_notes', 'manage_own_subscriptions',
            'view_resources_page'
        ],
        "staff": [
            'view_dashboard', 'view_approval_tracker', 'view_us_approval_tracker', 'add_approval_notes', 'manage_own_subscriptions',
            'view_resources_page'
        ]
    }
}

# --- GLOBAL VARIABLES ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(SCRIPT_DIR, 'config.json')
APPROVAL_DATA_FILE = os.path.join(SCRIPT_DIR, 'approval_data.json')
USER_FILE = os.path.join(SCRIPT_DIR, 'ct_tracker_users.json')
ACTIVITY_LOG_FILE = os.path.join(SCRIPT_DIR, 'ct_tracker_activity.log')
PROCESSED_FILES_LOG = os.path.join(SCRIPT_DIR, 'processed_files.log')
LOGO_PATH = os.path.join(SCRIPT_DIR, 'uploaded_logo.png')
MWL_DB_FILE = os.path.join(SCRIPT_DIR, 'mwl_data.db')
MWL_ACTIVITY_LOG_FILE = os.path.join(SCRIPT_DIR, 'mwl_activity.log')
REPORT_METADATA_FILE = os.path.join(SCRIPT_DIR, 'report_metadata.json')
# --- Quiz File Paths ---
QUIZ_DATA_FILE = os.path.join(SCRIPT_DIR, 'quiz_data.json')
QUIZ_LEADERBOARD_FILE = os.path.join(SCRIPT_DIR, 'quiz_leaderboard.json')
QUIZ_ATTEMPTS_FILE = os.path.join(SCRIPT_DIR, 'quiz_attempts.json')

# ------------------------------------------------------------------
# In-memory store for ongoing quizzes (per username)
quiz_sessions = {}
# ------------------------------------------------------------------




# --- Directory Setup ---
DOWNLOAD_DIR = os.path.join(SCRIPT_DIR, 'downloads')
REPORTS_DIR = os.path.join(SCRIPT_DIR, 'uploaded_reports')
DELETED_REPORTS_DIR = os.path.join(SCRIPT_DIR, 'deleted_reports')
TEMP_REPORTS_DIR = os.path.join(SCRIPT_DIR, 'temp_reports')
RESOURCES_DIR = os.path.join(SCRIPT_DIR, 'resource_uploads')
os.makedirs(DOWNLOAD_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(DELETED_REPORTS_DIR, exist_ok=True)
os.makedirs(TEMP_REPORTS_DIR, exist_ok=True)
os.makedirs(RESOURCES_DIR, exist_ok=True)


# --- Download Task Management & State ---
DOWNLOAD_TASKS = {}
tasks_lock = threading.Lock()

CONFIG = {}
pacs_data_cache = {'timestamp': None, 'data': [], 'status': 'Initializing...'}
approval_cache = {'requests': [], 'timestamp': None}
report_metadata = {}
quiz_data = [] # Will be loaded from quiz_data.json

config_lock = threading.Lock()
pacs_cache_lock = threading.Lock()
approval_cache_lock = threading.Lock()
approval_data_lock = threading.Lock()
mwl_db_lock = threading.Lock()
report_metadata_lock = threading.RLock()
emailed_studies_log = set()
emailed_studies_lock = threading.Lock()
initial_rad_poll_complete = False
rad_poll_lock = threading.Lock()

app = Flask(__name__)
# ... rest of the script

app = Flask(__name__)
app.config['SECRET_KEY'] = 'a-very-secret-key-for-ct-tracker-v26-final-ui-tweaks'
app.config['UPLOAD_FOLDER'] = REPORTS_DIR

@app.before_request
def before_request():
    session.permanent = True
    app.permanent_session_lifetime = datetime.timedelta(minutes=30)
    session.modified = True



@app.before_request
def require_email_profile():
    # Adjusted to allow access to new pre-registration endpoints
    allowed_endpoints = ('login', 'logout', 'serve_logo', 'profile', 'mwl_server',
                         'api_mwl_worklist', 'api_mwl_activity_log', 'add_mwl_patient',
                         'api_mwl_delete', 'api_mwl_get_record', 'api_mwl_update',
                         'preregister_patient', 'api_expected_patients', 'api_get_expected_patient', 'api_cancel_expected_patient',
                         'download_start', 'download_status', 'get_download',
                         'upload_report', 'delete_report', 'generate_report') # Add new report routes
    if 'username' in session and request.endpoint not in allowed_endpoints:
        users = load_users()
        user_data = users.get(session['username'])
        if user_data and not user_data.get('email'):
            flash("Please provide your email address to continue.", "info")
            return redirect(url_for('profile'))

# --- PERMISSIONS & AUTH DECORATORS ---
def has_permission(permission):
    """Checks if the current user's role has a specific permission."""
    if 'role' not in session:
        return False
    role = session['role']
    user_permissions = CONFIG.get('USER_ROLE_PERMISSIONS', {}).get(role, [])
    return permission in user_permissions

@app.context_processor
def inject_permission_checker():
    """Injects the has_permission function into Jinja templates."""
    return dict(has_permission=has_permission, ALL_PERMISSIONS=ALL_PERMISSIONS)

def permission_required(*permissions):
    """Decorator to protect routes based on user permissions."""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'username' not in session:
                return redirect(url_for('login'))
            for permission in permissions:
                if not has_permission(permission):
                    log_activity(f"AUTH_FAIL: User '{session['username']}' lacks permission '{permission}' for endpoint '{request.endpoint}'.")
                    abort(403)
            return f(*args, **kwargs)
        return decorated_function
    return decorator


# --- CONFIGURATION & STATE MANAGEMENT ---
# --- CONFIGURATION & STATE MANAGEMENT ---
def load_config():
    global CONFIG
    with config_lock:
        if not os.path.exists(CONFIG_FILE):
            print(f"INFO: Config file not found. Creating default config at {CONFIG_FILE}")
            with open(CONFIG_FILE, 'w') as f:
                json.dump(DEFAULT_CONFIG, f, indent=4)
        try:
            with open(CONFIG_FILE, 'r') as f:
                config_from_file = json.load(f)
                # Ensure all default keys exist
                for key, value in DEFAULT_CONFIG.items():
                    if key not in config_from_file:
                        config_from_file[key] = value
                    # Special handling for nested dictionaries
                    elif isinstance(value, dict):
                        for sub_key, sub_value in value.items():
                            if isinstance(sub_value, dict): # For USER_ROLE_PERMISSIONS and EMAIL_TEMPLATES
                                config_from_file[key].setdefault(sub_key, {})
                                for item_key, item_value in sub_value.items():
                                     config_from_file[key][sub_key].setdefault(item_key, item_value)
                            else:
                                config_from_file[key].setdefault(sub_key, sub_value)
                CONFIG = config_from_file
        except (json.JSONDecodeError, FileNotFoundError) as e:
            print(f"ERROR: Could not load config file: {e}. Using default.")
            CONFIG = DEFAULT_CONFIG


def save_config(new_config):
    with config_lock:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(new_config, f, indent=4)
    load_config()

def load_report_metadata():
    global report_metadata
    with report_metadata_lock:
        if not os.path.exists(REPORT_METADATA_FILE):
            report_metadata = {}
            return
        try:
            with open(REPORT_METADATA_FILE, 'r') as f:
                report_metadata = json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            log_activity("ERROR: Could not read report_metadata.json. Starting fresh.")
            report_metadata = {}

def save_report_metadata():
    with report_metadata_lock:
        with open(REPORT_METADATA_FILE, 'w') as f:
            json.dump(report_metadata, f, indent=4)

def load_leaderboard():
    """Loads quiz leaderboard data from a JSON file."""
    if not os.path.exists(QUIZ_LEADERBOARD_FILE):
        return []
    try:
        with open(QUIZ_LEADERBOARD_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return []

def save_leaderboard(data):
    """Saves quiz leaderboard data to a JSON file."""
    with open(QUIZ_LEADERBOARD_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)

def load_attempts():
    """Loads user quiz attempt timestamps from a JSON file."""
    if not os.path.exists(QUIZ_ATTEMPTS_FILE):
        return {}
    try:
        with open(QUIZ_ATTEMPTS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {}

def save_attempts(data):
    """Saves user quiz attempt timestamps to a JSON file."""
    with open(QUIZ_ATTEMPTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)

def load_quiz_data():
    """Loads the quiz question bank from the external JSON file."""
    global quiz_data
    if not os.path.exists(QUIZ_DATA_FILE):
        log_activity(f"ERROR: Quiz data file not found at {QUIZ_DATA_FILE}. Quiz will be disabled.")
        quiz_data = []
        return
    try:
        with open(QUIZ_DATA_FILE, 'r', encoding='utf-8') as f:
            quiz_data = json.load(f)
        log_activity(f"SUCCESS: Loaded {len(quiz_data)} questions from {os.path.basename(QUIZ_DATA_FILE)}.")
    except (json.JSONDecodeError, FileNotFoundError) as e:
        log_activity(f"ERROR: Could not load or parse {os.path.basename(QUIZ_DATA_FILE)}: {e}")
        quiz_data = []

# --- LOGGING, USER & APPROVAL MANAGEMENT ---
def log_activity(message):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    console_message = (message[:250] + '...') if len(message) > 250 else message
    print(f"LOG: {console_message}")
    log_entry = f"[{timestamp}] {message}\n"
    try:
        with open(ACTIVITY_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(log_entry)
    except IOError as e:
        print(f"Error: Could not write to activity log file {ACTIVITY_LOG_FILE}: {e}")

def log_mwl_activity(message):
    """Logs a message to the dedicated MWL activity log."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] {message}\n"
    print(f"MWL_LOG: {message}")
    try:
        with open(MWL_ACTIVITY_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(log_entry)
    except IOError as e:
        print(f"Error: Could not write to MWL activity log file {MWL_ACTIVITY_LOG_FILE}: {e}")

def load_users():
    if not os.path.exists(USER_FILE):
        default_users = {"ADMIN": {"password_hash": generate_password_hash("adminuser1"), "role": "admin", "full_name": "Administrator", "email": "", "notify_on_updates": True}}
        save_users(default_users)
        return default_users
    try:
        with open(USER_FILE, 'r') as f:
            users = json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {}
    
    users_updated = False
    for username, data in users.items():
        if 'email' not in data:
            data['email'] = ""
            users_updated = True
        if 'notify_on_updates' not in data:
            data['notify_on_updates'] = True
            users_updated = True
    
    if users_updated:
        log_activity("MIGRATION: Adding 'email' and 'notify_on_updates' fields to user profiles.")
        save_users(users)

    return users

def save_users(users):
    with open(USER_FILE, 'w') as f:
        json.dump(users, f, indent=4)

def load_approval_data():
    with approval_data_lock:
        if not os.path.exists(APPROVAL_DATA_FILE):
            return {}
        try:
            with open(APPROVAL_DATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {}

def save_approval_data(data):
    with approval_data_lock:
        with open(APPROVAL_DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4)


# --- MWL DATABASE & LOGIC FUNCTIONS ---
def init_mwl_db():
    """Initializes the MWL SQLite database and tables if they don't exist."""
    with mwl_db_lock:
        conn = sqlite3.connect(MWL_DB_FILE)
        cursor = conn.cursor()
        # Main worklist table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS patient_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_name TEXT NOT NULL,
                patient_id TEXT NOT NULL,
                accession_number TEXT NOT NULL UNIQUE,
                dob_yyyymmdd TEXT NOT NULL,
                sex TEXT NOT NULL,
                study_date TEXT NOT NULL,
                study_time TEXT NOT NULL,
                study_description TEXT NOT NULL,
                referred_from TEXT,
                modality TEXT NOT NULL,
                requesting_physician TEXT,
                requested_procedure_id TEXT,
                scheduled_station_ae_title TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        # Table for archived records
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS archived_records (
                id INTEGER PRIMARY KEY,
                patient_name TEXT, patient_id TEXT, accession_number TEXT,
                dob_yyyymmdd TEXT, sex TEXT, study_date TEXT, study_time TEXT,
                study_description TEXT, referred_from TEXT, modality TEXT,
                requesting_physician TEXT, requested_procedure_id TEXT,
                scheduled_station_ae_title TEXT,
                created_at TIMESTAMP,
                archived_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        # Table for pre-registered patients
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS expected_patients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_name TEXT NOT NULL,
                patient_id TEXT NOT NULL,
                dob_yyyymmdd TEXT,
                sex TEXT,
                study_description TEXT,
                modality TEXT,
                referred_from TEXT,
                requesting_physician TEXT,
                requesting_physician TEXT,
                scheduled_datetime TEXT NOT NULL,
                status TEXT DEFAULT 'pending', -- pending, registered
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        cursor.execute("CREATE INDEX IF NOT EXISTS idx_accession_number ON patient_records (accession_number)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_scheduled_datetime ON expected_patients (scheduled_datetime)")
        conn.commit()
        conn.close()
    log_mwl_activity(f"MWL database initialized/verified at {MWL_DB_FILE}")


def mwl_db_execute(query, params=(), fetchone=False, fetchall=False, commit=False):
    """Executes a query against the MWL database."""
    with mwl_db_lock:
        conn = sqlite3.connect(MWL_DB_FILE)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        try:
            cursor.execute(query, params)
            if commit:
                conn.commit()
                return cursor.lastrowid
            if fetchone:
                return cursor.fetchone()
            if fetchall:
                return cursor.fetchall()
            return True
        except sqlite3.Error as e:
            log_mwl_activity(f"MWL DB Error: {e} | Query: {query} | Params: {params}")
            return None
        finally:
            conn.close()

def get_modality_from_accession(accession_number):
    acc_upper = accession_number.upper()
    if acc_upper.startswith(CONFIG.get("DEFAULT_ACCESSION_PREFIX", "CRH") + "CT"): return "CT"
    if acc_upper.startswith(CONFIG.get("DEFAULT_ACCESSION_PREFIX", "CRH") + "DX"): return "DX"
    if acc_upper.startswith(CONFIG.get("DEFAULT_ACCESSION_PREFIX", "CRH") + "US"): return "US"
    if acc_upper.startswith(CONFIG.get("DEFAULT_ACCESSION_PREFIX", "CRH") + "MG"): return "MG"
    if acc_upper.startswith(CONFIG.get("DEFAULT_ACCESSION_PREFIX", "CRH") + "MR"): return "MR"
    return None

def check_duplicate_mwl_record(patient_name, patient_id, accession_number):
    now = datetime.datetime.now()
    threshold_dt = now - datetime.timedelta(hours=36)
    threshold_timestamp = threshold_dt.strftime("%Y-%m-%d %H:%M:%S")

    # Check for exact accession number match first (should be unique)
    exact_match = mwl_db_execute("SELECT * FROM patient_records WHERE accession_number = ?", (accession_number,), fetchone=True)
    if exact_match:
        return f"An entry with the same Accession Number '{accession_number}' already exists."

    # Check for recent similar patient
    query_general = """
        SELECT study_description, created_at FROM patient_records
        WHERE (patient_id = ? OR patient_name = ?) AND created_at > ?
        ORDER BY created_at DESC LIMIT 1
    """
    general_match = mwl_db_execute(query_general, (patient_id, patient_name, threshold_timestamp), fetchone=True)
    if general_match:
        return f"A recent record for Patient '{patient_name}' or ID '{patient_id}' was found from {general_match['created_at']}. Please confirm this is a new, distinct study."
    
    return None


# --- DICOM & FILE OPERATIONS ---
def push_dicom_to_pacs(dicom_dataset):
    try:
        ae = AE(ae_title=CONFIG.get('LOCAL_AE_TITLE'))
        ae.add_requested_context(EncapsulatedPDFStorage)
        ae.add_requested_context(SecondaryCaptureImageStorage)
        assoc = ae.associate(CONFIG.get('PACS_IP'), int(CONFIG.get('PACS_PORT')), ae_title=CONFIG.get('PACS_AE_TITLE'))
        if assoc.is_established:
            status = assoc.send_c_store(dicom_dataset)
            if status and status.Status == 0x0000:
                log_activity(f"DICOM PUSH SUCCESS: C-STORE status: 0x0000 for SOP UID {dicom_dataset.SOPInstanceUID}")
            else:
                log_activity(f"DICOM PUSH FAILED: C-STORE status: {status.Status if status else 'N/A'} for SOP UID {dicom_dataset.SOPInstanceUID}")
            assoc.release()
            return status
        else:
            log_activity("DICOM PUSH FAIL: Could not establish association.")
            return None
    except Exception as e:
        log_activity(f"DICOM PUSH CRITICAL ERROR: {e}")
        return None

def create_base_dicom_dataset(study_context):
    ds = Dataset()
    ds.PatientName = study_context.get('PatientName', '')
    ds.PatientID = study_context.get('PatientID', '')
    ds.StudyInstanceUID = study_context.get('StudyInstanceUID', '')
    ds.AccessionNumber = study_context.get('AccessionNumber', '')
    ds.StudyID = study_context.get('StudyID', '1')
    ds.StudyDate = study_context.get('StudyDate', '')
    ds.StudyTime = study_context.get('StudyTime', '')
    ds.ReferringPhysicianName = study_context.get('ReferringPhysicianName', '')
    ds.StudyDescription = study_context.get('StudyDescription', 'Attached Document')
    ds.SeriesInstanceUID = generate_uid()
    ds.SOPInstanceUID = generate_uid()
    ds.SeriesNumber = "999"
    ds.InstanceNumber = "1"
    ds.SeriesDate = datetime.datetime.now().strftime('%Y%m%d')
    ds.SeriesTime = datetime.datetime.now().strftime('%H%M%S')
    ds.file_meta = FileMetaDataset()
    ds.file_meta.FileMetaInformationGroupLength = 180
    ds.file_meta.FileMetaInformationVersion = b'\x00\x01'
    ds.file_meta.TransferSyntaxUID = ExplicitVRLittleEndian
    ds.file_meta.ImplementationClassUID = generate_uid()
    ds.file_meta.ImplementationVersionName = "CT_TRACKER_V10"
    ds.is_little_endian = True
    ds.is_implicit_VR = False
    return ds

def convert_pdf_to_dicom(file_path, study_context):
    log_activity(f"DICOM CONVERT: Creating DICOM Encapsulated PDF for {os.path.basename(file_path)}")
    ds = create_base_dicom_dataset(study_context)
    ds.SOPClassUID = EncapsulatedPDFStorage
    ds.file_meta.MediaStorageSOPClassUID = EncapsulatedPDFStorage
    ds.Modality = 'DOC'
    ds.ConversionType = 'WSD'
    ds.SeriesDescription = "Attached PDF Report"
    ds.MIMETypeOfEncapsulatedDocument = 'application/pdf'
    with open(file_path, 'rb') as f:
        ds.EncapsulatedDocument = f.read()
    return ds

def convert_jpg_to_dicom(file_path, study_context):
    log_activity(f"DICOM CONVERT: Creating DICOM Secondary Capture for {os.path.basename(file_path)}")
    try:
        img = Image.open(file_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        ds = create_base_dicom_dataset(study_context)
        # MODIFICATION: Set SeriesNumber for JPGs to 000 so they appear first.
        ds.SeriesNumber = "000"
        ds.SOPClassUID = SecondaryCaptureImageStorage
        ds.file_meta.MediaStorageSOPClassUID = SecondaryCaptureImageStorage
        ds.Modality = 'SC'
        ds.SeriesDescription = "Attached JPG Document"
        ds.ConversionType = 'WSD'
        ds.Rows, ds.Columns = img.height, img.width
        ds.SamplesPerPixel = 3
        ds.PhotometricInterpretation = "RGB"
        ds.PlanarConfiguration = 0
        ds.BitsAllocated, ds.BitsStored, ds.HighBit = 8, 8, 7
        ds.PixelRepresentation = 0
        ds.PixelData = img.tobytes()
        return ds
    except Exception as e:
        log_activity(f"ERROR: Failed to convert JPG {file_path}: {e}")
        return None

# --- BACKGROUND WORKERS ---

def notify_radiologists_of_new_studies(new_studies):
    """Emails a summary of brand-new studies to the radiologist list."""
    if not CONFIG.get('EMAIL_ENABLED') or not new_studies:
        return

    recipients = CONFIG.get('RADIOLOGIST_EMAILS', [])
    if not recipients:
        log_activity("EMAILER_WARN: Radiologist notification is enabled, but no recipient emails are configured.")
        return

    subject = f"RadTrac®: {len(new_studies)} New Studies Detected"
    
    # Build a simple HTML table for the email body
    studies_html = """
    <p>The following new studies were detected by the PACS poller:</p>
    <table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%;">
        <tr style="background-color: #f2f2f2;">
            <th>Patient Name</th><th>Accession #</th><th>Study Description</th>
        </tr>
    """
    for study in new_studies:
        studies_html += f"""
        <tr>
            <td>{study.get('patient_name', 'N/A')}</td>
            <td>{study.get('accession', 'N/A')}</td>
            <td>{study.get('study_desc', 'N/A')}</td>
        </tr>"""
    studies_html += "</table>"
    
    try:
        # Note: We are using a simplified context here, not the main templating engine
        # to avoid complexity with logos for this specific notification.
        html_body = f"<html><body>{studies_html}</body></html>"
        
        log_activity(f"NOTIFY_RADIOLOGISTS: Sending new study alert for {len(new_studies)} studies to {len(recipients)} recipients.")
        email_thread = threading.Thread(target=_send_email_worker, args=(subject, html_body, recipients, CONFIG))
        email_thread.start()
    except Exception as e:
        log_activity(f"NOTIFY_RADIOLOGISTS_ERROR: Failed to send email. Error: {e}")

def radiologist_update_poller():
    """
    A dedicated poller that runs every 30 minutes to email radiologists
    a summary of any new studies found since the last check.
    It skips the initial full list and only reports subsequent changes.
    """
    global initial_rad_poll_complete
    log_activity("RADIOLOGIST_POLLER: Starting radiologist update poller thread.")

    # A short delay on startup to ensure the main PACS poller runs at least once first.
    time.sleep(CONFIG.get('PACS_POLL_INTERVAL', 300) + 15)

    while True:
        try:
            with pacs_cache_lock:
                # Get a copy of the current studies from the main cache
                all_current_studies = list(pacs_data_cache.get('data', []))
                current_accessions = {s.get('accession') for s in all_current_studies if s.get('accession')}

            with rad_poll_lock:
                # Check if this is the first time the poller is running
                if not initial_rad_poll_complete:
                    with emailed_studies_lock:
                        # On the first run, just populate the log with all current studies
                        emailed_studies_log.update(current_accessions)
                    initial_rad_poll_complete = True
                    log_activity(f"RADIOLOGIST_POLLER: Initial run complete. Priming cache with {len(current_accessions)} studies. No email will be sent.")
                else:
                    # On subsequent runs, find studies that are in the new list but not in our log
                    with emailed_studies_lock:
                        new_accessions = current_accessions - emailed_studies_log

                    if new_accessions:
                        # If there are new studies, prepare them for the email
                        new_studies_for_email = [
                            s for s in all_current_studies if s.get('accession') in new_accessions
                        ]

                        # Send the notification email containing only the new studies
                        notify_radiologists_of_new_studies(new_studies_for_email)

                        # Update our log with the new accessions so they aren't sent again
                        with emailed_studies_lock:
                            emailed_studies_log.update(new_accessions)
                        log_activity(f"RADIOLOGIST_POLLER: Detected and emailed {len(new_studies_for_email)} new studies.")

        except Exception as e:
            log_activity(f"RADIOLOGIST_POLLER: An unexpected error occurred: {e}")

        # Wait for the configured interval before running again (e.g., 30 minutes)
        time.sleep(1800)

def pacs_poller():
    log_activity("POLLER: Starting PACS poller thread.")
    while True:
        try:
            log_activity("POLLER: Querying PACS for new studies.")
            recent_studies = query_pacs_for_ct_studies()

            if recent_studies:
                approval_data = load_approval_data()
                approval_data_updated = False

                # Create a dictionary of studies by patient ID for faster lookups
                studies_by_pid = {}
                for study in recent_studies:
                    # Use the raw patient ID from PACS for matching
                    pid = study.get('patient_id')
                    if pid:
                        studies_by_pid.setdefault(pid, []).append(study)

                for filename, req in approval_data.items():
                    # Skip if the request is already completed or hidden
                    if not req.get('visible', True) or req.get('status') == 'Completed':
                        continue

                    # Get the patient ID from the parsed filename info
                    req_pid = req.get('parsed_info', {}).get('patient_id')
                    if not req_pid or req_pid not in studies_by_pid:
                        continue

                    matching_studies_for_req = studies_by_pid[req_pid]
                    
                    try:
                        # Find a study for this patient that is NEWER than the approval request card
                        req_creation_dt = datetime.datetime.fromisoformat(req.get('creation_time'))
                        
                        new_enough_study = next(
                            (s for s in matching_studies_for_req if datetime.datetime.fromisoformat(s.get('log_timestamp', '')) > req_creation_dt), 
                            None
                        )
                        
                        if new_enough_study:
                            log_activity(f"AUTO-COMPLETE: Found PACS study for approval request '{filename}'. Marking as Completed.")
                            req['status'] = 'Completed'
                            now_iso = datetime.datetime.now().isoformat()
                            req['last_activity_on'] = now_iso
                            req['last_activity_by'] = 'System (Auto-Complete)'
                            
                            # Set it to be hidden from the main view after 24 hours
                            req['hide_after'] = (datetime.datetime.now() + datetime.timedelta(hours=24)).isoformat()
                            
                            req.setdefault('notes', []).append({
                                'text': f"Automatically marked 'Completed' (PACS study found: Acc# {new_enough_study.get('accession')}).",
                                'by': 'System (Auto)', 'on': now_iso
                            })
                            
                            patient_name_for_email = req.get('parsed_info', {}).get('patient_name', 'N/A')
                            subject = f"Request Completed: {patient_name_for_email}"
                            context = {
                                "patient_name": patient_name_for_email,
                                "accession_number": new_enough_study.get('accession', 'N/A')
                            }
                            _notify_subscribers(filename, subject, 'request_completed', context)
                            approval_data_updated = True
                    except (ValueError, TypeError) as e:
                        log_activity(f"AUTO-COMPLETE_WARN: Could not process timestamp for '{filename}'. Error: {e}")
                        continue

                if approval_data_updated:
                    save_approval_data(approval_data)
                
                # This part remains to update other warnings like recent scans
                update_approval_warnings(recent_studies, approval_data)
                cleanup_completed_approvals()

        except Exception as e:
            log_activity(f"POLLER: An unexpected error occurred: {e}")
        
        time.sleep(CONFIG.get('PACS_POLL_INTERVAL', 300))

def approval_poller():
    """Periodically re-scan the watch folders and enrich with warnings for the cache."""
    log_activity("APPROVAL_POLLER: starting up...")
    interval = CONFIG.get('APPROVAL_POLL_INTERVAL', 60)
    while True:
        try:
            # Step 1: Scan the folders for the latest request files
            raw_requests = scan_and_prepare_approvals()
            
            # Convert list to dictionary format that the warning function expects
            approval_dict = {req['filename']: req for req in raw_requests}

            # Step 2: Get the latest PACS data from its cache
            with pacs_cache_lock:
                recent_studies = pacs_data_cache.get('data', [])

            # Step 3: Add warnings to the request data
            enriched_approval_data = update_approval_warnings(recent_studies, approval_dict)

            # Step 4: Convert back to a list and sort for the final cache
            final_list = sorted(
                enriched_approval_data.values(),
                key=lambda x: x.get('last_activity_on', ''),
                reverse=True
            )

            with approval_cache_lock:
                approval_cache['requests'] = final_list
                approval_cache['timestamp'] = datetime.datetime.now().isoformat()
            
            log_activity(f"APPROVAL_POLLER: refreshed and enriched {len(final_list)} entries.")
        except Exception as e:
            log_activity(f"APPROVAL_POLLER_ERROR: {e}")
        
        time.sleep(interval)

def mwl_server_worker():
    """Main worker function to run the pynetdicom MWL server."""
    ae = None
    
    # Define handlers inside the worker
    def handle_find(event):
        log_mwl_activity(f"C-FIND request received from {event.assoc.requestor.ae_title}")
        req_identifier = event.identifier
        
        sql_query = "SELECT * FROM patient_records WHERE 1=1"
        params = []

        if 'PatientName' in req_identifier and req_identifier.PatientName:
            sql_query += " AND patient_name LIKE ?"
            params.append(f"%{req_identifier.PatientName.replace('*', '%')}%")
        if 'PatientID' in req_identifier and req_identifier.PatientID:
            sql_query += " AND patient_id = ?"
            params.append(str(req_identifier.PatientID))
        
        sps = req_identifier.ScheduledProcedureStepSequence[0]
        if 'ScheduledProcedureStepStartDate' in sps and sps.ScheduledProcedureStepStartDate:
            date_range = sps.ScheduledProcedureStepStartDate
            if '-' in date_range:
                start, end = date_range.split('-')
                sql_query += " AND study_date BETWEEN ? AND ?"
                params.extend([start, end])
            else:
                sql_query += " AND study_date = ?"
                params.append(date_range)

        if 'Modality' in sps and sps.Modality:
            sql_query += " AND modality = ?"
            params.append(sps.Modality)

        matching_records = mwl_db_execute(sql_query, tuple(params), fetchall=True)
        if matching_records is None:
            yield 0xA700, None # Out of resources
            return
        
        log_mwl_activity(f"Found {len(matching_records)} records matching C-FIND criteria.")

        for record in matching_records:
            ds = Dataset()
            ds.PatientName = record["patient_name"]
            ds.PatientID = record["patient_id"]
            ds.PatientBirthDate = record["dob_yyyymmdd"]
            ds.PatientSex = record["sex"]
            ds.AccessionNumber = record["accession_number"]
            ds.ReferringPhysicianName = record["referred_from"]
            ds.StudyInstanceUID = generate_uid()
            ds.RequestingPhysician = record["requesting_physician"]
            ds.RequestedProcedureDescription = record["study_description"]
            ds.RequestedProcedureID = record["requested_procedure_id"]
            
            sps_item = Dataset()
            sps_item.ScheduledStationAETitle = record["scheduled_station_ae_title"]
            sps_item.ScheduledProcedureStepStartDate = record["study_date"]
            sps_item.ScheduledProcedureStepStartTime = record["study_time"]
            sps_item.Modality = record["modality"]
            sps_item.ScheduledPerformingPhysicianName = ""
            sps_item.ScheduledProcedureStepDescription = record["study_description"]
            sps_item.ScheduledProcedureStepID = record["accession_number"]
            ds.ScheduledProcedureStepSequence = [sps_item]
            
            ds.SpecificCharacterSet = "ISO_IR 100"
            yield 0xFF00, ds

        yield 0x0000, None # Success

    def handle_mpps_create(event):
        ds = event.dataset
        patient_name = str(ds.get("PatientName", "N/A"))
        accession = ds.get("AccessionNumber", "N/A")
        log_mwl_activity(f"MPPS START received from {event.assoc.requestor.ae_title} for {patient_name} (Acc: {accession}).")
        
        response_ds = Dataset()
        response_ds.SOPInstanceUID = ds.SOPInstanceUID
        return 0x0000, response_ds

    def handle_mpps_set(event):
        request_ds = event.request.DataSet
        status = request_ds.get("PerformedProcedureStepStatus", "")
        
        sps_seq = request_ds.PerformedProcedureStepSequence[0]
        accession = sps_seq.ScheduledStepAttributesSequence[0].AccessionNumber

        log_msg = f"MPPS STATUS '{status}' received for Acc: {accession}."
        
        if status == "COMPLETED" and CONFIG.get("MWL_MPPS_ENABLED"):
            log_msg += " Deleting from worklist."
            delete_query = "DELETE FROM patient_records WHERE accession_number = ?"
            result = mwl_db_execute(delete_query, (accession,), commit=True)
            if result is not None:
                log_msg += " -> Success."
            else:
                log_msg += " -> FAILED or not found."
        
        log_mwl_activity(log_msg)
        return 0x0000, None

    def handle_echo(event):
        log_mwl_activity(f"C-ECHO request received from {event.assoc.requestor.ae_title}")
        return 0x0000 

    while True:
        current_config_state = (CONFIG.get('MWL_ENABLED'), CONFIG.get('MWL_PORT'), CONFIG.get('MWL_AE_TITLE'))
        
        if not CONFIG.get('MWL_ENABLED'):
            if ae and ae._server and ae._server.is_running:
                log_mwl_activity("MWL server is disabled in config. Shutting down.")
                ae.shutdown()
            time.sleep(10)
            continue
        
        server_needs_restart = False
        if ae and not ae._servers:
             log_mwl_activity("MWL server was stopped. Restarting based on config.")
             server_needs_restart = True
        
        if ae is None or server_needs_restart:
            try:
                ae = AE(ae_title=CONFIG.get('MWL_AE_TITLE'))
                ae.add_supported_context(ModalityWorklistInformationFind)
                ae.add_supported_context(Verification)
                if CONFIG.get('MWL_MPPS_ENABLED'):
                    ae.add_supported_context(ModalityPerformedProcedureStep)

                handlers = [
                    (evt.EVT_C_ECHO, handle_echo),
                    (evt.EVT_C_FIND, handle_find),
                    (evt.EVT_N_CREATE, handle_mpps_create),
                    (evt.EVT_N_SET, handle_mpps_set)
                ]

                port = int(CONFIG.get('MWL_PORT'))
                log_mwl_activity(f"Starting MWL Server on port {port} with AE Title {CONFIG.get('MWL_AE_TITLE')}...")
                
                ae.start_server(('', port), block=False, evt_handlers=handlers)
                
            except Exception as e:
                log_mwl_activity(f"MWL Server CRITICAL ERROR on startup: {e}. Retrying in 30 seconds.")
                ae = None
                time.sleep(30)
                continue

        # Monitor for config changes or if the server stops unexpectedly
        while ae and ae._servers:
            new_config_state = (CONFIG.get('MWL_ENABLED'), CONFIG.get('MWL_PORT'), CONFIG.get('MWL_AE_TITLE'))
            if new_config_state != current_config_state:
                log_mwl_activity("MWL config change detected. Restarting server...")
                ae.shutdown()
                ae = None 
                break
            time.sleep(5)
        
        if ae and not ae._servers:
            log_mwl_activity("MWL server stopped unexpectedly. Will attempt to restart.")
            ae = None 
            time.sleep(10)


def mwl_archive_worker():
    """Periodically archives old MWL entries."""
    log_activity("ARCHIVER: Starting MWL archive worker thread.")
    while True:
        try:
            cutoff_time = datetime.datetime.now() - datetime.timedelta(hours=4)
            cutoff_str = cutoff_time.strftime("%Y-%m-%d %H:%M:%S")

            # Find records older than the cutoff time
            old_records = mwl_db_execute(
                "SELECT * FROM patient_records WHERE created_at < ?",
                (cutoff_str,), fetchall=True
            )

            if old_records:
                log_mwl_activity(f"ARCHIVER: Found {len(old_records)} MWL entries to archive.")
                for record in old_records:
                    # Copy to archive table
                    archive_query = """
                        INSERT INTO archived_records (id, patient_name, patient_id, accession_number,
                        dob_yyyymmdd, sex, study_date, study_time, study_description, referred_from,
                        modality, requesting_physician, requested_procedure_id, scheduled_station_ae_title, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    archive_params = (
                        record['id'], record['patient_name'], record['patient_id'], record['accession_number'],
                        record['dob_yyyymmdd'], record['sex'], record['study_date'], record['study_time'],
                        record['study_description'], record['referred_from'], record['modality'],
                        record['requesting_physician'], record['requested_procedure_id'],
                        record['scheduled_station_ae_title'], record['created_at']
                    )
                    mwl_db_execute(archive_query, archive_params, commit=True)

                    # Delete from live table
                    mwl_db_execute("DELETE FROM patient_records WHERE id = ?", (record['id'],), commit=True)
                    log_mwl_activity(f"ARCHIVER: Archived and removed MWL entry for {record['patient_name']} (Acc: {record['accession_number']}).")

        except Exception as e:
            log_activity(f"ARCHIVER: An unexpected error occurred in the MWL archive worker: {e}")
        
        # Sleep for 15 minutes before checking again
        time.sleep(900)

def _send_email_worker(subject, html_body, recipients, config_dict, attachments=None):
    sender_email = config_dict.get('SMTP_SENDER_EMAIL')
    msg = MIMEMultipart('related')
    msg['Subject'] = subject
    msg['From'] = sender_email
    msg['To'] = ", ".join(recipients)
    msg.attach(MIMEText(html_body, 'html'))
    if attachments:
        for file_path in attachments:
            if not os.path.isfile(file_path):
                log_activity(f"EMAILER_WARN: Attachment file not found: {file_path}")
                continue
            try:
                with open(file_path, 'rb') as attachment:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(file_path)}"')
                msg.attach(part)
                log_activity(f"EMAILER: Attached {os.path.basename(file_path)} to email.")
            except Exception as e:
                log_activity(f"EMAILER_ERROR: Could not attach file {file_path}: {e}")
    smtp_host = config_dict.get('SMTP_HOST')
    smtp_port = int(config_dict.get('SMTP_PORT', 587))
    smtp_user = config_dict.get('SMTP_USER')
    smtp_password = config_dict.get('SMTP_PASSWORD')
    with smtplib.SMTP(smtp_host, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.sendmail(sender_email, recipients, msg.as_string())

def cleanup_zip_files_worker():
    """Periodically scans the download directory and deletes zip files older than 1 hour."""
    log_activity("CLEANUP_WORKER: Starting zip file cleanup thread.")
    while True:
        try:
            now = time.time()
            one_hour_in_seconds = 3600
            
            if not os.path.isdir(DOWNLOAD_DIR):
                log_activity(f"CLEANUP_WORKER: Download directory not found at {DOWNLOAD_DIR}, skipping.")
                time.sleep(600) # Wait 10 minutes before checking again
                continue

            for filename in os.listdir(DOWNLOAD_DIR):
                if filename.lower().endswith('.zip'):
                    file_path = os.path.join(DOWNLOAD_DIR, filename)
                    try:
                        file_mod_time = os.path.getmtime(file_path)
                        file_age = now - file_mod_time
                        
                        if file_age > one_hour_in_seconds:
                            os.remove(file_path)
                            log_activity(f"CLEANUP_WORKER: Deleted expired zip file: {filename}")

                    except (OSError, FileNotFoundError) as e:
                        log_activity(f"CLEANUP_WORKER_ERROR: Could not process file {filename}. Reason: {e}")
            
            # Sleep for 10 minutes before the next scan
            time.sleep(600)

        except Exception as e:
            log_activity(f"CLEANUP_WORKER: An unexpected critical error occurred: {e}")
            time.sleep(600) # Wait before retrying after a major error        


# --- TEMPLATES ---
RESOURCES_TEMPLATE = """
<!doctype html><html lang="en"><head><title>Resources - RadTrac®</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css">
<style>
    .navbar-custom { background: #003366 !important; }
    .navbar-custom .navbar-brand, .navbar-custom .navbar-text { color: #fff !important; }
    .banner-logo { height: 45px; margin-right: 14px; }
    .quiz-option-btn { width: 100%; text-align: left; }
    .leaderboard-table th { background-color: #e9ecef; }
    .modal-lg-quiz { max-width: 800px; }
    /* Styles for the new calculator */
    .calculator-result-label { font-weight: 500; }
    .calculator-result-value { font-weight: bold; color: #005A9C; }
    .highlight-cannula { color: white; padding: 2px 6px; border-radius: 4px; }
    .iv-site-highlight { background-color: #e7f1ff; border-left: 4px solid #005A9C; padding-left: 10px; }
    /* New style for cannula highlight */
    #iv-info-card ul li { transition: all 0.3s ease-in-out; border-radius: 5px; padding: 4px; }
</style>
</head>
<body>
<nav class="navbar navbar-expand-lg navbar-custom">
    <div class="container-fluid">
        <img src="/logo" alt="Logo" class="banner-logo"><span class="navbar-brand">CRH Radiology - Resources</span>
        <div class="d-flex align-items-center ms-auto">
            <span class="navbar-text me-3">User: {{ session.full_name }}</span>
            <a href="/" class="btn btn-outline-light me-2">Main Tracker</a>
            {% if has_permission('view_approval_tracker') %}<a href="/approval" class="btn btn-outline-info me-2">Approval Tracker</a>{% endif %}
            {% if has_permission('view_us_approval_tracker') %}<a href="/us_approval" class="btn btn-outline-info me-2">US Tracker</a>{% endif %}
            {% if has_permission('view_mwl_server') %}<a href="/mwl" class="btn btn-outline-info me-2">MWL Server</a>{% endif %}
            {% if has_permission('view_admin_page') %}<a href="/admin" class="btn btn-secondary me-2">Admin</a>{% endif %}
            {% if has_permission('view_settings_page') %}<a href="/settings" class="btn btn-info me-2">Settings</a>{% endif %}
            <a href="/logout" class="btn btn-warning">Logout</a>
        </div>
    </div>
</nav>

<div class="container mt-4">
    <!-- TOP ROW: CALCULATOR AND DOCUMENTS -->
    <div class="row">
        <div class="col-lg-7">
            <!-- CT CONTRAST PROTOCOL CALCULATOR START -->
            <div class="card mb-4">
                <div class="card-header bg-primary text-white">
                    <h5 class="mb-0"><i class="bi bi-calculator-fill me-2"></i>CT Contrast Protocol Calculator</h5>
                </div>
                <div class="card-body">
                    <form id="contrastCalculatorForm">
                        <div class="alert alert-danger d-none" id="calc-error-box"></div>
                        <div class="row">
                            <div class="col-md-6 mb-3"><label for="weight_lbs" class="form-label">Weight (lbs):</label><input type="number" class="form-control" id="weight_lbs" required></div>
                            <div class="col-md-6 mb-3"><label for="sex" class="form-label">Sex:</label><select class="form-select" id="sex"><option>Male</option><option>Female</option></select></div>
                            <div class="col-md-6 mb-3"><label for="age" class="form-label">Age (years):</label><input type="number" class="form-control" id="age" required></div>
                            <div class="col-md-6 mb-3"><label for="allergy" class="form-label">Contrast Allergy?:</label><select class="form-select" id="allergy"><option>No</option><option>Yes</option></select></div>
                            <div class="col-md-6 mb-3"><label for="creatinine" class="form-label">Creatinine (µmol/L):</label><input type="number" class="form-control" id="creatinine" step="any" placeholder="Leave blank if not available"></div>
                        </div>
                        <hr>
                        <div class="row">
                            <div class="col-md-6 mb-3"><label for="scan_type" class="form-label">Scan Type:</label><select class="form-select" id="scan_type"></select></div>
                            <div class="col-md-6 mb-3"><label for="ultravist_type" class="form-label">Ultravist Type:</label><select class="form-select" id="ultravist_type"><option>Ultravist 300</option><option>Ultravist 370</option></select></div>
                        </div>
                        <div class="text-center">
                            <button type="submit" class="btn btn-primary btn-lg">Calculate Protocol</button>
                        </div>
                    </form>
                </div>
            </div>
            <!-- RESULTS SECTIONS -->
            <div id="calculator-results-container">
                <div class="card mb-4 d-none" id="premed-card">
                    <div class="card-header bg-warning"><h5 class="mb-0"><i class="bi bi-shield-plus me-2"></i>IV Pre-medication Options</h5></div>
                    <div class="card-body">
                        <ul class="list-unstyled">
                            <li>- <strong>Option 1:</strong> Hydrocortisone 200mg IV q4-6h until scan + Diphenhydramine 50mg IV 1h prior.</li>
                            <li>- <strong>Option 2:</strong> Methylprednisolone 40mg IV q4h until scan + Diphenhydramine 50mg IV 1h prior.</li>
                            <li>- <strong>Option 3:</strong> Dexamethasone 7.5mg IV q4h until scan + Diphenhydramine 50mg IV 1h prior.</li>
                        </ul>
                    </div>
                </div>
                <div class="card mb-4 d-none" id="renal-card">
                    <div class="card-header bg-danger text-white"><h5 class="mb-0"><i class="bi bi-exclamation-triangle-fill me-2"></i>Renal Considerations</h5></div>
                    <div class="card-body"><p id="renal-message" class="mb-0" style="white-space: pre-wrap;"></p></div>
                </div>
                <div class="card mb-4 d-none" id="protocol-card">
                    <div class="card-header"><h5 class="mb-0"><i class="bi bi-clipboard2-pulse-fill me-2"></i>Calculated Protocol</h5></div>
                    <div class="card-body">
                        <div class="row">
                            <div class="col-sm-4"><p class="calculator-result-label">IV Dose:</p></div>
                            <div class="col-sm-8"><p class="calculator-result-value" id="iv-dose-result"></p></div>
                        </div>
                        <div class="row">
                            <div class="col-sm-4"><p class="calculator-result-label">Injection Rate:</p></div>
                            <div class="col-sm-8"><p class="calculator-result-value" id="rate-result"></p></div>
                        </div>
                        <div class="row d-none" id="oral-contrast-row">
                            <div class="col-sm-4"><p class="calculator-result-label">Oral Contrast:</p></div>
                            <div class="col-sm-8"><p id="oral-dose-result"></p></div>
                        </div>
                        <div class="row">
                            <div class="col-sm-4"><p class="calculator-result-label">Timing/Method:</p></div>
                            <div class="col-sm-8"><p id="timing-result" style="white-space: pre-wrap;"></p></div>
                        </div>
                    </div>
                </div>
                 <div class="card mb-4 d-none" id="iv-info-card">
                    <div class="card-header"><h5 class="mb-0"><i class="bi bi-bandaid-fill me-2"></i>IV Access Information</h5></div>
                    <div class="card-body">
                        <h6>Hierarchy of IV Access Sites:</h6>
                        <p id="iv-site-antecubital" class="ps-2">- <strong>1. Antecubital Fossa</strong> (Median Cubital, Basilic, Cephalic): Best for high flow rates.</p>
                        <p id="iv-site-forearm" class="ps-2">- <strong>2. Forearm Veins</strong> (Cephalic, Basilic): Good general-purpose sites.</p>
                        <p id="iv-site-hand" class="ps-2">- <strong>3. Dorsal Hand Veins:</strong> Use when other sites are unavailable.</p>
                        <h6 class="mt-3">Recommended Cannula Size:</h6>
                        <ul class="list-unstyled">
                            <li id="cannula-16G">- <span class="highlight-cannula" style="background-color:#808080;">16G</span>: For high-flow studies like angiograms.</li>
                            <li id="cannula-18G">- <span class="highlight-cannula" style="background-color:#009B77;">18G</span>: Excellent for most contrast studies.</li>
                            <li id="cannula-20G">- <span class="highlight-cannula" style="background-color:#F6546A;">20G</span>: Good for routine studies with adequate veins.</li>
                            <li id="cannula-22G">- <span class="highlight-cannula" style="background-color:#0077BE;">22G</span>: For smaller/fragile veins or slower rates.</li>
                            <li id="cannula-24G">- <span class="highlight-cannula" style="background-color:#FFC000;">24G</span>: Standard for pediatric patients.</li>
                        </ul>
                    </div>
                </div>
            </div>
            <!-- CT CONTRAST PROTOCOL CALCULATOR END -->
        </div>
        <div class="col-lg-5">
             <div class="card">
                <div class="card-header"><h5 class="mb-0">Resource Files & Consent Forms</h5></div>
                <div class="card-body">
                    {% if has_permission('manage_resources') %}
                    <div class="mb-4 p-3 border rounded bg-light">
                        <h6>Upload New Resource</h6>
                        <form action="{{ url_for('upload_resource') }}" method="post" enctype="multipart/form-data">
                            <div class="input-group"><input type="file" class="form-control" name="resource_file" required><button class="btn btn-success" type="submit">Upload</button></div>
                        </form>
                    </div>
                    {% endif %}
                    <h6>Available Documents:</h6>
                    <ul class="list-group" id="resource-files-list"></ul>
                </div>
            </div>
        </div>
    </div>
    <!-- BOTTOM ROW: QUIZ AND LEADERBOARD -->
    <div class="row mt-4">
        <div class="col-lg-7">
            <div class="card mb-4">
                <div class="card-header bg-primary text-white">
                    <h5 class="mb-0"><i class="bi bi-trophy-fill me-2"></i>Radiology Quiz Challenge</h5>
                </div>
                <div class="card-body">
                    <p class="card-text">Test your knowledge with our competitive, timed radiology quiz! Answer 15 questions as quickly and accurately as possible to earn points and climb the leaderboard.</p>
                    <p class="text-muted small">You can attempt the quiz once every 24 hours.</p>
                    <div id="quiz-controls"></div>
                </div>
            </div>
        </div>
        <div class="col-lg-5">
            <div class="card mb-4">
                <div class="card-header"><h5 class="mb-0"><i class="bi bi-bar-chart-line-fill me-2"></i>Leaderboard</h5></div>
                <div class="card-body">
                    <div class="table-responsive" style="max-height: 295px; overflow-y: auto;">
                        <table class="table table-striped table-sm leaderboard-table">
                            <thead><tr><th>Rank</th><th>Name</th><th>Score</th></tr></thead>
                            <tbody id="leaderboard-body"></tbody>
                        </table>
                    </div>
                </div>
            </div>
        </div>
    </div>
</div>

<div class="modal fade" id="quizModal" tabindex="-1" aria-labelledby="quizModalLabel" aria-hidden="true" data-bs-backdrop="static" data-bs-keyboard="false">
  <div class="modal-dialog modal-lg-quiz modal-dialog-centered">
    <div class="modal-content">
      <div class="modal-header">
        <h5 class="modal-title" id="quizModalLabel">Radiology Quiz</h5>
        <div class="ms-auto">
            <span id="quiz-timer" class="badge bg-primary fs-6 me-3">Time: 30</span>
            <span id="quiz-score" class="badge bg-success fs-6">Score: 0</span>
        </div>
      </div>
      <div class="modal-body">
        <div class="progress mb-3" style="height: 5px;"><div id="question-progress" class="progress-bar" role="progressbar" style="width: 0%;" aria-valuenow="0" aria-valuemin="0" aria-valuemax="15"></div></div>
        <h5 id="question-text" class="mb-4"></h5>
        <div id="options-container" class="d-grid gap-2"></div>
        <div id="feedback-container" class="mt-3 alert d-none"></div>
      </div>
      <div class="modal-footer"><button type="button" class="btn btn-primary" id="next-question-btn" disabled>Next Question</button></div>
    </div>
  </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script>
document.addEventListener('DOMContentLoaded', function() {
    // --- CT CONTRAST CALCULATOR LOGIC ---
    const contrastForm = document.getElementById('contrastCalculatorForm');
    const scanTypeSelect = document.getElementById('scan_type');
    
    const scanTypeOptions = [
        "CT Abdomen/Pelvis", "Pancreatic Protocol", "CT Enterography", "Triple Phase Urogram",
        "CT Chest/Abdo/Pelvis (CAP)", "Pulmonary Angiogram (CTPA)", 
        "CT Brain w/ Contrast", "Brain Venogram", "Split-Bolus Neck", "Lower Limb Angiogram"
    ];
    scanTypeOptions.forEach(opt => {
        const option = document.createElement('option');
        option.value = opt;
        option.textContent = opt;
        scanTypeSelect.appendChild(option);
    });

    contrastForm.addEventListener('submit', function(e) {
        e.preventDefault();
        calculateContrastProtocol();
    });

    const CANNULA_COLORS = {
        "16G": "#808080", "18G": "#009B77", "20G": "#F6546A",
        "22G": "#0077BE", "24G": "#FFC000"
    };

    function calculateContrastProtocol() {
        const errorBox = document.getElementById('calc-error-box');
        errorBox.classList.add('d-none');

        document.getElementById('premed-card').classList.add('d-none');
        document.getElementById('renal-card').classList.add('d-none');
        document.getElementById('protocol-card').classList.add('d-none');
        document.getElementById('iv-info-card').classList.add('d-none');
        
        try {
            const weight_lbs = parseFloat(document.getElementById('weight_lbs').value);
            const age = parseInt(document.getElementById('age').value);
            const creatinine_umol_l = parseFloat(document.getElementById('creatinine').value) || 0.0;
            
            if (isNaN(weight_lbs) || isNaN(age) || weight_lbs <= 0 || age <= 0) {
                throw new Error("Please enter valid, positive numbers for Weight and Age.");
            }

            const weight_kg = weight_lbs * 0.453592;
            const sex = document.getElementById('sex').value;
            const scan = document.getElementById('scan_type').value;
            const ultravist = document.getElementById('ultravist_type').value;
            const has_allergy = document.getElementById('allergy').value === "Yes";

            if (has_allergy) {
                document.getElementById('premed-card').classList.remove('d-none');
            }

            let iv_dose, injection_rate, oral_dose, timing, cannula;
            const is_pediatric = age <= 16;

            if (is_pediatric) {
                iv_dose = (age <= 2 ? 2.0 : 1.5) * weight_kg;
                iv_dose = Math.min(iv_dose, 100);
                if (ultravist === "Ultravist 370") iv_dose *= 0.8;
                
                injection_rate = weight_kg < 10 ? 1.0 : 1.5;
                
                const is_ultravist_370 = (ultravist === "Ultravist 370");
                let oral_volume, contrast_vol;
                if (weight_kg < 5) {
                    oral_volume = weight_kg * 15;
                    contrast_vol = is_ultravist_370 ? 2.0 : 2.5;
                    oral_dose = `Mix ${contrast_vol.toFixed(1)}ml ${ultravist} in ${oral_volume.toFixed(0)}ml water/juice. Give 30-60 min prior.`;
                } else if (weight_kg < 10) {
                    oral_volume = weight_kg * 10;
                    contrast_vol = is_ultravist_370 ? 4.0 : 5.0;
                    oral_dose = `Mix ${contrast_vol.toFixed(1)}ml ${ultravist} in ${oral_volume.toFixed(0)}ml water/juice. Give 30-60 min prior.`;
                } else if (weight_kg < 20) {
                    contrast_vol = is_ultravist_370 ? 4.0 : 5.0;
                    oral_dose = `Mix ${contrast_vol.toFixed(1)}ml ${ultravist} in 200ml water/juice. Give 60-90 min prior.`;
                } else {
                    contrast_vol = is_ultravist_370 ? 8.0 : 10.0;
                    oral_dose = `Mix ${contrast_vol.toFixed(1)}ml ${ultravist} in 300-400ml water/juice. Give 60-90 min prior.`;
                }
                const saline_flush_vol = Math.min(1.5 * weight_kg, 20);
                timing = `Scan at 60s post-injection. ${saline_flush_vol.toFixed(1)}ml Saline flush.`;
                cannula = ["24G", "22G"];
                
                document.getElementById('iv-dose-result').textContent = `${iv_dose.toFixed(1)} ml (PEDIATRIC PROTOCOL APPLIED)`;

            } else { // Adult Protocol
                const split_bolus_timing = "1. Bolus 1: Inject 60ml Contrast\\n2. Pause: Wait for 70 seconds\\n3. Bolus 2: Inject 40ml Contrast\\n4. Saline Flush: 40ml Saline\\n5. Scan: Acquire images 25s after starting Bolus 2";
                const oral_contrast_vol = ultravist === "Ultravist 370" ? 16 : 20;
                const adult_oral_string = `${oral_contrast_vol}ml ${ultravist} in 1000ml water. Give 500ml 1hr prior, 500ml 30min prior.`;

                const protocols = {
                    "CT Abdomen/Pelvis": {"iv": 1.25, "max_iv": 100, "oral": adult_oral_string, "timing": "Scan at 70s post-injection. 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "Pancreatic Protocol": {"iv": 1.5, "max_iv": 100, "oral": "1000ml Water. Give over 30-45 min prior.", "timing": "Dual Phase: Scan at 35-40s (late arterial) & 60-70s (portal venous). 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "CT Enterography": {"iv": 1.25, "max_iv": 100, "oral": "1350ml Volumen over 60 mins.", "timing": "Scan at 50s (enteric phase). 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "Triple Phase Urogram": {"iv": 1.0, "max_iv": 100, "oral": "N/A", "timing": "Multiphase: Non-con, 70s Nephrographic, 10-15min Delayed Excretory. 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "CT Chest/Abdo/Pelvis (CAP)": {"iv": 1.5, "max_iv": 100, "oral": adult_oral_string, "timing": "Scan at 70s post-injection. 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "Pulmonary Angiogram (CTPA)": {"iv_fixed": 80, "oral": "N/A", "timing": "Fixed dose for consistent enhancement. Rapid bolus. Use bolus tracking in main pulmonary artery. 50ml Saline flush.", "cannula": ["18G", "16G"]},
                    "CT Brain w/ Contrast": {"iv_fixed": 75, "oral": "N/A", "timing": "Scan 5 minutes post-injection for routine studies. 30ml Saline flush.", "cannula": ["22G", "20G"]},
                    "Brain Venogram": {"iv_fixed": 90, "oral": "N/A", "timing": "Scan at 55-60s (venous phase). 40ml Saline flush.", "cannula": ["20G", "18G"]},
                    "Split-Bolus Neck": {"iv_fixed": 100, "oral": "N/A", "timing": split_bolus_timing, "cannula": ["20G", "18G"]},
                    "Lower Limb Angiogram": {"iv_fixed": 100, "oral": "N/A", "timing": "High rate injection. Use bolus tracking over distal aorta. 50ml Saline flush.", "cannula": ["18G", "16G"]}
                };
                const protocol = protocols[scan];
                iv_dose = protocol.iv_fixed || Math.min(protocol.iv * weight_kg, protocol.max_iv);
                if (ultravist === "Ultravist 370") iv_dose *= 0.8;
                
                const base_rate = scan.includes("Angiogram") ? 5.0 : 3.5;
                injection_rate = age <= 65 ? base_rate : base_rate * 0.8;
                
                oral_dose = protocol.oral;
                timing = protocol.timing;
                cannula = protocol.cannula;

                if (creatinine_umol_l > 0) {
                    const creatinine_mg_dl = creatinine_umol_l / 88.4;
                    const kappa = sex === "Female" ? 0.7 : 0.9;
                    const alpha = sex === "Female" ? -0.241 : -0.302;
                    const sex_factor = sex === "Female" ? 1.012 : 1.0;
                    const egfr = 142 * (Math.min(creatinine_mg_dl / kappa, 1)**alpha) * (Math.max(creatinine_mg_dl / kappa, 1)**-1.200) * (0.9938**age) * sex_factor;

                    if (egfr < 60) {
                        document.getElementById('renal-card').classList.remove('d-none');
                        let msg = `eGFR ${egfr.toFixed(1)} ml/min/1.73m². MODERATE RISK.\\nRecommend pre- and post-CT hydration (e.g., 1L Normal Saline or oral fluids).`;
                        if (egfr < 30) {
                            msg = `eGFR ${egfr.toFixed(1)} ml/min/1.73m². SEVERE RISK.\\nContrast is relatively contraindicated. Discuss with Radiologist. Consider alternative imaging. If scan is necessary, consider post-CT dialysis.`;
                        }
                        document.getElementById('renal-message').textContent = msg;
                    }
                }
                document.getElementById('iv-dose-result').textContent = `${iv_dose.toFixed(1)} ml`;
            }

            document.getElementById('rate-result').textContent = `${injection_rate.toFixed(1)} ml/s`;
            const oralRow = document.getElementById('oral-contrast-row');
            if (oral_dose && oral_dose !== "N/A") {
                document.getElementById('oral-dose-result').textContent = oral_dose;
                oralRow.classList.remove('d-none');
            } else {
                oralRow.classList.add('d-none');
            }
            document.getElementById('timing-result').textContent = timing;

            resetHighlights();
            applyHighlights(cannula);

            document.getElementById('protocol-card').classList.remove('d-none');
            document.getElementById('iv-info-card').classList.remove('d-none');

        } catch (err) {
            errorBox.textContent = err.message;
            errorBox.classList.remove('d-none');
        }
    }

    function resetHighlights() {
        document.querySelectorAll('#iv-info-card li, #iv-info-card p').forEach(el => {
            el.classList.remove('iv-site-highlight');
            el.style.boxShadow = 'none';
        });
    }

    function applyHighlights(recommended_cannulas) {
        const primary_cannula = recommended_cannulas[0];
        const primary_el = document.getElementById(`cannula-${primary_cannula}`);
        if (primary_el) {
            const color = CANNULA_COLORS[primary_cannula] || '#000000';
            primary_el.style.boxShadow = `0 0 8px 2px ${color}`;
        }

        let site_el;
        if (["16G", "18G"].includes(primary_cannula)) {
            site_el = document.getElementById('iv-site-antecubital');
        } else if (["20G", "22G"].includes(primary_cannula)) {
            site_el = document.getElementById('iv-site-forearm');
        } else { // 24G
            site_el = document.getElementById('iv-site-hand');
        }
        if (site_el) {
            site_el.classList.add('iv-site-highlight');
        }
    }

    // --- QUIZ LOGIC ---
    const quizModalEl = document.getElementById('quizModal');
    const quizModal = new bootstrap.Modal(quizModalEl);
    let quizState = {};
    let timerInterval;

    async function fetchPageData() {
        try {
            const response = await fetch("{{ url_for('quiz_data_endpoint') }}");
            if (!response.ok) throw new Error('Failed to fetch page data');
            const data = await response.json();
            renderLeaderboard(data.leaderboard);
            renderResourceFiles(data.files);
            renderQuizControls(data.user_can_play, data.cooldown_message);
        } catch (error) {
            console.error("Error fetching page data:", error);
        }
    }
    
    function renderResourceFiles(files) {
        const listEl = document.getElementById('resource-files-list');
        if (!listEl) return;
        if (!files || files.length === 0) {
            listEl.innerHTML = '<li class="list-group-item text-muted">No resource files have been uploaded.</li>';
            return;
        }
        let filesHtml = '';
        const canManage = {{ has_permission('manage_resources')|tojson }};
        files.forEach(file => {
            filesHtml += `
            <li class="list-group-item d-flex justify-content-between align-items-center">
                <span><i class="bi bi-file-earmark-text me-2"></i>${file}</span>
                <div>
                    <a href="/download_resource/${encodeURIComponent(file)}" class="btn btn-sm btn-outline-primary">Download</a>
                    ${canManage ? `
                    <form action="/delete_resource/${encodeURIComponent(file)}" method="post" class="d-inline ms-1" onsubmit="return confirm('Are you sure?');">
                        <button type="submit" class="btn btn-sm btn-outline-danger">Delete</button>
                    </form>` : ''}
                </div>
            </li>`;
        });
        listEl.innerHTML = filesHtml;
    }

    function renderLeaderboard(leaderboard) {
        const leaderboardBody = document.getElementById('leaderboard-body');
        if (!leaderboardBody) return;
        if (!leaderboard || leaderboard.length === 0) {
            leaderboardBody.innerHTML = '<tr><td colspan="3" class="text-center text-muted">No scores yet. Be the first!</td></tr>';
            return;
        }
        let boardHtml = '';
        leaderboard.forEach((entry, index) => {
            boardHtml += `<tr><td>${index + 1}</td><td>${entry.full_name}</td><td>${entry.score}</td></tr>`;
        });
        leaderboardBody.innerHTML = boardHtml;
    }

    function renderQuizControls(canPlay, message) {
        const controlsContainer = document.getElementById('quiz-controls');
        if (!controlsContainer) return;
        if (canPlay) {
            controlsContainer.innerHTML = '<button id="start-quiz-btn" class="btn btn-primary btn-lg">Start Today\\'s Quiz</button>';
            document.getElementById('start-quiz-btn').addEventListener('click', startQuiz);
        } else {
            controlsContainer.innerHTML = `<div class="alert alert-info">${message}</div>`;
        }
    }

    async function startQuiz() {
        document.getElementById('start-quiz-btn').disabled = true;
        document.getElementById('start-quiz-btn').textContent = 'Loading...';
        try {
            const response = await fetch("{{ url_for('quiz_start') }}", { method: 'POST' });
            if (!response.ok) {
                 const error = await response.json();
                 alert(error.message || 'Could not start quiz.');
                 fetchPageData(); return;
            }
            const data = await response.json();
            quizState = { questions: data.questions, current_question_index: 0, score: 0, answers: [] };
            document.getElementById('quiz-score').textContent = 'Score: 0';
            quizModal.show();
            displayQuestion();
        } catch (err) { alert('An error occurred. Please try again later.'); fetchPageData(); }
    }
    
    function displayQuestion() {
        if (timerInterval) clearInterval(timerInterval);
        const questionData = quizState.questions[quizState.current_question_index];
        document.getElementById('question-text').textContent = questionData.question;
        const optionsContainer = document.getElementById('options-container');
        optionsContainer.innerHTML = '';
        questionData.options.forEach(option => {
            const button = document.createElement('button');
            button.className = 'btn btn-outline-secondary quiz-option-btn';
            button.textContent = option;
            button.addEventListener('click', () => selectAnswer(option, questionData.answer));
            optionsContainer.appendChild(button);
        });
        document.getElementById('feedback-container').classList.add('d-none');
        document.getElementById('next-question-btn').disabled = true;
        const progress = ((quizState.current_question_index) / quizState.questions.length) * 100;
        document.getElementById('question-progress').style.width = `${progress}%`;
        let timeLeft = 30;
        const timerEl = document.getElementById('quiz-timer');
        timerEl.textContent = `Time: ${timeLeft}`;
        timerEl.className = 'badge bg-primary fs-6 me-3';
        timerInterval = setInterval(() => {
            timeLeft--;
            timerEl.textContent = `Time: ${timeLeft}`;
            if (timeLeft <= 10) timerEl.className = 'badge bg-danger fs-6 me-3';
            if (timeLeft <= 0) { clearInterval(timerInterval); selectAnswer(null, questionData.answer); }
        }, 1000);
    }

    function selectAnswer(selectedOption, correctAnswer) {
        clearInterval(timerInterval);
        const feedbackContainer = document.getElementById('feedback-container');
        feedbackContainer.classList.remove('d-none');
        let points = 0;
        if (selectedOption === correctAnswer) {
            const timeLeft = parseInt(document.getElementById('quiz-timer').textContent.split(' ')[1]);
            points = 1000 + (timeLeft * 50);
            quizState.score += points;
            feedbackContainer.className = 'mt-3 alert alert-success';
            feedbackContainer.innerHTML = `<strong>Correct!</strong> +${points} points.`;
        } else {
            feedbackContainer.className = 'mt-3 alert alert-danger';
            feedbackContainer.innerHTML = `<strong>${selectedOption === null ? "Time's up!" : "Incorrect."}</strong> The correct answer was: <strong>${correctAnswer}</strong>`;
        }
        quizState.answers.push({ question: quizState.questions[quizState.current_question_index].question, answer: selectedOption, is_correct: selectedOption === correctAnswer });
        document.getElementById('quiz-score').textContent = `Score: ${quizState.score}`;
        document.querySelectorAll('.quiz-option-btn').forEach(btn => {
            btn.disabled = true;
            if (btn.textContent === correctAnswer) { btn.classList.remove('btn-outline-secondary'); btn.classList.add('btn-success'); }
        });
        const nextButton = document.getElementById('next-question-btn');
        nextButton.disabled = false;
        nextButton.textContent = quizState.current_question_index === quizState.questions.length - 1 ? 'Finish Quiz' : 'Next Question';
    }

    document.getElementById('next-question-btn').addEventListener('click', () => {
        if (quizState.current_question_index < quizState.questions.length - 1) {
            quizState.current_question_index++;
            displayQuestion();
        } else {
            finishQuiz();
        }
    });

    async function finishQuiz() {
      quizModal.hide();
      const response = await fetch("{{ url_for('quiz_submit') }}", {
        method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ answers: quizState.answers })
      });
      const result = await response.json();
      if (!response.ok) { alert(result.message || 'Could not submit quiz. Please try again.'); }
      else { alert(`Quiz Finished!\\n\\nYour final official score is: ${result.final_score}`); }
      fetchPageData();
    }
    
    fetchPageData();
});
</script>
</body></html>
"""



LOGIN_TEMPLATE = """
<!doctype html><html lang="en"><head><title>Login - RadTrac®</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><style>body{display:flex;align-items:center;justify-content:center;min-height:100vh;background-color:#f0f2f5}.login-card{width:100%;max-width:400px;padding:2rem;border-radius:.5rem;box-shadow:0 4px 12px rgba(0,0,0,.1)}.logo-title{font-size:1.6rem;font-weight:400;color:#fff;letter-spacing:-1px;text-align:center;margin-bottom:1.2rem;background:#003366;padding:.8rem;border-radius:.5rem}.banner-logo{display:block;margin:auto;height:45px}</style></head><body><div class="card login-card"><img src="/logo" alt="Logo" class="banner-logo mb-3"><div class="logo-title">RadTrac®</div>{% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, message in messages %}<div class="alert alert-{{ category }}">{{ message }}</div>{% endfor %}{% endif %}{% endwith %}<form method="post"><div class="mb-3"><label for="username" class="form-label">Username</label><input type="text" class="form-control" id="username" name="username" required autofocus></div><div class="mb-3"><label for="password" class="form-label">Password</label><input type="password" class="form-control" id="password" name="password" required></div><button type="submit" class="btn btn-primary w-100">Login</button></form></div></body></html>
"""
ADMIN_TEMPLATE = """
<!doctype html><html lang="en"><head><title>Admin Settings - CRH CT Tracker</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"></head><body><nav class="navbar navbar-expand-lg" style="background:#003366"><div class="container-fluid"><img src="/logo" alt="Logo" style="height:38px;margin-right:14px"><span class="navbar-brand text-white">Admin Settings</span><div class="d-flex ms-auto"><span class="navbar-text text-white me-3">User: {{ session.full_name }}</span><a href="/" class="btn btn-outline-light me-2">Main Tracker</a><a href="/mwl" class="btn btn-outline-info me-2">MWL Server</a><a href="/settings" class="btn btn-outline-info me-2">Settings</a><a href="/logout" class="btn btn-outline-warning">Logout</a></div></div></nav><div class="container mt-4">{% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, message in messages %}<div class="alert alert-{{ category }}">{{ message }}</div>{% endfor %}{% endif %}{% endwith %}<div class="card mb-4"><div class="card-header">Create New User</div><div class="card-body"><form method="post" action="{{ url_for('admin_add_user') }}"><div class="row g-3 align-items-end"><div class="col-md"><label class="form-label">Full Name</label><input type="text" class="form-control" name="full_name" placeholder="e.g. John Smith" required></div><div class="col-md"><label class="form-label">Username</label><input type="text" class="form-control" name="username" placeholder="e.g. jsmith" required></div><div class="col-md"><label class="form-label">Password</label><input type="password" class="form-control" name="password" placeholder="Password" required></div><div class="col-md"><label class="form-label">Role</label><select name="role" class="form-select"><option value="staff">Staff</option><option value="doctor">Doctor</option><option value="radiology_staff">Radiology Staff</option><option value="admin">Admin</option></select></div><div class="col-md-auto"><button type="submit" class="btn btn-primary">Add User</button></div></div></form></div></div><div class="card"><div class="card-header">Manage Users</div><div class="card-body"><ul class="list-group">{% for username, user_data in users.items() %}<li class="list-group-item d-flex justify-content-between align-items-center"><span><strong>{{ user_data.full_name }}</strong> ({{ username }}) <span class="badge bg-secondary ms-2">{{ user_data.role.replace('_', ' ')|title }}</span></span>{% if username.upper() != 'ADMIN' %}<form method="post" action="{{ url_for('admin_delete_user') }}" style="margin:0"><input type="hidden" name="username" value="{{ username }}"><button type="submit" class="btn btn-danger btn-sm" onclick="return confirm('Are you sure you want to delete this user?');">Delete</button></form>{% endif %}</li>{% endfor %}</ul></div></div></div></body></html>
"""
DASHBOARD_TEMPLATE = """
<!doctype html><html lang="en"><head><title>RadTrac® Dashboard</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/flatpickr/dist/flatpickr.min.css">
<style>
    body { background: #f6f8fa; height: 100vh; display: flex; flex-direction: column; }
    .navbar-custom { background: #003366 !important; }
    .navbar-custom .navbar-brand, .navbar-custom .navbar-text, .navbar-custom .nav-link { color: #fff !important; }
    .banner-logo { height: 45px; margin-right: 14px; }
    .btn-report-ready { background: #2fff35 !important; color: #222 !important; font-weight: 600; border: none; }
    .table th, .table td { vertical-align: middle; }
    #pacsStatus { font-weight: 500; }
    .status-ok { color: #198754; } .status-fail { color: #dc3545; }
    .header-container { position: sticky; top: 0; z-index: 1030; background-color: #f6f8fa; box-shadow: 0 2px 4px rgba(0,0,0,.1); }
    .table-container { flex-grow: 1; overflow-y: auto; }
    .table thead th { background: #003366; color: #fff; white-space: nowrap; position: sticky; top: 0; z-index: 10; }
    .report-actions form { display: inline-block; margin-left: 5px; }
    .report-options { display: flex; flex-wrap: wrap; align-items: center; gap: 5px; }
</style>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/flatpickr"></script>
</head>
<body>
<div class="container mt-3">
  {% with messages = get_flashed_messages(with_categories=true) %}
    {% if messages %}
      {% for category, message in messages %}
        <div class="alert alert-{{ category }} alert-dismissible fade show" role="alert">
          {{ message }}
          <button type="button" class="btn-close" data-bs-dismiss="alert" aria-label="Close"></button>
        </div>
      {% endfor %}
    {% endif %}
  {% endwith %}
</div>
    <div class="header-container">
        <nav class="navbar navbar-expand-lg navbar-custom">
            <div class="container-fluid">
                <img src="/logo" alt="Logo" class="banner-logo">
                <span class="navbar-brand">{{ config.INSTITUTION_NAME }} - RadTrac®</span>
                <div class="d-flex align-items-center ms-auto">
                    <span class="navbar-text me-3">User: {{ session.full_name }}</span>
                    {% if has_permission('view_approval_tracker') %}<a href="/approval" class="btn btn-info me-2">Approval Tracker</a>{% endif %}
                    {% if has_permission('view_us_approval_tracker') %}<a href="/us_approval" class="btn btn-info me-2">US Tracker</a>{% endif %}
                    {% if has_permission('view_mwl_server') %}<a href="/mwl" class="btn btn-info me-2">MWL Server</a>{% endif %}
                    {% if has_permission('view_admin_page') %}<a href="/admin" class="btn btn-secondary me-2">Admin</a>{% endif %}
                    {% if has_permission('view_settings_page') %}<a href="/settings" class="btn btn-info me-2">Settings</a>{% endif %}
                    {% if has_permission('view_resources_page') %}<a href="/resources" class="btn btn-outline-success me-2">Resources</a>{% endif %}
                    <a href="/logout" class="btn btn-warning">Logout</a>
                </div>
            </div>
        </nav>
        <div class="container-fluid py-3">
            <div class="row gy-2 gx-3 align-items-center">
                <div class="col-auto"><select class="form-select" id="search_field"><option value="patient_name">Patient Name</option><option value="patient_id">Patient ID</option><option value="accession">Accession #</option><option value="referred_from">Referred From</option></select></div>
                <div class="col-auto"><input class="form-control" style="min-width:300px" type="search" id="search_query" placeholder="Search..."></div>
                <div class="col-auto"><input class="form-control" type="text" id="dateFilter" placeholder="Filter by Study Date"></div>
                <div class="col-auto ms-auto text-muted small text-end"><strong>PACS Status:</strong> <span id="pacsStatus">Initializing...</span></div>
            </div>
        </div>
    </div>
    <div class="container-fluid mt-2 table-container">
        <div class="table-responsive">
            <table class="table table-striped table-hover table-bordered shadow-sm" id="ctTable">
                <thead><tr><th>Study Date/Time</th><th>Modality</th><th>Patient Name</th><th>Patient ID</th><th>Accession</th><th>Referred From</th><th>Study Description</th><th>Generate Report</th><th>Manage Report</th><th>Download</th></tr></thead>
                <tbody id="ctTableBody"><tr><td colspan="10" class="text-center p-5"><div class="spinner-border" role="status"><span class="visually-hidden">Loading...</span></div><p class="mt-2">Fetching live data from PACS...</p></td></tr></tbody>
            </table>
        </div>
    </div>
    <footer class="container-fluid text-center mt-auto mb-2 small text-muted">&copy; {{ year }} {{ config.INSTITUTION_NAME }}</footer>
    
<script>
    document.addEventListener('DOMContentLoaded', () => {
        const searchField = document.getElementById('search_field');
        const searchQuery = document.getElementById('search_query');
        const dateFilter = document.getElementById('dateFilter');
        const tableBody = document.getElementById('ctTableBody');
        const pacsStatusEl = document.getElementById('pacsStatus');
        let allEntries = [];

        function renderTable(entries) {
            if (!entries || entries.length === 0) {
                tableBody.innerHTML = '<tr><td colspan="10" class="text-center p-4">No studies found matching your criteria.</td></tr>';
                return;
            }
            entries.sort((a,b) => (b.log_timestamp || '').localeCompare(a.log_timestamp || ''));
            let rowsHtml = '';
            for (const row of entries) {
                
                // --- Create a colored badge for the modality ---
                const modality = (row.modality || 'N/A').trim().toUpperCase();
                let modalityBadge = `<span class="badge bg-secondary">${modality}</span>`; // Default
                
                if (modality === 'CT') {
                    modalityBadge = `<span class="badge bg-primary">${modality}</span>`;
                } else if (modality === 'US') {
                    // Custom light pink badge for Ultrasound
                    modalityBadge = `<span class="badge" style="background-color: #FFEBF0; color: #d1386b; border: 1px solid #d1386b;">${modality}</span>`;
                } else if (modality === 'DX') {
                    // Green badge for DX
                    modalityBadge = `<span class="badge bg-success">${modality}</span>`;
                } else if (modality === 'CR') {
                    modalityBadge = `<span class="badge bg-light text-dark">${modality}</span>`;
                } else if (modality === 'MG') {
                    modalityBadge = `<span class="badge" style="background-color: #ff69b4; color: white;">${modality}</span>`;
                }

                // --- Generate Report Cell ---
                let generateCell = '';
                if (!row.report_path && {{ has_permission('upload_reports')|tojson }}) {
                    const queryParams = new URLSearchParams({
                        accession: row.accession
                    }).toString();
                    generateCell = `<a href="/generate_report?${queryParams}" class="btn btn-sm btn-info" target="_blank">Generate Blank</a>`;
                }

                // --- Manage Report Cell ---
                let manageCell = '';
                if (row.report_path) {
                  manageCell = `
                    <div class="report-options">
                      <a href="/report?path=${encodeURIComponent(row.report_path)}" class="btn btn-report-ready btn-sm" target="_blank">View Report</a>`;
                  if ({{ has_permission('delete_reports')|tojson }}) {
                    manageCell += `<form method="POST" action="/delete_report" onsubmit="return confirm('Are you sure?');" class="d-inline ms-1"><input type="hidden" name="accession" value="${row.accession}"><input type="hidden" name="report_path" value="${row.report_path}"><input type="hidden" name="patient_name" value="${row.patient_name}"><button type="submit" class="btn btn-danger btn-sm">Delete</button></form>`;
                  }
                  manageCell += `</div>`;
                } else {
                  if ({{ has_permission('upload_reports')|tojson }}) {
                    manageCell = `<form method="POST" action="/upload_report" enctype="multipart/form-data"><input type="hidden" name="accession" value="${row.accession}"><input type="hidden" name="patient_name" value="${row.patient_name}"><input type="hidden" name="patient_id" value="${row.patient_id}"><div class="input-group"><input type="file" name="report_file" class="form-control form-control-sm" accept=".pdf" required><button type="submit" class="btn btn-sm btn-primary">Upload</button></div></form>`;
                  } else {
                    manageCell = `<span class="badge bg-secondary">Report Pending</span>`;
                  }
                }
                
                // --- Download Images Cell ---
                let downloadBtn = '';
                if ({{ has_permission('download_images')|tojson }}) {
                    downloadBtn = `<button class="btn btn-outline-dark btn-sm download-btn" data-accession="${row.accession}" title="Download all images as a ZIP file"><span class="btn-text">Download</span><span class="spinner-border spinner-border-sm d-none" role="status" aria-hidden="true"></span></button>`;
                }

                // --- Assemble the final row (no special class needed anymore) ---
                rowsHtml += `<tr class="report-actions">
                    <td>${row.study_date_fmt} ${row.log_time||'N/A'}</td>
                    <td>${modalityBadge}</td>
                    <td>${row.patient_name}</td>
                    <td>${row.patient_id}</td>
                    <td>${row.accession}</td>
                    <td>${row.referred_from}</td>
                    <td>${row.study_desc}</td>
                    <td class="text-center">${generateCell}</td>
                    <td>${manageCell}</td>
                    <td class="text-center">${downloadBtn}</td>
                </tr>`;
            }
            tableBody.innerHTML = rowsHtml;
        }

        function filterAndRender() {
            let filtered = [...allEntries];
            const query = searchQuery.value.toLowerCase();
            const field = searchField.value;
            const selectedDates = dateFilter._flatpickr.selectedDates;
            if (query) { filtered = filtered.filter(e => e[field] && String(e[field]).toLowerCase().includes(query)); }
            if (selectedDates.length === 2) {
                const startStr = selectedDates[0].toISOString().slice(0,10).replace(/-/g,'');
                const endStr = selectedDates[1].toISOString().slice(0,10).replace(/-/g,'');
                filtered = filtered.filter(e => e.study_date && e.study_date >= startStr && e.study_date <= endStr);
            }
            renderTable(filtered);
        }

        async function fetchData() {
            try {
                const response = await fetch('/api/studies');
                if (!response.ok) { pacsStatusEl.textContent = 'Error fetching data.'; pacsStatusEl.className = 'status-fail'; return; }
                const result = await response.json();
                pacsStatusEl.textContent = result.status;
                pacsStatusEl.className = result.success ? 'status-ok' : 'status-fail';
                if (result.data && result.data.length > 0) { allEntries = result.data; filterAndRender(); }
                else if (!result.success) { tableBody.innerHTML = '<tr><td colspan="10" class="text-center p-4 text-danger">Could not retrieve data from PACS and no cached results are available.</td></tr>'; }
                else { renderTable([]); }
            } catch (error) { console.error("Fetch error:", error); pacsStatusEl.textContent = 'Network error.'; pacsStatusEl.className = 'status-fail'; }
        }

        flatpickr(dateFilter, { mode: "range", altInput: true, altFormat: "M j, Y", dateFormat: "Y-m-d", onClose: filterAndRender });
        searchQuery.addEventListener('input', filterAndRender);
        searchField.addEventListener('change', filterAndRender);
        
        document.addEventListener('click', function(e) {
            const downloadButton = e.target.closest('.download-btn');
            if (!downloadButton) return;
            e.preventDefault();
            if (downloadButton.disabled) return;
            downloadButton.disabled = true;

            const btnText = downloadButton.querySelector('.btn-text');
            const spinner = downloadButton.querySelector('.spinner-border');
            btnText.textContent = 'Preparing...';
            spinner.classList.remove('d-none');

            const accession = downloadButton.dataset.accession;

            fetch(`/download/start/${encodeURIComponent(accession)}`)
                .then(response => response.json())
                .then(data => {
                    const taskId = data.task_id;
                    if (!taskId) throw new Error('Failed to start download task.');

                    const interval = setInterval(() => {
                        fetch(`/download/status/${taskId}`)
                            .then(res => res.json())
                            .then(statusData => {
                                if (statusData.state === 'SUCCESS') {
                                    clearInterval(interval);
                                    btnText.textContent = 'Ready!';
                                    spinner.classList.add('d-none');
                                    downloadButton.disabled = false;
                                    window.location.href = `/download/get/${statusData.filename}`;
                                } else if (statusData.state === 'FAILURE') {
                                    clearInterval(interval);
                                    btnText.textContent = 'Error';
                                    spinner.classList.add('d-none');
                                    downloadButton.disabled = false;
                                    alert(`Download failed: ${statusData.status}`);
                                } else {
                                    btnText.textContent = statusData.status || 'Processing...';
                                }
                            });
                    }, 2500);
                })
                .catch(error => {
                    console.error('Download initiation error:', error);
                    alert('Could not start the download. Check console for details.');
                    btnText.textContent = 'Download';
                    spinner.classList.add('d-none');
                    downloadButton.disabled = false;
                });
        });

        fetchData();
        setInterval(fetchData, 30000);
    });
</script>
</body></html>
"""


APPROVAL_TEMPLATE = """
<!doctype html><html lang="en"><head><title>Approval Tracker - RadTrac®</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/flatpickr/dist/flatpickr.min.css">
<style>
    body { display: flex; flex-direction: column; height: 100vh; background-color: #f8f9fa; }
    .header-container { position: sticky; top: 0; z-index: 1030; background-color: #f8f9fa; box-shadow: 0 2px 4px rgba(0,0,0,.1); }
    .navbar-custom { background: #003366 !important; }
    .navbar-custom .navbar-brand, .navbar-custom .navbar-text { color: #fff !important; }
    .banner-logo { height: 45px; margin-right: 14px; }
    .table-container { flex-grow: 1; overflow-y: auto; }
    .table thead th { background: #003366; color: #fff; white-space: nowrap; position: sticky; top: 0; z-index: 10; }
    .table th, .table td { vertical-align: middle; }
    .activity-log { max-height: 180px; overflow-y: auto; padding: 5px; border: 1px solid #eee; border-radius: 4px; background: #fff; }
    .note-entry { line-height: 1.2; }
    .subscriber-list { list-style-type: none; padding-left: 0; margin-bottom: 0.5rem; max-height: 80px; overflow-y: auto;}
    .subscriber-list li { display: flex; justify-content: space-between; align-items: center; padding: 0.2rem 0; font-size: 0.9em; }
    .toast-container { z-index: 1090; }
    .filter-btn.active { font-weight: bold; }
    #request-count { font-weight: 500; }
    /* Custom style for the info needed status */
    .badge-info-needed { background-color: #FFE0B2; color: #8C530C; }
</style>
</head>
<body>
<div class="header-container">
    <nav class="navbar navbar-expand-lg navbar-custom">
        <div class="container-fluid">
            <img src="/logo" alt="Logo" class="banner-logo"><span class="navbar-brand">CRH CT Approval Tracker</span>
            <div class="d-flex align-items-center ms-auto">
                <span class="navbar-text me-3">User: {{ session.full_name }}</span>
                {% if has_permission('view_dashboard') %}<a href="/" class="btn btn-outline-light me-2">Main Tracker</a>{% endif %}
                {% if has_permission('view_us_approval_tracker') %}<a href="/us_approval" class="btn btn-outline-info me-2">US Tracker</a>{% endif %}
                {% if has_permission('view_mwl_server') %}<a href="/mwl" class="btn btn-outline-info me-2">MWL Server</a>{% endif %}
                {% if has_permission('view_admin_page') %}<a href="/admin" class="btn btn-secondary me-2">Admin</a>{% endif %}
                {% if has_permission('view_settings_page') %}<a href="/settings" class="btn btn-info me-2">Settings</a>{% endif %}
                {% if has_permission('view_resources_page') %}<a href="/resources" class="btn btn-outline-success me-2">Resources</a>{% endif %}
                <a href="/logout" class="btn btn-warning">Logout</a>
            </div>
        </div>
    </nav>
    <div class="container-fluid py-3">
        <div class="row g-2 align-items-center">
            <div class="col-md-4">
                <input type="search" id="approvalSearch" class="form-control" placeholder="Search by Patient Name, ID, Ward...">
            </div>
            <div class="col-md-8 d-flex justify-content-start align-items-center">
                <div class="btn-group me-3" role="group" aria-label="Status Filters">
                  <button type="button" class="btn btn-primary filter-btn active" data-filter="All">All</button>
                  <button type="button" class="btn btn-outline-primary filter-btn" data-filter="Pending Review">Pending</button>
                  <button type="button" class="btn btn-outline-primary filter-btn" data-filter="Info Needed">Info Needed</button>
                  <button type="button" class="btn btn-outline-primary filter-btn" data-filter="Approved">Approved</button>
                  <button type="button" class="btn btn-outline-primary filter-btn" data-filter="Completed">Completed</button>
                </div>
                <div class="me-3">
                    <select class="form-select" id="sort-select">
                        <option value="activity-desc">Sort by Last Activity (Newest)</option>
                        <option value="activity-asc">Sort by Last Activity (Oldest)</option>
                        <option value="name-asc">Sort by Patient Name (A-Z)</option>
                    </select>
                </div>
                <div id="request-count" class="text-muted"></div>
            </div>
        </div>
    </div>
</div>
<div class="container-fluid table-container">
    <div id="toast-container" class="toast-container position-fixed top-0 end-0 p-3"></div>
    <div class="table-responsive"><table class="table table-bordered table-hover align-middle">
            <thead><tr>
                <th style="width: 10%;">Notifications</th>
                <th style="width: 18%;">Patient Info</th>
                <th>Card</th>
                {% if has_permission('preregister_patient') %}
                <th class="text-center" style="width: 5%; white-space: nowrap;">Pre-Register</th>
                {% endif %}
                <th>Status</th>
                <th style="width: 25%;">Activity Log</th>
                <th style="width: 15%;">Actions</th>
            </tr></thead>
            <tbody id="approvalTableBody">
                <tr><td colspan="7" class="text-center p-5"><div class="spinner-border" role="status"><span class="visually-hidden">Loading...</span></div><p class="mt-2">Loading Approval Requests...</p></td></tr>
            </tbody></table></div>
</div>

<div class="modal fade" id="preregisterModal" tabindex="-1" aria-labelledby="preregisterModalLabel" aria-hidden="true">
  <div class="modal-dialog">
    <div class="modal-content">
      <div class="modal-header">
        <h5 class="modal-title" id="preregisterModalLabel">Pre-register Patient for MWL</h5>
        <button type="button" class="btn-close" data-bs-dismiss="modal" aria-label="Close"></button>
      </div>
      <form id="preregisterForm">
        <div class="modal-body">
            <p>Add patient to 'Expected Patients' list.</p>
            <input type="hidden" id="prereg_filename" name="filename">
            <div class="row">
                <div class="col-md-8 mb-3"><label for="prereg_patient_name" class="form-label">Patient Name</label><input type="text" class="form-control" id="prereg_patient_name" name="patient_name"></div>
                <div class="col-md-4 mb-3"><label for="prereg_sex" class="form-label">Sex</label><select class="form-select" id="prereg_sex" name="sex"><option value="" selected disabled>Select...</option><option value="M">Male</option><option value="F">Female</option><option value="O">Other</option></select></div>
            </div>
            <div class="row">
                <div class="col-md-6 mb-3"><label for="prereg_patient_id" class="form-label">Patient ID</label><input type="text" class="form-control" id="prereg_patient_id" name="patient_id"></div>
                <div class="col-md-6 mb-3"><label for="prereg_dob" class="form-label">Date of Birth</label><input type="text" class="form-control" id="prereg_dob" name="dob" placeholder="DD/MM/YYYY"></div>
            </div>
            <div class="row">
                <div class="col-md-7 mb-3"><label for="prereg_study_description" class="form-label">Study Description</label><input type="text" class="form-control" id="prereg_study_description" name="study_description"></div>
                <div class="col-md-5 mb-3"><label for="prereg_requesting_physician" class="form-label">Requesting Physician</label><input type="text" class="form-control" id="prereg_requesting_physician" name="requesting_physician"></div>
            </div>
            <div class="row">
                <div class="col-md-6 mb-3"><label for="prereg_modality" class="form-label">Modality</label><input type="text" class="form-control" id="prereg_modality" name="modality" readonly></div>
                <div class="col-md-6 mb-3"><label for="prereg_referred_from" class="form-label">Referred From</label><input type="text" class="form-control" id="prereg_referred_from" name="referred_from"></div>
            </div>
            <div class="mb-3"><label for="prereg_scheduled_datetime" class="form-label">Scheduled Date & Time</label><input type="text" class="form-control" id="prereg_scheduled_datetime" name="scheduled_datetime" required></div>
        </div>
        <div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button><button type="submit" class="btn btn-primary">Save to Expected List</button></div>
      </form>
    </div>
  </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/flatpickr"></script>
<script>
document.addEventListener('DOMContentLoaded', () => {
    // --- Performance Debounce Function ---
    function debounce(func, delay) {
        let timeout;
        return function(...args) {
            const context = this;
            clearTimeout(timeout);
            timeout = setTimeout(() => func.apply(context, args), delay);
        };
    }

    // --- Element Selectors ---
    const searchInput = document.getElementById('approvalSearch');
    const tableBody = document.getElementById('approvalTableBody');
    const filterButtons = document.querySelectorAll('.filter-btn');
    const sortSelect = document.getElementById('sort-select');
    const requestCountEl = document.getElementById('request-count');
    
    // --- State Variables ---
    let allRows = []; // Will be populated after fetch
    let allRequestsData = [];
    let activeStatusFilter = 'All';
    let currentSort = 'activity-desc';
    let userPermissions = [];
    let currentUser = '';

    // --- Client-Side Template Function ---
    function createRowHtml(req, allUsers) {
        const pInfo = req.parsed_info || {};
        const subscribers = req.subscribers || [];
        const notes = req.notes || [];

        const searchText = `${pInfo.patient_name || ''} ${pInfo.patient_id_part || ''} ${pInfo.ward_part || ''}`.toLowerCase();
        
        // --- Notifications Column ---
        let notificationsHtml = `<span class="text-muted small">No subscription options for your role.</span>`;
        if (userPermissions.includes('manage_all_subscriptions')) {
            const subsHtml = subscribers.length > 0 ? subscribers.map(sub => `
                <li>
                    <span>${allUsers[sub] ? allUsers[sub].full_name : sub}</span>
                    <form method="POST" action="/approval/unsubscribe_user" class="d-inline">
                        <input type="hidden" name="filename" value="${req.filename}">
                        <input type="hidden" name="username_to_unsubscribe" value="${sub}">
                        <button type="submit" class="btn btn-danger btn-sm p-0 px-1" title="Remove">&times;</button>
                    </form>
                </li>`).join('') : '<li class="text-muted small">No subscribers.</li>';

            const userOptions = Object.keys(allUsers).map(username => {
                if (!subscribers.includes(username) && allUsers[username] && allUsers[username].email) {
                    return `<option value="${username}">${allUsers[username].full_name}</option>`;
                }
                return '';
            }).join('');

            notificationsHtml = `
                <div class="mb-2"><ul class="subscriber-list">${subsHtml}</ul></div>
                <form method="POST" action="/approval/subscribe_user">
                    <input type="hidden" name="filename" value="${req.filename}">
                    <div class="input-group">
                        <select name="username_to_subscribe" class="form-select form-select-sm">
                            <option value="">-- Select User --</option>${userOptions}
                        </select>
                        <button type="submit" class="btn btn-info btn-sm">Add</button>
                    </div>
                </form>`;
        } else if (userPermissions.includes('manage_own_subscriptions')) {
            const isSubscribed = subscribers.includes(currentUser);
            notificationsHtml = `
                <form method="POST" action="/approval/toggle_approval_subscription" class="d-grid">
                    <input type="hidden" name="filename" value="${req.filename}">
                    <button type="submit" class="btn btn-sm ${isSubscribed ? 'btn-secondary' : 'btn-info'}">
                        ${isSubscribed ? 'Stop Notifying Me' : 'Notify Me'}
                    </button>
                </form>`;
        }

        // --- Patient Info Column ---
        const warningsHtml = (req.warnings || []).map(w => `<div class="alert alert-warning p-1 small mb-1" role="alert"><i class="bi bi-exclamation-triangle-fill me-1"></i>${w}</div>`).join('');
        const patientInfoHtml = `
            <strong>${pInfo.patient_name || 'N/A'}</strong><br>
            <small class="text-muted">
                ID: ${pInfo.patient_id_part || 'N/A'}
                ${pInfo.ward_part ? `| Ward: ${pInfo.ward_part}` : ''}
                | Mod: ${pInfo.modality || 'N/A'}
            </small>
            ${warningsHtml ? `<div class="mt-2">${warningsHtml}</div>` : ''}`;

        // --- Status Badge ---
        let statusBadge;
        if (req.status === 'Completed') {
            statusBadge = `<span class="badge bg-success fs-6 fw-bold">Completed</span>`;
        } else if ((req.status || '').includes('Approved')) {
            statusBadge = `<span class="badge bg-info fs-6 fw-bold">${req.status}</span>`;
        } else if (req.status === 'Info Needed') {
            statusBadge = `<span class="badge badge-info-needed fs-6 fw-bold">Info Needed</span>`;
        } else if (req.status === 'Rejected' || req.status === 'Scan Failed') {
            statusBadge = `<span class="badge bg-danger fs-6 fw-bold">${req.status}</span>`;
        } else {
            statusBadge = `<span class="badge bg-warning text-dark fs-6">${req.status || 'N/A'}</span>`;
        }
        
        // --- Activity Log ---
        const notesHtml = [...notes].reverse().map(note => {
            const noteDate = new Date(note.on).toLocaleString('en-US', { month: 'short', day: 'numeric', hour: '2-digit', minute: '2-digit', hour12: true });
            return `<div class="note-entry mb-2"><div>${note.text}</div><div class="text-muted" style="font-size: 0.9em;"><em>- ${note.by} on ${noteDate}</em></div></div>`;
        }).join('');
        const creationDate = new Date(req.creation_time).toLocaleString('en-US', { month: 'short', day: 'numeric', hour: '2-digit', minute: '2-digit', hour12: true });
        const activityLogHtml = `<div class="small activity-log">${notesHtml}<div class="note-entry mb-2"><div>Card Uploaded</div><div class="text-muted" style="font-size: 0.9em;"><em>- System on ${creationDate}</em></div></div></div>`;

        // --- Actions Column ---
        let actionsHtml = '';
        if (userPermissions.includes('manage_approvals')) {
            actionsHtml += `
                <form method="POST" action="/approval/update_status" class="mb-2">
                    <input type="hidden" name="filename" value="${req.filename}">
                    <div class="input-group">
                        <select name="new_status" class="form-select form-select-sm">
                            <option>-- Set Status --</option><option value="Pending Review">Pending</option><option value="Info Needed">Info Needed</option><option value="Approved">Approved</option><option value="Approved Send ASAP">Approved ASAP</option><option value="Approved Wait for Time">Approved Wait</option><option value="Scan Failed">Scan Failed</option><option value="Completed">Completed</option><option value="Rejected">Reject</option>
                        </select>
                        <button type="submit" class="btn btn-secondary btn-sm">Set</button>
                    </div>
                </form>`;
        }
        if (userPermissions.includes('add_approval_notes')) {
            actionsHtml += `
                <form method="POST" action="/approval/add_note" class="mb-2">
                    <input type="hidden" name="filename" value="${req.filename}">
                    <textarea name="note_text" class="form-control form-control-sm" rows="1" placeholder="Add a brief note..."></textarea>
                    <div class="d-grid"><button type="submit" class="btn btn-primary btn-sm mt-1">Add Note</button></div>
                </form>`;
        }
        
        // --- Pre-Register Button ---
        const preRegButton = userPermissions.includes('preregister_patient') ? `<td class="text-center"><button class="btn btn-sm btn-outline-success preregister-btn">Pre-register</button></td>` : '';

        // --- Assemble Row ---
        return `
            <tr class="approval-row"
                data-search-text="${searchText}"
                data-status="${req.status || ''}"
                data-activity-time="${req.last_activity_on || ''}"
                data-patient-name="${(pInfo.patient_name || '').toLowerCase()}"
                data-filename="${req.filename}">
                <td>${notificationsHtml}</td>
                <td>${patientInfoHtml}</td>
                <td class="text-center"><a href="/request_card?path=${encodeURIComponent(req.file_path_safe)}" class="btn btn-outline-primary btn-sm" target="_blank">View Card</a></td>
                ${preRegButton}
                <td>${statusBadge}</td>
                <td>${activityLogHtml}</td>
                <td>${actionsHtml}</td>
            </tr>`;
    }

    // --- Main Filter and Sort Function ---
    function filterAndSort() {
        const searchQuery = searchInput.value.toLowerCase().trim();
        
        let visibleRows = allRows.filter(row => {
            const searchText = row.dataset.searchText || '';
            const rowStatus = row.dataset.status || '';
            const matchesSearch = searchText.includes(searchQuery);
            
            let matchesStatus = false;
            if (activeStatusFilter === 'All') {
                matchesStatus = true;
            } else if (activeStatusFilter === 'Approved') {
                matchesStatus = rowStatus.startsWith('Approved');
            } else {
                matchesStatus = rowStatus === activeStatusFilter;
            }
            return matchesSearch && matchesStatus;
        });

        visibleRows.sort((a, b) => {
            const nameA = a.dataset.patientName || '';
            const nameB = b.dataset.patientName || '';
            const timeA = a.dataset.activityTime || '';
            const timeB = b.dataset.activityTime || '';

            switch (currentSort) {
                case 'name-asc': return nameA.localeCompare(nameB);
                case 'activity-asc': return timeA.localeCompare(timeB);
                case 'activity-desc':
                default: return timeB.localeCompare(timeA);
            }
        });

        allRows.forEach(row => row.style.display = 'none');
        const fragment = document.createDocumentFragment();
        visibleRows.forEach(row => {
            row.style.display = '';
            fragment.appendChild(row);
        });
        tableBody.appendChild(fragment);

        requestCountEl.textContent = `Showing ${visibleRows.length} of ${allRows.length} requests.`;
        const noResultsRow = document.getElementById('no-results-row');
        if (noResultsRow) noResultsRow.style.display = visibleRows.length === 0 ? '' : 'none';
    }
    
    // --- Initial Data Load ---
    async function initialLoad() {
        try {
            const response = await fetch('/api/approvals');
            if (!response.ok) throw new Error(`HTTP error! status: ${response.status}`);
            const data = await response.json();

            allRequestsData = data.requests || [];
            userPermissions = data.permissions || [];
            currentUser = data.current_user || '';
            const allUsers = data.all_users || {};

            let tableHtml = '';
            if (allRequestsData.length > 0) {
                allRequestsData.forEach(req => {
                    tableHtml += createRowHtml(req, allUsers);
                });
            } else {
                tableHtml = `<tr><td colspan="7" class="text-center p-4">No active approval requests found.</td></tr>`;
            }
            tableHtml += `<tr id="no-results-row" style="display: none;"><td colspan="7" class="text-center p-4">No requests match the current filters.</td></tr>`;
            
            tableBody.innerHTML = tableHtml;
            allRows = Array.from(tableBody.querySelectorAll('.approval-row'));
            filterAndSort();

        } catch (error) {
            console.error("Failed to load approval data:", error);
            tableBody.innerHTML = `<tr><td colspan="7" class="text-center p-4 text-danger">Failed to load data. Please try refreshing the page.</td></tr>`;
        }
    }

    // --- Event Listeners ---
    searchInput.addEventListener('input', debounce(filterAndSort, 300));
    sortSelect.addEventListener('change', (e) => {
        currentSort = e.target.value;
        filterAndSort();
    });

    filterButtons.forEach(button => {
        button.addEventListener('click', () => {
            filterButtons.forEach(btn => {
                btn.classList.remove('active', 'btn-primary');
                btn.classList.add('btn-outline-primary');
            });
            button.classList.add('active', 'btn-primary');
            button.classList.remove('btn-outline-primary');
            activeStatusFilter = button.dataset.filter;
            filterAndSort();
        });
    });

    // --- Pre-registration Modal Logic ---
    const preregisterModalEl = document.getElementById('preregisterModal');
    if (preregisterModalEl) {
        const preregisterModal = new bootstrap.Modal(preregisterModalEl);
        const preregisterForm = document.getElementById('preregisterForm');

        flatpickr("#prereg_scheduled_datetime", {
            enableTime: true,
            dateFormat: "Y-m-d H:i",
            defaultDate: new Date()
        });
        
        tableBody.addEventListener('click', event => {
            if (!event.target.classList.contains('preregister-btn')) return;

            const row = event.target.closest('.approval-row');
            const filename = row.dataset.filename;
            const reqData = allRequestsData.find(r => r.filename === filename);
            
            if (!reqData || !reqData.parsed_info) {
                showToast('Error', 'Could not find data for this request.', true);
                return;
            }
            
            const info = reqData.parsed_info;
            document.getElementById('prereg_filename').value = filename;
            document.getElementById('prereg_patient_name').value = info.patient_name || '';
            document.getElementById('prereg_patient_id').value = info.patient_id_part || '';
            document.getElementById('prereg_modality').value = info.modality || '';
            document.getElementById('prereg_referred_from').value = info.ward_part || '';
            preregisterModal.show();
        });

        preregisterForm.addEventListener('submit', async (e) => {
            e.preventDefault();
            const formData = new FormData(preregisterForm);
            
            const response = await fetch("/approval/preregister", {
                method: 'POST',
                body: formData
            });
            
            const result = await response.json();
            showToast(result.success ? 'Success' : 'Error', result.message, !result.success);

            if (result.success) {
                preregisterModal.hide();
                preregisterForm.reset();
            }
        });
    }

    function showToast(title, body, isError = false) {
        const toastContainer = document.getElementById('toast-container');
        const toastId = 'toast-' + Date.now();
        const toastHtml = `
            <div id="${toastId}" class="toast" role="alert" aria-live="assertive" aria-atomic="true">
              <div class="toast-header ${isError ? 'bg-danger text-white' : 'bg-success text-white'}">
                <strong class="me-auto">${title}</strong>
                <button type="button" class="btn-close btn-close-white" data-bs-dismiss="toast" aria-label="Close"></button>
              </div>
              <div class="toast-body">${body}</div>
            </div>`;
        toastContainer.insertAdjacentHTML('beforeend', toastHtml);
        const toastElement = document.getElementById(toastId);
        const toast = new bootstrap.Toast(toastElement);
        toast.show();
        toastElement.addEventListener('hidden.bs.toast', () => toastElement.remove());
    }

    // --- Kick off the initial load ---
    initialLoad();
});
</script>
</body></html>
"""

US_APPROVAL_TEMPLATE = """<!doctype html><html lang="en"><head><title>Ultrasound Approval Tracker</title><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/flatpickr/dist/flatpickr.min.css">
<style>
  body { background: #fff5f8; }
  .navbar-custom { background: #9370DB !important; }
  .navbar-custom .navbar-brand, .navbar-custom .navbar-text { color: #fff!important; }
  .table thead th { background: #9370DB; color: #fff; vertical-align: middle; }
  .btn-mauve { background-color: #9370DB; color: white; }
  .btn-outline-mauve { color: #9370DB; border-color: #9370DB; }
  .btn-outline-mauve:hover, .btn-check:active+.btn-outline-mauve, .btn-check:checked+.btn-outline-mauve { background-color: #9370DB; color: white; }
  .toast-container { z-index: 1090; }
  .inline-activity-log { max-height: 120px; overflow-y: auto; font-size: 0.85em; background-color: #fdfdfd; padding: 0.5rem; border-radius: 4px; border: 1px solid #eee; }
  .inline-note-entry { border-bottom: 1px solid #f0f0f0; padding-bottom: 0.4rem; margin-bottom: 0.4rem; }
  .inline-note-entry:last-child { border-bottom: none; }
  .subscriber-list { list-style-type: none; padding-left: 0; margin-bottom: 0.5rem; max-height: 80px; overflow-y: auto;}
</style>
</head><body>
<nav class="navbar navbar-expand-lg navbar-custom mb-3">
  <div class="container-fluid"><span class="navbar-brand">US Approval Tracker</span>
    <div class="ms-auto"><span class="navbar-text me-2">User: {{ session.full_name }}</span>
      <a href="/" class="btn btn-light btn-sm me-2">Main</a><a href="/approval" class="btn btn-light btn-sm me-2">CT Tracker</a><a href="/logout" class="btn btn-warning btn-sm">Logout</a>
    </div>
  </div>
</nav>
<div class="container-fluid">
    <div id="toast-container" class="toast-container position-fixed top-0 end-0 p-3"></div>
    <div class="row mb-2 align-items-center">
        <div class="col-md-4"><input id="usApprovalSearch" type="search" class="form-control" placeholder="Search..."></div>
        <div class="col-md-8 d-flex align-items-center">
            <div class="btn-group me-3" role="group" id="statusFilters">
                <button type="button" class="btn btn-mauve filter-btn active" data-filter="All">All</button>
                <button type="button" class="btn btn-outline-mauve filter-btn" data-filter="Pending Review">Pending</button>
                <button type="button" class="btn btn-outline-mauve filter-btn" data-filter="Info Needed">Info Needed</button>
                <button type="button" class="btn btn-outline-mauve filter-btn" data-filter="Approved">Approved</button>
                <button type="button" class="btn btn-outline-mauve filter-btn" data-filter="Completed">Completed</button>
            </div>
            <select class="form-select w-auto" id="usSortSelect">
                <option value="activity-desc">Sort by Last Activity (Newest)</option>
                <option value="activity-asc">Sort by Last Activity (Oldest)</option>
                <option value="name-asc">Sort by Patient Name (A-Z)</option>
            </select>
        </div>
    </div>
    <div class="table-responsive">
        <table class="table table-bordered table-hover align-middle" id="usTable">
            <thead><tr>
                <th style="width: 15%;">Notifications</th>
                <th style="width: 20%;">Patient Info</th><th style="width: 10%;">Ward</th><th style="width: 10%;">Status</th>
                <th style="width: 25%;">Activity Log</th>
                <th style="width: 20%;">Actions</th>
            </tr></thead>
            <tbody id="usTableBody"></tbody>
        </table>
    </div>
</div>

<div class="modal fade" id="usManageModal" tabindex="-1">
  <div class="modal-dialog">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title" id="usManageModalLabel">Manage Request</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body">
        <div id="modal-status-form-container"></div>
        <div id="modal-prereg-form-container"></div>
      </div>
    </div>
  </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/flatpickr"></script>
<script>
  document.addEventListener('DOMContentLoaded', () => {
    let allRequests = []; let userPermissions = []; let allUsers = {}; let currentUser = '';
    let activeStatusFilter = 'All'; let activePopoverTrigger = null;
    let currentSort = 'activity-desc'; // ADDED: State for sorting

    const tbody = document.getElementById('usTableBody');
    const search = document.getElementById('usApprovalSearch');
    const filterButtons = document.querySelectorAll('.filter-btn');
    const sortSelect = document.getElementById('usSortSelect'); // ADDED: Sort select element
    const usManageModal = new bootstrap.Modal(document.getElementById('usManageModal'));
    
    window.showToast = function(title, body, isError = false) {
      const toastContainer = document.getElementById('toast-container');
      const toastId = 'toast-' + Date.now();
      const toastHtml = `<div id="${toastId}" class="toast" role="alert"><div class="toast-header ${isError ? 'bg-danger text-white' : 'bg-success text-white'}"><strong class="me-auto">${title}</strong><button type="button" class="btn-close" data-bs-dismiss="toast"></button></div><div class="toast-body">${body}</div></div>`;
      toastContainer.insertAdjacentHTML('beforeend', toastHtml);
      new bootstrap.Toast(document.getElementById(toastId)).show();
    }

    function renderTable() {
      // --- SORTING LOGIC ADDED ---
      allRequests.sort((a, b) => {
          const nameA = a.parsed_info?.patient_name || '';
          const nameB = b.parsed_info?.patient_name || '';
          const timeA = a.last_activity_on || '';
          const timeB = b.last_activity_on || '';
          switch (currentSort) {
              case 'name-asc': return nameA.localeCompare(nameB);
              case 'activity-asc': return timeA.localeCompare(timeB);
              case 'activity-desc': default: return timeB.localeCompare(timeA);
          }
      });

      const query = search.value.toLowerCase();
      let rowsHtml = allRequests.map(req => {
          const p = req.parsed_info || {};
          const searchText = `${p.patient_name || ''} ${p.patient_id || ''} ${p.ward_part || ''}`.toLowerCase();
          const matchesSearch = searchText.includes(query);
          const matchesStatus = (activeStatusFilter === 'All' || req.status === activeStatusFilter || (activeStatusFilter === 'Approved' && req.status.startsWith('Approved')));
          if (!matchesSearch || !matchesStatus) return ''; // Filter out here instead of using display:none

          // --- ACTIVITY LOG LOGIC UPDATED ---
          // Now includes the "Card Uploaded" event by default
          let events = [...(req.notes || [])].map(n => ({ ...n, type: 'note' }));
          events.push({ on: req.creation_time, text: 'Request Card Uploaded', by: 'System', type: 'system' });
          events.sort((a, b) => new Date(b.on) - new Date(a.on));
          let logHtml = '<div class="inline-activity-log">';
          if (events.length > 0) {
              logHtml += events.map(event => `<div class="inline-note-entry"><div>${event.text}</div><div class="text-muted"><em>- ${event.by} on ${new Date(event.on).toLocaleDateString()}</em></div></div>`).join('');
          } else { logHtml += '<span class="text-muted">No activity yet.</span>'; }
          logHtml += '</div>';

          let notificationsHtml = '<span class="text-muted small">No subscription options.</span>';
          const subscribers = req.subscribers || [];
          if (userPermissions.includes('manage_all_subscriptions')) {
              const subsHtml = subscribers.length > 0 ? subscribers.map(sub => `<li><span>${allUsers[sub] ? allUsers[sub].full_name : sub}</span><form method="POST" action="/approval/unsubscribe_user" class="d-inline"><input type="hidden" name="filename" value="${req.filename}"><input type="hidden" name="username_to_unsubscribe" value="${sub}"><button type="submit" class="btn btn-danger btn-sm p-0 px-1">&times;</button></form></li>`).join('') : '<li class="text-muted small">None</li>';
              const userOptions = Object.keys(allUsers).map(username => !subscribers.includes(username) && allUsers[username]?.email ? `<option value="${username}">${allUsers[username].full_name}</option>` : '').join('');
              notificationsHtml = `<ul class="subscriber-list">${subsHtml}</ul><form method="POST" action="/approval/subscribe_user"><input type="hidden" name="filename" value="${req.filename}"><div class="input-group"><select name="username_to_subscribe" class="form-select form-select-sm"><option value="">-- Add --</option>${userOptions}</select><button type="submit" class="btn btn-info btn-sm">Add</button></div></form>`;
          } else if (userPermissions.includes('manage_own_subscriptions')) {
              const isSubscribed = subscribers.includes(currentUser);
              notificationsHtml = `<form method="POST" action="/approval/toggle_approval_subscription" class="d-grid"><input type="hidden" name="filename" value="${req.filename}"><button type="submit" class="btn btn-sm ${isSubscribed ? 'btn-secondary' : 'btn-info'}">${isSubscribed ? 'Unsubscribe' : 'Notify Me'}</button></form>`;
          }

          let actionsHtml = `<a href="/request_card?path=${encodeURIComponent(req.file_path_safe)}" class="btn btn-sm btn-outline-secondary" target="_blank">View Card</a>`;
          if (userPermissions.includes('add_approval_notes')) {
              const formHtml = `<form class="popover-note-form" data-filename="${req.filename}"><textarea class="form-control form-control-sm mb-2" name="note_text" rows="3"></textarea><button type="submit" class="btn btn-primary btn-sm">Save</button></form>`;
              actionsHtml += `<button type="button" class="btn btn-sm btn-outline-info ms-1" data-bs-toggle="popover" data-bs-html="true" data-bs-title="Add a Note" data-bs-content="${formHtml.replace(/"/g, '&quot;')}">Note</button>`;
          }
          const canManage = userPermissions.includes('manage_approvals') || userPermissions.includes('preregister_patient');
          if (canManage) { actionsHtml += `<button class="btn btn-sm btn-mauve manage-btn ms-1" data-filename="${req.filename}">Manage</button>`; }

          return `<tr>
              <td>${notificationsHtml}</td>
              <td><strong>${p.patient_name || 'N/A'}</strong><br><small class="text-muted">${p.patient_id || 'N/A'}</small></td>
              <td>${p.ward_part || 'N/A'}</td>
              <td><span class="badge ${req.status === 'Pending Review' ? 'bg-warning text-dark' : req.status === 'Completed' ? 'bg-success' : 'bg-info'}">${req.status}</span></td>
              <td>${logHtml}</td>
              <td>${actionsHtml}</td>
            </tr>`;
      }).join('');
      tbody.innerHTML = rowsHtml || '<tr><td colspan="6" class="text-center py-5 text-muted">No Ultrasound requests found.</td></tr>';
      [...document.querySelectorAll('[data-bs-toggle="popover"]')].map(el => new bootstrap.Popover(el, { sanitize: false }));
    }

    window.loadData = function() {
      fetch('/api/us_approvals').then(r => r.json()).then(data => {
        allRequests = data.requests || []; userPermissions = data.permissions || []; allUsers = data.all_users || {}; currentUser = data.current_user || '';
        renderTable();
      });
    }
    
    tbody.addEventListener('click', function(e) {
        const manageButton = e.target.closest('.manage-btn');
        if (!manageButton) return;
        const filename = manageButton.dataset.filename;
        const requestData = allRequests.find(r => r.filename === filename);
        if (!requestData) return;
        const pInfo = requestData.parsed_info || {};
        document.getElementById('usManageModalLabel').textContent = `Manage: ${pInfo.patient_name}`;
        document.getElementById('modal-status-form-container').innerHTML = userPermissions.includes('manage_approvals') ? `<form method="POST" action="/approval/update_status" class="mb-3"><input type="hidden" name="filename" value="${filename}"><label class="form-label">Update Status</label><div class="input-group"><select name="new_status" class="form-select"><option>Pending Review</option><option>Info Needed</option><option>Approved</option><option>Completed</option><option>Rejected</option></select><button type="submit" class="btn btn-mauve">Set</button></div></form>` : '';
        document.getElementById('modal-prereg-form-container').innerHTML = userPermissions.includes('preregister_patient') ? `<div class="accordion" id="appointmentAccordion"><div class="accordion-item"><h2 class="accordion-header"><button class="accordion-button collapsed" type="button" data-bs-toggle="collapse" data-bs-target="#collapseAppointment">Set Appointment</button></h2><div id="collapseAppointment" class="accordion-collapse collapse" data-bs-parent="#appointmentAccordion"><div class="accordion-body"><form id="usPreregisterForm"><input type="hidden" name="filename" value="${filename}"><input type="hidden" name="modality" value="US"><div class="mb-2"><label>Patient Name</label><input type="text" class="form-control" name="patient_name" value="${pInfo.patient_name || ''}" readonly></div><div class="mb-2"><label>Patient ID</label><input type="text" class="form-control" name="patient_id" value="${pInfo.patient_id || ''}" readonly></div><div class="mb-2"><label>Appointment Time</label><input type="text" class="form-control prereg-time-input" name="scheduled_datetime" required></div><button type="submit" class="btn btn-success">Save Appointment</button></form></div></div></div></div>` : '';
        flatpickr(".prereg-time-input", { enableTime: true, dateFormat: "Y-m-d H:i", defaultDate: new Date() });
        usManageModal.show();
    });
    
    tbody.addEventListener('shown.bs.popover', (event) => { activePopoverTrigger = event.target; });

    document.body.addEventListener('submit', function(e) {
      if (e.target && e.target.id === 'usPreregisterForm') {
        e.preventDefault();
        const formData = new FormData(e.target);
        fetch('/approval/preregister', { method: 'POST', body: formData })
          .then(res => res.json()).then(result => {
            showToast(result.success ? 'Success' : 'Error', result.message, !result.success);
            if (result.success) { usManageModal.hide(); loadData(); }
          }).catch(err => showToast('Network Error', 'Could not submit appointment.', true));
      }
      if (e.target && e.target.classList.contains('popover-note-form')) {
        e.preventDefault();
        const formData = new FormData(e.target);
        if (activePopoverTrigger) { bootstrap.Popover.getInstance(activePopoverTrigger)?.hide(); }
        fetch('/approval/add_note', { method: 'POST', body: formData })
          .then(res => {
            showToast('Success', 'Note added successfully!');
            loadData();
          }).catch(err => showToast('Error', 'Could not add note.', true));
      }
    });

    search.addEventListener('input', renderTable);
    filterButtons.forEach(button => {
      button.addEventListener('click', () => {
        filterButtons.forEach(btn => {btn.classList.remove('active', 'btn-mauve'); btn.classList.add('btn-outline-mauve');});
        button.classList.add('active', 'btn-mauve');
        activeStatusFilter = button.dataset.filter;
        renderTable();
      });
    });
    // ADDED: Event listener for the new sort dropdown
    sortSelect.addEventListener('change', (e) => {
        currentSort = e.target.value;
        renderTable();
    });

    loadData();
    setInterval(loadData, 60000);
  });
</script>
</body></html>
"""



SETTINGS_TEMPLATE = """
<!doctype html><html lang="en"><head><title>System Settings - CRH CT Tracker</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
<style>.form-check-label{cursor:pointer;}.form-switch .form-check-input{cursor:pointer;}</style>
</head><body><nav class="navbar navbar-expand-lg" style="background:#003366"><div class="container-fluid"><img src="/logo" alt="Logo" style="height:38px;margin-right:14px"><span class="navbar-brand text-white">System Settings</span><div class="d-flex ms-auto"><span class="navbar-text text-white me-3">User: {{ session.full_name }}</span><a href="/" class="btn btn-outline-light me-2">Main Tracker</a><a href="/mwl" class="btn btn-outline-info me-2">MWL Server</a><a href="/admin" class="btn btn-outline-secondary me-2">User Admin</a><a href="/logout" class="btn btn-outline-warning">Logout</a></div></div></nav>
<div class="container mt-4">
    {% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, message in messages %}<div class="alert alert-{{ category }} alert-dismissible fade show" role="alert">{{ message }}<button type="button" class="btn-close" data-bs-dismiss="alert" aria-label="Close"></button></div>{% endfor %}{% endif %}{% endwith %}
    <form id="mainSettingsForm" action="{{ url_for('settings') }}" method="post">
        <div class="row">
            <div class="col-lg-6">
                <div class="card mb-4"><div class="card-header">User Role Permissions</div><div class="card-body">
                    <p class="card-text small text-muted">Control what each user role can see and do within the application.</p>
                    {% for role, permissions in config.USER_ROLE_PERMISSIONS.items() %}
                    <h6 class="mt-3">Role: <span class="badge bg-secondary">{{ role.replace('_', ' ')|title }}</span></h6>
                    <div class="row row-cols-2">
                        {% for p in ALL_PERMISSIONS %}
                        <div class="col">
                            <div class="form-check form-switch">
                                <input class="form-check-input" type="checkbox" role="switch" name="perm_{{ role }}_{{ p }}" id="perm-{{ role }}-{{ p }}" value="True" {% if p in permissions %}checked{% endif %} {% if role == 'admin' %}disabled{% endif %}>
                                <label class="form-check-label small" for="perm-{{ role }}-{{ p }}">{{ p.replace('_', ' ')|title }}</label>
                            </div>
                        </div>
                        {% endfor %}
                    </div>
                    {% endfor %}
                </div></div>
                <div class="card mb-4"><div class="card-header">Branding Settings</div><div class="card-body">
                    <div class="mb-3"><label class="form-label">Institution Name</label><input type="text" class="form-control" name="INSTITUTION_NAME" value="{{ config.INSTITUTION_NAME }}"></div>
                </div></div>
                <div class="card mb-4"><div class="card-header">PACS & DICOM Settings</div><div class="card-body">
                    <div class="mb-3"><label class="form-label">PACS Root Directory</label><input type="text" class="form-control" name="DICOM_ROOT" value="{{ config.DICOM_ROOT }}"></div>
                    <div class="mb-3"><label class="form-label">PACS IP Address</label><input type="text" class="form-control" name="PACS_IP" value="{{ config.PACS_IP }}"></div>
                    <div class="mb-3"><label class="form-label">PACS Port</label><input type="number" class="form-control" name="PACS_PORT" value="{{ config.PACS_PORT }}"></div>
                    <div class="mb-3"><label class="form-label">PACS AE Title (Their AE Title)</label><input type="text" class="form-control" name="PACS_AE_TITLE" value="{{ config.PACS_AE_TITLE }}"></div>
                    <div class="mb-3"><label class="form-label">Local AE Title (Our AE Title)</label><input type="text" class="form-control" name="LOCAL_AE_TITLE" value="{{ config.LOCAL_AE_TITLE }}"></div>
                </div></div>
                <div class="card mb-4"><div class="card-header">Folder & Polling Settings</div><div class="card-body">
                    <div class="mb-3"><label class="form-label">CT Request Card (JPG) Watch Folders (one per line)</label><textarea class="form-control" name="JPG_WATCH_FOLDERS" rows="3">{{ config.JPG_WATCH_FOLDERS|join('\\n') }}</textarea></div>
                    <div class="mb-3"><label class="form-label">Ultrasound Request Card (JPG) Watch Folders (one per line)</label><textarea class="form-control" name="US_WATCH_FOLDERS" rows="3">{{ config.US_WATCH_FOLDERS|join('\\n') }}</textarea></div>
                    <hr>
                    <div class="mb-3"><label class="form-label">PACS Poller Interval (seconds)</label><input type="number" class="form-control" name="PACS_POLL_INTERVAL" value="{{ config.PACS_POLL_INTERVAL }}"></div>
                    <div class="mb-3"><label class="form-label">Approval Poller Interval (seconds)</label><input type="number" class="form-control" name="APPROVAL_POLL_INTERVAL" value="{{ config.APPROVAL_POLL_INTERVAL }}"></div>
                </div></div>
            </div>
            <div class="col-lg-6">
                <div class="card mb-4"><div class="card-header">Email Notification Settings</div><div class="card-body">
                    <div class="form-check form-switch mb-3"><input class="form-check-input" type="checkbox" role="switch" id="emailEnabledCheck" name="EMAIL_ENABLED" value="True" {% if config.EMAIL_ENABLED %}checked{% endif %}><label class="form-check-label" for="emailEnabledCheck">Enable Email Notifications</label></div>
                    <div class="mb-3"><label class="form-label">Radiologist Notification Emails (one per line)</label><textarea class="form-control" name="RADIOLOGIST_EMAILS" rows="2" placeholder="email1@example.com&#x0a;email2@example.com">{{ config.RADIOLOGIST_EMAILS|join('\\n') }}</textarea></div>
                    <div class="mb-3"><label class="form-label">Archiving Notification Emails (one per line)</label><textarea class="form-control" name="ARCHIVING_EMAILS" rows="2" placeholder="archive1@example.com&#x0a;archive2@example.com">{{ config.ARCHIVING_EMAILS|join('\\n') }}</textarea></div>
                    <hr>
                    <h5 class="card-title mt-3">SMTP Server Settings</h5><div class="row">
                        <div class="col-md-8"><div class="mb-3"><label class="form-label">SMTP Host</label><input type="text" class="form-control" name="SMTP_HOST" value="{{ config.SMTP_HOST }}"></div></div>
                        <div class="col-md-4"><div class="mb-3"><label class="form-label">SMTP Port</label><input type="number" class="form-control" name="SMTP_PORT" value="{{ config.SMTP_PORT }}"></div></div></div>
                    <div class="mb-3"><label class="form-label">Sender Email Address</label><input type="email" class="form-control" name="SMTP_SENDER_EMAIL" value="{{ config.SMTP_SENDER_EMAIL }}"></div>
                    <div class="mb-3"><label class="form-label">SMTP Username</label><input type="text" class="form-control" name="SMTP_USER" value="{{ config.SMTP_USER }}"></div>
                    <div class="mb-3"><label class="form-label">SMTP Password</label><input type="password" class="form-control" name="SMTP_PASSWORD" value="{{ config.SMTP_PASSWORD }}"></div><hr>
                    <div class="d-grid"><button type="button" class="btn btn-primary" data-bs-toggle="modal" data-bs-target="#testSmtpModal">Test SMTP Settings</button></div>
                </div></div>
                <div class="card mb-4"><div class="card-header">Modality Worklist (MWL) Server Settings</div><div class="card-body">
                    <div class="form-check form-switch mb-3"><input class="form-check-input" type="checkbox" role="switch" id="mwlEnabledCheck" name="MWL_ENABLED" value="True" {% if config.MWL_ENABLED %}checked{% endif %}><label class="form-check-label" for="mwlEnabledCheck">Enable MWL Server</label></div>
                    <div class="form-check form-switch mb-3"><input class="form-check-input" type="checkbox" role="switch" id="mppsEnabledCheck" name="MWL_MPPS_ENABLED" value="True" {% if config.MWL_MPPS_ENABLED %}checked{% endif %}><label class="form-check-label" for="mppsEnabledCheck">Enable MPPS Support (auto-removes completed studies)</label></div>
                    <div class="mb-3"><label class="form-label">MWL Server AE Title</label><input type="text" class="form-control" name="MWL_AE_TITLE" value="{{ config.MWL_AE_TITLE }}"></div>
                    <div class="mb-3"><label class="form-label">MWL Server Port</label><input type="number" class="form-control" name="MWL_PORT" value="{{ config.MWL_PORT }}"></div><hr>
                    <div class="mb-3"><label class="form-label">Default Accession Prefix</label><input type="text" class="form-control" name="DEFAULT_ACCESSION_PREFIX" value="{{ config.DEFAULT_ACCESSION_PREFIX }}"></div>
                    <div class="mb-3"><label class="form-label">Default Scheduled Station AE</label><input type="text" class="form-control" name="DEFAULT_SCHEDULED_STATION_AE" value="{{ config.DEFAULT_SCHEDULED_STATION_AE }}"></div>
                </div></div>
                 <div class="card mb-4"><div class="card-header">DOCX Report Settings</div><div class="card-body">
                    <div class="mb-3"><label class="form-label">Report DOCX Template Path</label><input type="text" class="form-control" name="DOCX_TEMPLATE_PATH" value="{{ config.DOCX_TEMPLATE_PATH }}"></div>
                    <hr><h6 class="card-subtitle mb-2 text-muted">DOCX Report Output Folders by Modality</h6>
                    {% for mod in ['General', 'CT', 'DX', 'US', 'MG', 'MR'] %}
                    <div class="mb-2 row"><label class="col-sm-2 col-form-label">{{ mod }}</label><div class="col-sm-10"><input type="text" class="form-control form-control-sm" name="DOCX_OUTPUT_FOLDERS_{{mod}}" value="{{ config.DOCX_OUTPUT_FOLDERS[mod] }}"></div></div>
                    {% endfor %}
                </div></div>
            </div>
        </div>
        <button type="submit" class="btn btn-success btn-lg mb-4">Save All Settings</button>
    </form>
</div>
<div class="modal fade" id="testSmtpModal" tabindex="-1"><div class="modal-dialog"><div class="modal-content">
    <div class="modal-header"><h5 class="modal-title">Test SMTP Settings</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
    <form id="testSmtpForm" action="{{ url_for('test_smtp_settings') }}" method="post"><div class="modal-body">
        <p>This will send a test email using the settings from the main form to the address below.</p>
        <div class="mb-3"><label for="test_email_recipient" class="form-label">Recipient Email Address</label><input type="email" class="form-control" id="test_email_recipient" name="test_email_recipient" required value="{{ config.SMTP_USER }}"></div>
        <input type="hidden" name="SMTP_HOST_TEST"><input type="hidden" name="SMTP_PORT_TEST"><input type="hidden" name="SMTP_SENDER_EMAIL_TEST">
        <input type="hidden" name="SMTP_USER_TEST"><input type="hidden" name="SMTP_PASSWORD_TEST">
    </div><div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button><button type="submit" class="btn btn-primary">Send Test Email</button></div></form>
</div></div></div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script>
    document.getElementById('testSmtpModal').addEventListener('show.bs.modal', function () {
        document.querySelector('input[name="SMTP_HOST_TEST"]').value = document.querySelector('input[name="SMTP_HOST"]').value;
        document.querySelector('input[name="SMTP_PORT_TEST"]').value = document.querySelector('input[name="SMTP_PORT"]').value;
        document.querySelector('input[name="SMTP_SENDER_EMAIL_TEST"]').value = document.querySelector('input[name="SMTP_SENDER_EMAIL"]').value;
        document.querySelector('input[name="SMTP_USER_TEST"]').value = document.querySelector('input[name="SMTP_USER"]').value;
        document.querySelector('input[name="SMTP_PASSWORD_TEST"]').value = document.querySelector('input[name="SMTP_PASSWORD"]').value;
    });
</script></body></html>
"""
PROFILE_TEMPLATE = """
<!doctype html><html lang="en"><head><title>Complete Profile - CRH CT Tracker</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><style>body{display:flex;align-items:center;justify-content:center;min-height:100vh;background-color:#f0f2f5}.profile-card{width:100%;max-width:450px;padding:2rem;border-radius:.5rem;box-shadow:0 4px 12px rgba(0,0,0,.1)}</style></head><body><div class="card profile-card"><h3 class="card-title mb-4 text-center">Complete Your Profile</h3><p class="text-muted text-center">An email address is required to use the CT Tracker and receive important notifications.</p><form method="post"><div class="mb-3"><label for="email" class="form-label">Your Email Address</label><input type="email" class="form-control" id="email" name="email" value="{{ user.email }}" required autofocus></div><div class="form-check form-switch mb-4"><input class="form-check-input" type="checkbox" role="switch" id="notify" name="notify_on_updates" value="True" {% if user.notify_on_updates %}checked{% endif %}><label class="form-check-label" for="notify">Receive email notifications for approvals and reports</label></div><button type="submit" class="btn btn-primary w-100">Save Profile and Continue</button></form></div></body></html>
"""
MWL_TEMPLATE = """
<!doctype html><html lang="en"><head><title>MWL Server - CRH CT Tracker</title><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css">
<style>
    body { background-color: #f8f9fa; }
    .navbar-custom { background: #003366 !important; }
    .navbar-custom .navbar-brand, .navbar-custom .navbar-text { color: #fff !important; }
    .banner-logo { height: 45px; margin-right: 14px; }
    .table-container { max-height: 50vh; overflow-y: auto; }
    .table thead th { background: #003366; color: #fff; white-space: nowrap; position: sticky; top: 0; z-index: 10; }
    .activity-log { height: 250px; overflow-y: scroll; background-color: #2B3035; color: #00FF41; font-family: 'Consolas', 'Monaco', monospace; font-size: 0.85rem; padding: 10px; border-radius: 5px; white-space: pre-wrap; }
    .server-status { font-weight: 500; }
    .status-enabled { color: #198754; } .status-disabled { color: #dc3545; }
    .form-label { font-weight: 500; }
    .action-btns .btn { padding: 0.2rem 0.5rem; font-size: 0.8rem; }
    .toast-container { z-index: 1090; }
    .expected-patient-row:hover { background-color: #e9ecef; cursor: pointer; }
</style>
</head>
<body>
<nav class="navbar navbar-expand-lg navbar-custom">
    <div class="container-fluid">
        <img src="/logo" alt="Logo" class="banner-logo"><span class="navbar-brand">CRH Radiology - MWL Server</span>
        <div class="d-flex align-items-center ms-auto">
            <span class="navbar-text me-3">User: {{ session.full_name }}</span>
            {% if has_permission('view_dashboard') %}<a href="/" class="btn btn-outline-light me-2">Main Tracker</a>{% endif %}
            {% if has_permission('view_approval_tracker') %}<a href="/approval" class="btn btn-outline-info me-2">Approval Tracker</a>{% endif %}
            {% if has_permission('view_us_approval_tracker') %}<a href="/us_approval" class="btn btn-outline-info me-2">US Tracker</a>{% endif %}
            {% if has_permission('view_admin_page') %}<a href="/admin" class="btn btn-secondary me-2">Admin</a>{% endif %}
            {% if has_permission('view_settings_page') %}<a href="/settings" class="btn btn-info me-2">Settings</a>{% endif %}
            {% if has_permission('view_resources_page') %}<a href="/resources" class="btn btn-outline-success me-2">Resources</a>{% endif %}
            <a href="/logout" class="btn btn-warning">Logout</a>
        </div>
    </div>
</nav>
<div class="container-fluid mt-3">
    <div class="row">
        <div class="col-lg-4">
            <div class="card mb-3">
                <div class="card-header">
                    <a class="h5 text-decoration-none" data-bs-toggle="collapse" href="#patientFormCollapse" role="button" aria-expanded="false" aria-controls="patientFormCollapse">
                        Patient Registration
                    </a>
                </div>
                <div class="collapse" id="patientFormCollapse">
                    <div class="card-body">
                        <form id="mwlPatientForm">
                             <input type="hidden" name="record_id" id="record_id">
                             <input type="hidden" name="expected_id" id="expected_id">
                             <div class="row">
                                <div class="col-md-6 mb-2"><label for="patient_name" class="form-label">Patient Name</label><input type="text" class="form-control" id="patient_name" name="patient_name" required></div>
                                <div class="col-md-6 mb-2"><label for="patient_id" class="form-label">Patient ID</label><input type="text" class="form-control" id="patient_id" name="patient_id" required></div>
                                <div class="col-md-6 mb-2"><label for="dob" class="form-label">Date of Birth</label><input type="text" class="form-control" id="dob" name="dob" placeholder="DD/MM/YYYY or DDMMYYYY" required></div>
                                <div class="col-md-6 mb-2"><label for="sex" class="form-label">Sex</label><select class="form-select" id="sex" name="sex" required><option value="">Select...</option><option value="M">Male</option><option value="F">Female</option><option value="O">Other</option></select></div>
                                <div class="col-md-12 mb-2"><label for="accession_number" class="form-label">Accession Number</label><input type="text" class="form-control" id="accession_number" name="accession_number" placeholder="{{config.DEFAULT_ACCESSION_PREFIX}}CT12345" value="{{config.DEFAULT_ACCESSION_PREFIX}}" required></div>
                                <div class="col-md-12 mb-2"><label for="study_description" class="form-label">Study Description</label><input type="text" class="form-control" id="study_description" name="study_description" required></div>
                                <div class="col-md-6 mb-2"><label for="referred_from" class="form-label">Referred From</label><input type="text" class="form-control" id="referred_from" name="referred_from" required></div>
                                <div class="col-md-6 mb-2"><label for="requesting_physician" class="form-label">Requesting Physician</label><input type="text" class="form-control" id="requesting_physician" name="requesting_physician" required></div>
                            </div>
                            <div class="d-flex justify-content-end mt-3">
                                <button type="button" class="btn btn-secondary me-2" id="clearFormBtn">Clear</button>
                                <button type="submit" class="btn btn-primary" id="submitBtn">Add Patient</button>
                            </div>
                        </form>
                    </div>
                </div>
            </div>
             <div class="card">
                <div class="card-header"><h5 class="mb-0">Today's Expected Patients</h5></div>
                <div class="card-body p-0">
                    <div class="table-responsive" style="max-height: 250px; overflow-y: auto;">
                        <table class="table table-sm table-hover mb-0">
                            <thead><tr><th>Name</th><th>ID</th><th>Time</th><th>Modality</th><th>Select</th><th>Action</th></tr></thead>
                            <tbody id="expectedPatientsTableBody"></tbody>
                        </table>
                    </div>
                </div>
            </div>
        </div>

        <div class="col-lg-8 d-flex flex-column">
            <div class="card flex-grow-1">
                <div class="card-header d-flex justify-content-between align-items-center">
                    <h5 class="mb-0">Current Worklist</h5>
                    <div>
                        <strong>Server Status:</strong>
                        {% if config.MWL_ENABLED %}
                        <span class="server-status status-enabled">ENABLED</span>
                        <small class="text-muted"> ({{config.MWL_AE_TITLE}} @ {{config.MWL_PORT}})</small>
                        {% else %}
                        <span class="server-status status-disabled">DISABLED</span>
                        {% endif %}
                    </div>
                </div>
                <div class="card-body table-container">
                    <table class="table table-sm table-bordered table-hover">
                        <thead><tr><th>Patient Name</th><th>Patient ID</th><th>Accession</th><th>DOB</th><th>Sex</th><th>Modality</th><th>Study</th><th>Actions</th></tr></thead>
                        <tbody id="mwlTableBody"><tr><td colspan="8" class="text-center p-5">Loading worklist...</td></tr></tbody>
                    </table>
                </div>
            </div>
            <div class="accordion mt-3" id="logAccordion">
                <div class="accordion-item">
                    <h2 class="accordion-header" id="headingOne">
                        <button class="accordion-button collapsed" type="button" data-bs-toggle="collapse" data-bs-target="#collapseLog" aria-expanded="false" aria-controls="collapseLog">
                            MWL/MPPS Activity Log
                        </button>
                    </h2>
                    <div id="collapseLog" class="accordion-collapse collapse" aria-labelledby="headingOne" data-bs-parent="#logAccordion">
                        <div class="accordion-body p-0">
                            <div id="activityLog" class="activity-log">Connecting to activity log...</div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>
</div>

<div class="toast-container position-fixed bottom-0 end-0 p-3">
  <div id="liveToast" class="toast" role="alert" aria-live="assertive" aria-atomic="true">
    <div class="toast-header">
      <strong class="me-auto" id="toastTitle"></strong>
      <button type="button" class="btn-close" data-bs-dismiss="toast" aria-label="Close"></button>
    </div>
    <div class="toast-body" id="toastBody"></div>
  </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
<script>
    const mwlTableBody = document.getElementById('mwlTableBody');
    const activityLog = document.getElementById('activityLog');
    const patientForm = document.getElementById('mwlPatientForm');
    const submitBtn = document.getElementById('submitBtn');
    const clearFormBtn = document.getElementById('clearFormBtn');
    const recordIdField = document.getElementById('record_id');
    const expectedIdField = document.getElementById('expected_id');
    const toastElement = document.getElementById('liveToast');
    const toast = new bootstrap.Toast(toastElement);
    let lastUsedAccession = '';

    function showToast(title, body, isError = false) {
        document.getElementById('toastTitle').textContent = title;
        document.getElementById('toastBody').textContent = body;
        toastElement.classList.toggle('bg-danger', isError);
        toastElement.classList.toggle('text-white', isError);
        toast.show();
    }

    async function fetchWorklist() {
        try {
            const response = await fetch('/api/mwl_worklist');
            const worklist = await response.json();
            if (!worklist || worklist.length === 0) {
                mwlTableBody.innerHTML = '<tr><td colspan="8" class="text-center p-4">Worklist is currently empty.</td></tr>';
                return;
            }
            let rowsHtml = '';
            for (const item of worklist) {
                rowsHtml += `<tr>
                    <td>${item.patient_name}</td><td>${item.patient_id}</td><td>${item.accession_number}</td>
                    <td>${item.dob_yyyymmdd}</td><td>${item.sex}</td><td>${item.modality}</td>
                    <td>${item.study_description}</td>
                    <td class="action-btns">
                        <button class="btn btn-sm btn-outline-primary" onclick="editRecord(${item.id})"><i class="bi bi-pencil-square"></i></button>
                        <button class="btn btn-sm btn-outline-danger" onclick="deleteRecord(${item.id}, '${item.patient_name}')"><i class="bi bi-trash"></i></button>
                    </td>
                </tr>`;
            }
            mwlTableBody.innerHTML = rowsHtml;
        } catch (error) {
            console.error("Error fetching MWL:", error);
            mwlTableBody.innerHTML = '<tr><td colspan="8" class="text-center p-4 text-danger">Error fetching worklist data.</td></tr>';
        }
    }

    async function fetchActivityLog() {
        try {
            const response = await fetch('/api/mwl_activity_log');
            const logData = await response.text();
            if (activityLog.textContent !== logData) {
                activityLog.textContent = logData;
                activityLog.scrollTop = activityLog.scrollHeight;
            }
        } catch (error) {
            console.error("Error fetching activity log:", error);
        }
    }
    
    async function fetchExpectedPatients() {
        const tableBody = document.getElementById('expectedPatientsTableBody');
        try {
            const response = await fetch("{{ url_for('api_expected_patients') }}");
            const patients = await response.json();
            if (!patients || patients.length === 0) {
                tableBody.innerHTML = '<tr><td colspan="6" class="text-center small text-muted p-2">No expected patients for today.</td></tr>';
                return;
            }
            let rowsHtml = '';
            for (const p of patients) {
                rowsHtml += `<tr class="expected-patient-row" onclick="selectExpectedPatient(${p.id})">
                    <td>${p.patient_name}</td>
                    <td>${p.patient_id}</td>
                    <td>${p.scheduled_time}</td>
                    <td>${p.modality}</td>
                    <td><button type="button" class="btn btn-xs btn-outline-primary py-0 px-1">Select</button></td>
                    <td><button type="button" class="btn btn-xs btn-outline-danger py-0 px-1" onclick="event.stopPropagation(); cancelExpectedPatient(${p.id}, '${p.patient_name.replace(/'/g, "\\'")}');">Cancel</button></td>
                </tr>`;
            }
            tableBody.innerHTML = rowsHtml;
        } catch (error) {
            console.error("Error fetching expected patients:", error);
            tableBody.innerHTML = '<tr><td colspan="6" class="text-center text-danger p-2">Error loading list.</td></tr>';
        }
    }
    
    async function selectExpectedPatient(id) {
        clearForm();
        try {
            const response = await fetch(`/api/get_expected_patient/${id}`);
            const p = await response.json();
            if (p.error) {
                showToast('Error', p.error, true);
                return;
            }
            document.getElementById('patient_name').value = p.patient_name;
            document.getElementById('patient_id').value = p.patient_id;
            document.getElementById('dob').value = p.dob;
            document.getElementById('sex').value = p.sex;
            document.getElementById('study_description').value = p.study_description;
            document.getElementById('referred_from').value = p.referred_from;
            document.getElementById('requesting_physician').value = p.requesting_physician || p.referred_from;
            document.getElementById('expected_id').value = p.id;
            
            // Auto-expand the patient registration form
            const formCollapseElem = document.getElementById('patientFormCollapse');
            const formCollapse = new bootstrap.Collapse(formCollapseElem, { toggle: false });
            formCollapse.show();
            
            // Set focus to accession number field
            document.getElementById('accession_number').focus();
            showToast('Patient Loaded', `${p.patient_name} is ready for registration. Please enter the Accession Number.`, false);
        } catch (error) {
            showToast('Error', 'Could not load patient details.', true);
        }
    }

    function clearForm() {
        patientForm.reset();
        recordIdField.value = '';
        expectedIdField.value = '';
        submitBtn.textContent = 'Add Patient';
        submitBtn.classList.remove('btn-success');
        submitBtn.classList.add('btn-primary');
        document.getElementById('accession_number').readOnly = false;
        
        // Auto-populate next sequential accession number
        if (lastUsedAccession) {
            const match = lastUsedAccession.match(/(.*?)(\\d+)$/);
            if (match) {
                const prefix = match[1];
                const numberPart = match[2];
                const nextNumber = parseInt(numberPart, 10) + 1;
                const paddedNextNumber = String(nextNumber).padStart(numberPart.length, '0');
                document.getElementById('accession_number').value = prefix + paddedNextNumber;
                lastUsedAccession = ''; // Clear after use
                return; // Exit
            }
        }
        // Fallback to default prefix
        document.getElementById('accession_number').value = "{{config.DEFAULT_ACCESSION_PREFIX}}";
    }
    
    async function editRecord(id) {
        clearForm();
        const response = await fetch(`/api/mwl/record/${id}`);
        const record = await response.json();
        if (record.error) {
            showToast('Error', record.error, true);
            return;
        }
        
        for (const key in record) {
            const field = document.getElementById(key);
            if (field) {
                field.value = record[key];
            }
        }
        recordIdField.value = record.id;
        submitBtn.textContent = 'Update Patient';
        submitBtn.classList.remove('btn-primary');
        submitBtn.classList.add('btn-success');
        document.getElementById('accession_number').readOnly = true;

        const formCollapse = new bootstrap.Collapse(document.getElementById('patientFormCollapse'), {
            toggle: false
        });
        formCollapse.show();
        window.scrollTo(0, 0);
    }
    
    async function deleteRecord(id, name) {
        if (confirm(`Are you sure you want to delete the worklist entry for ${name}?`)) {
            const response = await fetch(`/api/mwl/delete/${id}`, { method: 'DELETE' });
            const result = await response.json();
            if (result.success) {
                showToast('Success', 'Record deleted successfully.');
                fetchWorklist();
            } else {
                showToast('Error', result.error, true);
            }
        }
    }

    async function cancelExpectedPatient(id, name) {
        if (confirm(`Are you sure you want to cancel the expected registration for ${name}?`)) {
            try {
                const response = await fetch(`/api/expected_patient/cancel/${id}`, {
                    method: 'DELETE'
                });
                const result = await response.json();
                if (result.success) {
                    showToast('Success', result.message, false);
                    fetchExpectedPatients(); // Refresh the list
                } else {
                    showToast('Error', result.error, true);
                }
            } catch (error) {
                console.error("Error cancelling patient:", error);
                showToast('Error', 'A network error occurred.', true);
            }
        }
    }

    document.addEventListener('DOMContentLoaded', () => {
        fetchWorklist();
        fetchActivityLog();
        fetchExpectedPatients();
        setInterval(fetchWorklist, 15000);
        setInterval(fetchActivityLog, 5000);
        setInterval(fetchExpectedPatients, 60000); // Refresh expected list every minute
        
        clearFormBtn.addEventListener('click', clearForm);

        patientForm.addEventListener('submit', async (e) => {
            e.preventDefault();
            lastUsedAccession = document.getElementById('accession_number').value;
            const formData = new FormData(patientForm);
            const recordId = recordIdField.value;
            let url = '/mwl/add_patient';
            
            if (recordId) {
                url = `/api/mwl/update/${recordId}`;
            } else {
                // Pre-flight duplicate check for new entries
                const checkFD = new FormData(patientForm);
                checkFD.append('check_only', 'true');
                const checkResponse = await fetch(url, { method: 'POST', body: checkFD });
                const checkResult = await checkResponse.json();

                if (checkResult.duplicate_warning) {
                    if (!confirm(checkResult.duplicate_warning + "\\n\\nProceed anyway?")) {
                        return; // User cancelled
                    }
                    formData.append('bypass_duplicate_check', 'true');
                } else if (!checkResult.success) {
                    showToast('Validation Error', checkResult.error, true);
                    return;
                }
            }
            
            // Final submission for insert or update
            const finalResponse = await fetch(url, { method: 'POST', body: formData });
            const finalResult = await finalResponse.json();
            
            if (finalResult.success) {
                showToast('Success', finalResult.message || 'Operation successful!');
                clearForm();
                fetchWorklist();
                fetchExpectedPatients(); // Refresh expected list after registration
            } else {
                showToast('Error', finalResult.error, true);
            }
        });
    });
</script>
</body></html>
"""

# --- ROUTES ---


@app.route('/logo')
def serve_logo():
    if os.path.isfile(LOGO_PATH): return send_file(LOGO_PATH, mimetype='image/png')
    abort(404)

@app.route('/admin')
@permission_required('view_admin_page')
def admin():
    """Renders the user administration page."""
    users = load_users()
    return render_template_string(ADMIN_TEMPLATE, users=users)

@app.route("/admin_add_user", methods=["POST"])
@permission_required('manage_users')
def admin_add_user():
    """Handles the creation of a new user from the admin panel."""
    users = load_users()
    username = request.form.get('username', '').upper().strip()
    password = request.form.get('password')
    role = request.form.get('role')
    full_name = request.form.get('full_name', '').strip()

    if not all([username, password, role, full_name]):
        flash("All fields are required to create a new user.", "danger")
        return redirect(url_for('admin'))

    if username in users:
        flash(f"Username '{username}' already exists.", "danger")
        return redirect(url_for('admin'))

    users[username] = {
        "password_hash": generate_password_hash(password),
        "role": role,
        "full_name": full_name,
        "email": "",
        "notify_on_updates": True
    }
    save_users(users)
    log_activity(f"ADMIN: User '{session['username']}' created new user '{username}' with role '{role}'.")
    flash(f"User '{full_name}' created successfully.", "success")
    return redirect(url_for('admin'))

@app.route("/admin_delete_user", methods=["POST"])
@permission_required('manage_users')
def admin_delete_user():
    """Handles the deletion of a user from the admin panel."""
    users = load_users()
    username_to_delete = request.form.get('username')

    if not username_to_delete:
        flash("Username not provided for deletion.", "danger")
        return redirect(url_for('admin'))

    if username_to_delete.upper() == 'ADMIN':
        flash("The default ADMIN user cannot be deleted.", "danger")
        return redirect(url_for('admin'))

    if username_to_delete in users:
        del users[username_to_delete]
        save_users(users)
        log_activity(f"ADMIN: User '{session['username']}' deleted user '{username_to_delete}'.")
        flash(f"User '{username_to_delete}' has been deleted.", "success")
    else:
        flash(f"User '{username_to_delete}' not found.", "danger")

    return redirect(url_for('admin'))

@app.template_filter('fromisoformat')
def fromisoformat_filter(s):
    try: return datetime.datetime.fromisoformat(s)
    except (TypeError, ValueError): return datetime.datetime.now()

@app.template_filter('strftime')
def strftime_filter(dt): return dt.strftime('%Y-%m-%d %H:%M')

@app.template_filter('strftime_note')
def strftime_note_filter(dt): return dt.strftime('%b %d, %H:%M')

@app.route("/profile", methods=["GET", "POST"])
def profile():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    users = load_users()
    username = session['username']
    user_data = users[username]

    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        notify = 'notify_on_updates' in request.form
        
        if email:
            user_data['email'] = email
            user_data['notify_on_updates'] = notify
            users[username] = user_data
            save_users(users)
            log_activity(f"PROFILE: User '{username}' updated their profile. Email: {email}, Notify: {notify}")
            flash("Profile updated successfully!", "success")
            return redirect(url_for('dashboard'))
        else:
            flash("Email address cannot be empty.", "danger")

    return render_template_string(PROFILE_TEMPLATE, user=user_data)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username, password = request.form.get("username", "").upper(), request.form.get("password")
        users = load_users()
        user_data = users.get(username)

        # First, check for valid username and password
        if user_data and check_password_hash(user_data['password_hash'], password):

            # --- IP RESTRICTION LOGIC ---
            user_role = user_data.get('role')
            remote_ip = request.remote_addr
            privileged_roles = ['admin', 'radiology_staff']

            if remote_ip.startswith('192.168.0'):
                if user_role not in privileged_roles:
                    log_activity(f"LOGIN_FAIL: User '{username}' with role '{user_role}' attempted login from a restricted IP ({remote_ip}). Access denied.")
                    flash(f"The {user_role.title()} role cannot log in from the 192.168.0.x network.", "danger")
                    return redirect(url_for('login'))
            
            session['username'] = username
            session['role'] = user_role
            session['full_name'] = user_data.get('full_name', username)
            log_activity(f"LOGIN: User '{username}' logged in successfully from {remote_ip}.")
            return redirect(url_for('dashboard'))
        else:
            log_activity(f"LOGIN_FAIL: Invalid credentials for username '{username}'.")
            flash("Invalid username or password.", "danger")

    return render_template_string(LOGIN_TEMPLATE)

@app.route("/api/approvals")
@permission_required('view_approval_tracker')
def api_approvals():
    """API endpoint to fetch approval requests data asynchronously."""
    with approval_cache_lock:
        requests_list = approval_cache.get('requests', [])
    
    all_users = load_users()
    
    user_permissions = CONFIG.get('USER_ROLE_PERMISSIONS', {}).get(session.get('role'), [])

    return jsonify(
        requests=requests_list,
        all_users=all_users,
        permissions=user_permissions,
        current_user=session.get('username')
    )

@app.route("/logout")
def logout():
    log_activity(f"LOGOUT: User '{session.get('username')}' logged out.")
    session.clear(); flash("You have been logged out.", "success"); return redirect(url_for('login'))

@app.route("/")
@permission_required('view_dashboard')
def dashboard():
    return render_template_string(DASHBOARD_TEMPLATE, year=datetime.datetime.now().year)

@app.route("/api/studies")
def api_studies():
    if 'username' not in session: abort(403)
    with pacs_cache_lock:
        success = "failed" not in pacs_data_cache['status'].lower() and "error" not in pacs_data_cache['status'].lower()
        return jsonify(success=success, data=pacs_data_cache['data'], status=pacs_data_cache['status'])

# --- CORE LOGIC FUNCTIONS ---
def query_pacs_for_ct_studies():
    """Queries PACS for studies of specified modalities and updates cache."""
    try:
        ae = AE(ae_title=CONFIG.get('LOCAL_AE_TITLE'))
        ae.add_requested_context(StudyRootQueryRetrieveInformationModelFind)
        
        ds = Dataset()
        ds.QueryRetrieveLevel = "STUDY"
        ds.PatientName = '*'
        today = datetime.datetime.now()
        start = (today - datetime.timedelta(days=21)).strftime('%Y%m%d')
        end   = today.strftime('%Y%m%d')
        ds.StudyDate = f"{start}-{end}"
        ds.ModalitiesInStudy = ['CT', 'CR', 'DX', 'MG', 'US']
        
        tags_to_get = ['PatientID', 'PatientName', 'AccessionNumber', 'StudyInstanceUID', 
                       'StudyDate', 'StudyTime', 'StudyDescription', 'ReferringPhysicianName', 
                       'StudyID', 'PatientSex', 'PatientBirthDate', 'Modality', 'RequestingPhysician']
        for tag in tags_to_get:
            setattr(ds, tag, '')
        
        assoc = ae.associate(CONFIG.get('PACS_IP'), int(CONFIG.get('PACS_PORT')), ae_title=CONFIG.get('PACS_AE_TITLE'))
        if not assoc.is_established:
            status_msg = "Connection Failed"
            log_activity(f"PACS_QUERY_FAIL: {status_msg}")
            with pacs_cache_lock: pacs_data_cache['status'] = status_msg
            return None
        
        responses = assoc.send_c_find(ds, StudyRootQueryRetrieveInformationModelFind)
        study_list_for_ui = []
        seen_accessions = set()
        
        for status, identifier in responses:
            if status and status.Status in (0xFF00, 0xFF01) and identifier:
                accession = getattr(identifier, 'AccessionNumber', '').strip()
                if accession and accession not in seen_accessions:
                    parsed_study_for_ui = parse_pacs_identifier(identifier)
                    if parsed_study_for_ui:
                        study_list_for_ui.append(parsed_study_for_ui)
                        seen_accessions.add(accession)
        assoc.release()

        with pacs_cache_lock:
            pacs_data_cache.update({
                'timestamp': time.time(), 'data': study_list_for_ui,
                'status': f"Updated at {datetime.datetime.now():%H:%M:%S}"
            })
        
        study_list_for_ui.sort(key=lambda x: (x.get('log_timestamp', '')), reverse=True)
        log_activity(f"PACS_QUERY: Finished. Found {len(study_list_for_ui)} unique studies.")
        return study_list_for_ui

    except Exception as e:
        status_msg = f"CRITICAL Error: {type(e).__name__}"
        log_activity(f"PACS_QUERY_CRITICAL_ERROR: {e}")
        with pacs_cache_lock: pacs_data_cache['status'] = status_msg
        return None

def update_approval_warnings(recent_studies, approval_data):
    """
    Analyzes all approval requests to add warnings for potential duplicates
    and recent existing scans for the same patient.
    """
    log_activity("SYSTEM: Cross-referencing requests for warnings.")
    
    # --- Part 1: Find all recent studies by patient ID ---
    studies_by_pid = {}
    if recent_studies:
        for study in recent_studies:
            pid = study.get('patient_id')
            if pid:
                studies_by_pid.setdefault(pid, []).append(study)

    # --- Part 2: Find all pending approval requests by patient ID ---
    pending_by_pid = {}
    for filename, req in approval_data.items():
        if req.get('visible', True) and req.get('status') == 'Pending Review':
            pid = req.get('parsed_info', {}).get('patient_id')
            if pid:
                pending_by_pid.setdefault(pid, []).append(filename)

    # --- Part 3: Iterate through all requests and add warnings ---
    for filename, req_data in approval_data.items():
        if not req_data.get('visible', True):
            continue

        req_data['warnings'] = []  # Reset warnings each time
        req_info = req_data.get('parsed_info', {})
        req_pid = req_info.get('patient_id')

        if not req_pid:
            continue

        # WARNING 1: Check for other pending request cards for the same patient
        if len(pending_by_pid.get(req_pid, [])) > 1:
            req_data['warnings'].append("DUPLICATE REQUEST: Another pending card exists for this patient.")

        # WARNING 2: Check for recent scans in PACS (within last 7 days)
        if req_pid in studies_by_pid:
            seven_days_ago = datetime.datetime.now() - datetime.timedelta(days=7)
            for study in studies_by_pid[req_pid]:
                try:
                    study_time = datetime.datetime.fromisoformat(study['log_timestamp'])
                    if study_time > seven_days_ago:
                        warning_msg = f"RECENT SCAN: Scan on {study['study_date_fmt']} (Acc# {study['accession']})."
                        if warning_msg not in req_data['warnings']:
                            req_data['warnings'].append(warning_msg)
                        # We only need to show the most recent one, so we can stop after finding the first.
                        break 
                except (ValueError, TypeError):
                    continue
    
    return approval_data

def _extract_pid_prefix(pid_str):
    if not pid_str:
        return ""
    return pid_str.strip().split()[0]

def cleanup_completed_approvals():
    log_activity("CLEANUP: Checking for old completed approval entries to hide.")
    approval_data = load_approval_data()
    updated = False
    now = datetime.datetime.now()
    for filename, request_data in list(approval_data.items()):
        if request_data.get('visible') and request_data.get('status') == 'Completed' and 'hide_after' in request_data:
            try:
                hide_after_time = datetime.datetime.fromisoformat(request_data['hide_after'])
                if now >= hide_after_time:
                    log_activity(f"CLEANUP: Hiding entry '{filename}' after 24-hour grace period.")
                    request_data['visible'] = False
                    updated = True
            except (ValueError, TypeError):
                log_activity(f"CLEANUP_WARNING: Invalid 'hide_after' timestamp for '{filename}'. Hiding now.")
                request_data['visible'] = False
                updated = True
    if updated: save_approval_data(approval_data)

def parse_pacs_identifier(identifier):
    """
    Parses a pynetdicom C-FIND identifier dataset into a clean dictionary.
    This version includes robust, multi-step modality parsing.
    """
    accession = getattr(identifier, 'AccessionNumber', '').strip()
    if not accession:
        return None  # Skip records without an accession number

    # --- MODALITY PARSING ---
    # 1. Prioritize 'ModalitiesInStudy', then fallback to 'Modality'
    modality_raw = str(getattr(identifier, 'ModalitiesInStudy', '') or getattr(identifier, 'Modality', '')).strip().upper()
    
    final_modality = 'N/A'
    primary_modalities = ['CT', 'US', 'MR', 'DX', 'CR', 'MG']

    # 2. If we have a modality string, parse it
    if modality_raw:
        # Split by common separators like '\', '/', or ' '
        parts = re.split(r'[\\/ ]', modality_raw)
        # Find the first primary modality in the parts
        for part in parts:
            if part in primary_modalities:
                final_modality = part
                break

    # 3. If still not found, check the accession number as a fallback
    if final_modality == 'N/A':
        accession_upper = accession.upper()
        # This logic assumes the modality is part of the accession, e.g., "CRHCT12345"
        for mod in primary_modalities:
            if mod in accession_upper:
                final_modality = mod
                break

    # --- DATE & TIME PARSING ---
    study_date_raw = getattr(identifier, 'StudyDate', '')
    study_time_raw = getattr(identifier, 'StudyTime', '000000')
    log_timestamp, dt_obj = '', None
    if study_date_raw:
        try:
            time_str = str(study_time_raw).split('.')[0].ljust(6, '0')
            dt_obj = datetime.datetime.strptime(f"{study_date_raw}{time_str}", "%Y%m%d%H%M%S")
            log_timestamp = dt_obj.isoformat()
        except (ValueError, AttributeError):
            pass

    # --- REPORT METADATA ---
    with report_metadata_lock:
        report_path = report_metadata.get(accession)

    # --- ASSEMBLE FINAL DICTIONARY ---
    study_data = {
        'log_timestamp': log_timestamp,
        'log_time': dt_obj.strftime('%H:%M') if dt_obj else 'N/A',
        'patient_name': str(getattr(identifier, 'PatientName', 'N/A')).replace('^', ' ').title(),
        'patient_id': getattr(identifier, 'PatientID', 'N/A'),
        'accession': accession,
        'referred_from': str(getattr(identifier, 'ReferringPhysicianName', 'N/A')).replace('^', ' ').title(),
        'study_date': study_date_raw,
        'study_date_fmt': format_pacs_date(study_date_raw, "%B %d, %Y"),
        'study_desc': str(getattr(identifier, 'StudyDescription', 'No Description')).strip(),
        'status': "Completed",  # This is static for the PACS tracker view
        'status_class': "completed",
        'report_path': report_path,
        'modality': final_modality,  # Use the cleaned-up modality
        'clinic_referred_from': str(getattr(identifier, 'RequestingPhysician', 'N/A')).replace('^', ' ').title(),
        'sex': getattr(identifier, 'PatientSex', 'N/A'),
        'date_of_birth': format_pacs_date(getattr(identifier, 'PatientBirthDate', ''), "%B %d, %Y").upper(),
    }
    return study_data

def format_pacs_date(pacs_date_str, output_format="%B %d, %Y"):
    if not pacs_date_str: return "N/A"
    try: return datetime.datetime.strptime(pacs_date_str, "%Y%m%d").strftime(output_format)
    except (ValueError, TypeError): return pacs_date_str

def find_matching_jpg_card(study_data):
    patient_id = study_data.get('patient_id')
    patient_name_str = study_data.get('patient_name', '').lower()
    if not patient_id and not patient_name_str: return None
    potential_matches = []
    for folder in CONFIG.get('JPG_WATCH_FOLDERS', []):
        if not os.path.isdir(folder): continue
        try:
            for root, _, files in os.walk(folder):
                for filename in files:
                    if not filename.lower().endswith(('.jpg', '.jpeg')): continue
                    file_path = os.path.join(root, filename)
                    norm_filename = filename.lower()
                    score = 0
                    if patient_id and patient_id.lower() in norm_filename: score += 5
                    if patient_name_str and patient_name_str in norm_filename: score += 3
                    if score > 0:
                        try: score += (os.path.getmtime(file_path) / (10**10))
                        except OSError: pass
                        potential_matches.append({'path': file_path, 'score': score})
        except OSError as e: log_activity(f"Warning: Could not read JPG directory {folder}: {e}")
    if potential_matches:
        potential_matches.sort(key=lambda x: x['score'], reverse=True)
        best_match = potential_matches[0]
        log_activity(f"JPG_MATCH: Found request card for Patient ID {patient_id} (Score: {best_match['score']:.2f}): {os.path.basename(best_match['path'])}")
        return best_match['path']
    return None

def scan_and_prepare_approvals():
    approval_data = load_approval_data()
    data_was_updated = False
    all_current_jpgs = {}

    for folder in CONFIG.get('JPG_WATCH_FOLDERS', []):
        if not os.path.isdir(folder):
            continue
        try:
            for root, _, files in os.walk(folder):
                for filename in files:
                    if filename.lower().endswith(('.jpg', '.jpeg')):
                        all_current_jpgs[filename] = os.path.join(root, filename)
        except OSError as e:
            log_activity(f"APPROVAL_SCAN_ERROR: Could not read directory {folder}: {e}")
    
    visible_requests = []
    for filename, current_path in all_current_jpgs.items():
        try:
            ctime_ts = os.path.getctime(current_path)
            creation_iso = datetime.datetime.fromtimestamp(ctime_ts).isoformat()
        except Exception as e:
            log_activity(f"APPROVAL_TIME_WARN: Could not get ctime for {current_path}: {e}")
            creation_iso = datetime.datetime.now().isoformat()

        if filename in approval_data:
            request_data = approval_data[filename]
            
            if not request_data.get('visible', True):
                continue
            
            if request_data.get('file_path') != current_path:
                request_data['file_path'] = current_path
                data_was_updated = True

            if request_data.get('creation_time') != creation_iso:
                request_data['creation_time'] = creation_iso
                data_was_updated = True
            
            new_info = parse_jpg_filename(filename)
            if new_info:
                request_data['parsed_info'] = new_info
                data_was_updated = True

        else:
            parsed_info = parse_jpg_filename(filename)
            if not parsed_info:
                continue
            
            request_data = {
                'parsed_info':      parsed_info,
                'file_path':        current_path,
                'status':           'Pending Review',
                'notes':            [],
                'visible':          True,
                'creation_time':    creation_iso,
                'last_activity_on': creation_iso,
                'last_activity_by': 'System',
                'warnings':         [],
                'subscribers':      []
            }
            approval_data[filename] = request_data
            log_activity(f"APPROVAL: New request card detected: {filename}")
            data_was_updated = True
        
        request_data['filename'] = filename
        request_data['file_path_safe'] = current_path
        visible_requests.append(request_data)

    if data_was_updated: 
        save_approval_data(approval_data)
    
    visible_requests.sort(key=lambda x: x.get('last_activity_on', ''), reverse=True)
    return visible_requests


def scan_us_approvals():
    """
    Exactly the same as scan_and_prepare_approvals but only for modality 'US',
    from US_WATCH_FOLDERS.
    """
    approval_data = load_approval_data()
    data_updated = False
    jpgs = {}
    for folder in CONFIG.get('US_WATCH_FOLDERS', []):
        if not os.path.isdir(folder): continue
        for root, _, files in os.walk(folder):
            for fn in files:
                if fn.lower().endswith(('.jpg','jpeg')):
                     jpgs[fn] = os.path.join(root,fn)

    visible = []
    for fn, path in jpgs.items():
        # parse
        info = parse_jpg_filename(fn)
        if not info or info.get('modality') != 'US':
             continue
        ct = os.path.getctime(path)
        iso = datetime.datetime.fromtimestamp(ct).isoformat()
        if fn not in approval_data:
            approval_data[fn] = dict(
                parsed_info=info, file_path=path, status='Pending Review',
                notes=[], visible=True, creation_time=iso,
                last_activity_on=iso, last_activity_by='System',
                warnings=[], subscribers=[], filename=fn,
                file_path_safe=path
            )
            log_activity(f"US_APPROVAL: New US request card: {fn}")
            data_updated = True
        else:
            r = approval_data[fn]
            if r.get('file_path')!=path:
                r['file_path']=path; data_updated=True
            if r.get('creation_time')!=iso:
                r['creation_time']=iso; data_updated=True
            if r.get('parsed_info')!=info:
                r['parsed_info']=info; data_updated=True
        approval_data[fn]['filename']=fn
        approval_data[fn]['file_path_safe']=path
        if approval_data[fn].get('visible', True):
            visible.append(approval_data[fn])
    if data_updated:
        save_approval_data(approval_data)
    visible.sort(key=lambda x:x.get('last_activity_on',''), reverse=True)
    return visible


def parse_jpg_filename(filename):
    """
    Parses complex JPG filenames to extract patient information.
    This version specifically looks for a patient ID pattern of 5-6 digits, a space, and 2-3 letters.
    """
    KNOWN_MODALITIES = {'CT', 'DX', 'US', 'MG', 'MR', 'CR'}
    WARD_TRANSFORMS = {'AE': 'A&E'}
    
    # --- ADJUSTMENT: New regex to match the '123456 AB' pattern ---
    # This pattern finds a word boundary, 5 or 6 digits, one or more spaces,
    # 2 or 3 letters, and a final word boundary. It's case-insensitive.
    patient_id_pattern = re.compile(r'(\b\d{5,6}\s+[a-zA-Z]{2,3}\b)', re.IGNORECASE)

    name_part = os.path.splitext(filename)[0]
    match = patient_id_pattern.search(name_part)

    if not match:
        log_activity(f"FILENAME_PARSE_WARN: Could not find a valid patient ID pattern in '{filename}'.")
        return None

    patient_id_full = match.group(0).upper() # Capture the full ID, e.g., "123456 AB"
    id_start_index = match.start()
    
    patient_name = name_part[:id_start_index].strip().upper().replace("_", " ")
    if not patient_name:
        log_activity(f"FILENAME_PARSE_WARN: No name found before patient ID in '{filename}'.")
        return None

    post_id_part = name_part[match.end():].strip()
    tokens = post_id_part.split()
    
    modality = "Unknown"
    ward_parts = []

    for token in tokens:
        token_upper = token.upper()
        if token_upper in KNOWN_MODALITIES:
            modality = token_upper
        elif token_upper in WARD_TRANSFORMS:
            ward_parts.append(WARD_TRANSFORMS[token_upper])
        elif len(token) > 1: # Avoid single letter tokens unless they are a known modality/ward
             ward_parts.append(token)
    
    return {
        'patient_name':      patient_name,
        'patient_id':        patient_id_full,
        'patient_id_part':   patient_id_full, # Using the full captured ID here
        'ward_part':         ' '.join(ward_parts),
        'modality':          modality,
    }


def generate_docx_report_internal(study_data, template_path_override=None, output_dir=None):
    """
    Generates a DOCX report from a template, populating it with study data.
    This version includes robust filename sanitization.
    """
    template_path = template_path_override or CONFIG.get('DOCX_TEMPLATE_PATH')
    if not os.path.exists(template_path):
        log_activity(f"ERROR: DOCX template not found at {template_path}.")
        return None, "Template file not found."

    # Prepare all replacement values from the study_data dictionary
    patient_name = study_data.get('patient_name', 'N/A')
    patient_id = study_data.get('patient_id', 'N/A')
    accession = study_data.get('accession', 'N/A')
    dob = study_data.get('date_of_birth', 'N/A')
    sex = study_data.get('sex', 'N/A')
    exam_date = study_data.get('study_date_fmt', 'N/A')
    modality = study_data.get('modality', 'N/A')
    clinic_from = study_data.get('clinic_referred_from', 'N/A')
    body_part = study_data.get('study_desc', 'N/A')
    referring = study_data.get('referred_from', 'N/A')

    replacements = {
        '{Patient Name}': patient_name, '{Docket Number}': patient_id,
        '{Date of Birth}': dob, '{Accession Number}': accession,
        '{Study Description}': body_part, '{Referring Physician}': referring,
        '{Date of Exam}': exam_date, '{Clinic Referred From}': clinic_from,
        '{Body Part Done}': body_part, '{Modality Done}': modality,
        '{Sex}': "Male" if sex == "M" else "Female" if sex == "F" else sex,
        '{GENDER}': "Male" if sex == "M" else "Female" if sex == "F" else sex
    }

    try:
        doc = Document(template_path)
        # Populate the document template
        for p in doc.paragraphs:
            for r in p.runs:
                for k, v in replacements.items():
                    if k in r.text: r.text = r.text.replace(k, str(v))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p_cell in cell.paragraphs:
                        for r_cell in p_cell.runs:
                            for k, v in replacements.items():
                                if k in r_cell.text: r_cell.text = r_cell.text.replace(k, str(v))

        # --- SECURE FILENAME GENERATION ---
        # Use secure_filename on each component to remove invalid characters
        s_name = secure_filename(patient_name) or "NoName"
        s_id = secure_filename(patient_id) or "NoID"
        s_acc = secure_filename(accession) or "NoAccession"
        s_mod = secure_filename(modality) or "NoMod"
        s_desc = secure_filename(body_part) or "NoDesc"
        filename = f"{s_name}_{s_id}_{s_acc}_{s_mod}_{s_desc}.docx"

        # Save the populated document
        save_dir = output_dir or REPORTS_DIR
        os.makedirs(save_dir, exist_ok=True)
        out_path = os.path.join(save_dir, filename)
        doc.save(out_path)
        log_activity(f"DOCX_GEN: Created report -> {out_path}")
        return out_path, None

    except Exception as e:
        log_activity(f"DOCX_GEN_ERROR: {e}")
        return None, str(e)


def generate_and_save_mwl_docx(patient_data):
    modality = patient_data.get("modality")
    if not modality:
        return "Modality not found in patient data."

    # Special logic for DX modality
    if modality == "DX":
        referred_from = patient_data.get("referred_from", "").upper()
        if not (referred_from.startswith("H/C") or referred_from.startswith("HEALTH")):
            msg = f"Report for DX modality not generated because referrer '{patient_data.get('referred_from')}' does not meet criteria."
            log_mwl_activity(msg)
            return msg

    # Determine output directory
    output_folders = CONFIG.get("DOCX_OUTPUT_FOLDERS", {})
    base_output_dir = output_folders.get(modality, output_folders.get("General"))
    if not base_output_dir:
        return f"DOCX output directory for modality '{modality}' or 'General' not configured."

    # Create patient-specific subfolder
    safe_pname = re.sub(r"[^\w\- ]", "", patient_data.get("patient_name", "UnknownPatient")).strip().replace(" ", "_")
    safe_pid = re.sub(r"[^\w\- ]", "", patient_data.get("patient_id", "NoID")).strip().replace(" ", "_")
    patient_subfolder = f"{safe_pname}_{safe_pid}"
    final_output_dir = os.path.join(base_output_dir, patient_subfolder)

    # Prepare data for the generator function
    docx_data = {
        'patient_name': patient_data.get('patient_name'),
        'patient_id': patient_data.get('patient_id'),
        'accession': patient_data.get('accession_number'),
        'date_of_birth': format_pacs_date(patient_data.get('dob_yyyymmdd')),
        'sex': patient_data.get('sex'),
        'study_date_fmt': format_pacs_date(patient_data.get('study_date')),
        'modality': modality,
        'study_desc': patient_data.get('study_description'),
        'referred_from': patient_data.get('referred_from'),
        'clinic_referred_from': patient_data.get('referred_from') 
    }

    _, error = generate_docx_report_internal(docx_data, output_dir=final_output_dir)
    return error


# --- ADMIN & SETTINGS ROUTES ---
@app.route("/settings", methods=["GET", "POST"])
@permission_required('view_settings_page')
def settings():
    if request.method == 'POST':
        new_config = CONFIG.copy()
        form_keys_text = ['INSTITUTION_NAME', 'DICOM_ROOT', 'PACS_IP', 'PACS_AE_TITLE', 'LOCAL_AE_TITLE', 'SMTP_HOST', 'SMTP_SENDER_EMAIL', 'SMTP_USER', 'SMTP_PASSWORD', 'DOCX_TEMPLATE_PATH', 'MWL_AE_TITLE', 'DEFAULT_ACCESSION_PREFIX', 'DEFAULT_SCHEDULED_STATION_AE']
        form_keys_int = ['PACS_PORT', 'PACS_POLL_INTERVAL', 'APPROVAL_POLL_INTERVAL', 'SMTP_PORT', 'MWL_PORT']
        form_keys_list = ['JPG_WATCH_FOLDERS', 'US_WATCH_FOLDERS', 'RADIOLOGIST_EMAILS', 'ARCHIVING_EMAILS'] # Updated list
        for key in form_keys_text: new_config[key] = request.form.get(key)
        for key in form_keys_int:
            try: new_config[key] = int(request.form.get(key))
            except (ValueError, TypeError): flash(f"Invalid value for {key}. It must be a number.", "danger"); return redirect(url_for('settings'))
        for key in form_keys_list: new_config[key] = [line.strip() for line in request.form.get(key, '').splitlines() if line.strip()]

        # Handle checkboxes
        new_config['EMAIL_ENABLED'] = 'EMAIL_ENABLED' in request.form
        new_config['MWL_ENABLED'] = 'MWL_ENABLED' in request.form
        new_config['MWL_MPPS_ENABLED'] = 'MWL_MPPS_ENABLED' in request.form

        # Handle DOCX output folder dictionary
        new_config['DOCX_OUTPUT_FOLDERS'] = {}
        for mod in ['General', 'CT', 'DX', 'US', 'MG', 'MR']:
            new_config['DOCX_OUTPUT_FOLDERS'][mod] = request.form.get(f'DOCX_OUTPUT_FOLDERS_{mod}')

        # Handle User Role Permissions
        new_permissions = {"admin": ALL_PERMISSIONS} # Admin always has all permissions
        defined_roles = DEFAULT_CONFIG['USER_ROLE_PERMISSIONS'].keys()
        for role in defined_roles:
            if role == 'admin': continue
            new_permissions[role] = []
            for perm in ALL_PERMISSIONS:
                if f"perm_{role}_{perm}" in request.form:
                    new_permissions[role].append(perm)
        new_config['USER_ROLE_PERMISSIONS'] = new_permissions

        save_config(new_config)
        flash("Settings saved successfully! Some changes may require an application restart.", "success")
        log_activity(f"SETTINGS: Admin '{session['username']}' updated system configuration.")
        return redirect(url_for('settings'))
    return render_template_string(SETTINGS_TEMPLATE, config=CONFIG, ALL_PERMISSIONS=ALL_PERMISSIONS)

@app.route("/test_smtp", methods=["POST"])
@permission_required('test_smtp')
def test_smtp_settings():
    """Endpoint to handle sending a test email from the settings page."""
    form = request.form
    # Create a temporary config dict from the form data posted by the modal
    test_config = {
        "SMTP_HOST": form.get('SMTP_HOST_TEST'),
        "SMTP_PORT": form.get('SMTP_PORT_TEST'),
        "SMTP_SENDER_EMAIL": form.get('SMTP_SENDER_EMAIL_TEST'),
        "SMTP_USER": form.get('SMTP_USER_TEST'),
        "SMTP_PASSWORD": form.get('SMTP_PASSWORD_TEST')
    }
    recipient = form.get('test_email_recipient')

    if not all(test_config.values()) or not recipient:
        flash("One or more SMTP fields were missing from the test form.", "danger")
        return redirect(url_for('settings'))

    subject = f"RadTrac® SMTP Test - {datetime.datetime.now():%Y-%m-%d %H:%M:%S}"
    html_body = "<h1>SMTP Test Successful!</h1><p>If you have received this email, your SMTP settings are configured correctly in RadTrac®.</p>"

    try:
        # Use a thread to avoid blocking the UI while sending
        email_thread = threading.Thread(
            target=_send_email_worker,
            args=(subject, html_body, [recipient], test_config)
        )
        email_thread.start()
        log_activity(f"SMTP_TEST: User '{session['username']}' sent test email to '{recipient}'.")
        flash(f"Test email sent to {recipient}. Please check the inbox.", "success")
    except Exception as e:
        log_activity(f"SMTP_TEST_ERROR: {e}")
        flash(f"Failed to send test email: {e}", "danger")

    return redirect(url_for('settings'))


@app.route("/approval")
@permission_required('view_approval_tracker')
def approval_tracker():
    """Renders the approval tracker page shell. Data is loaded via API."""
    return render_template_string(APPROVAL_TEMPLATE)

@app.route("/us_approval")
@permission_required('view_us_approval_tracker')
def us_approval_tracker():
    """Render the Ultrasound-only tracker."""
    return render_template_string(US_APPROVAL_TEMPLATE, config=CONFIG)

@app.route("/api/us_approvals")
@permission_required('view_us_approval_tracker')
def api_us_approvals():
    """Return only modality=='US' approval requests."""
    all_reqs = scan_us_approvals()
    user_permissions = CONFIG.get('USER_ROLE_PERMISSIONS', {}).get(session.get('role'), [])
    all_users = {u: d for u, d in load_users().items()} # Get all users
    return jsonify(
        requests=all_reqs,
        permissions=user_permissions,
        all_users=all_users, # Add all users to the response
        current_user=session.get('username') # Add current user's username
    )

@app.route("/request_card")
@permission_required('view_approval_tracker')
def serve_request_card():
    file_path = unquote(request.args.get("path", ''))
    # Combine both CT and US watch folders for validation
    allowed_folders_ct = [os.path.abspath(f) for f in CONFIG.get('JPG_WATCH_FOLDERS', [])]
    allowed_folders_us = [os.path.abspath(f) for f in CONFIG.get('US_WATCH_FOLDERS', [])]
    allowed_folders = allowed_folders_ct + allowed_folders_us
    
    if os.path.isfile(file_path) and any(os.path.abspath(file_path).startswith(f) for f in allowed_folders):
        return send_file(file_path, mimetype='image/jpeg')
    log_activity(f"ACCESS_DENIED: User '{session.get('username')}' tried to access forbidden path: {file_path}")
    abort(404)

# --- APPROVAL TRACKER ROUTES ---
@app.route("/approval/add_note", methods=["POST"])
@permission_required('add_approval_notes')
def add_approval_note():
    filename, note_text = request.form.get('filename'), request.form.get('note_text', '').strip()
    if not filename or not note_text: flash("Cannot add an empty note.", "warning"); return redirect(url_for('approval_tracker'))
    approval_data = load_approval_data()
    if filename in approval_data:
        new_note = {'text': note_text, 'by': session['full_name'], 'on': datetime.datetime.now().isoformat()}
        approval_data[filename]['notes'].append(new_note)
        approval_data[filename]['last_activity_on'] = datetime.datetime.now().isoformat()
        approval_data[filename]['last_activity_by'] = session['full_name']
        save_approval_data(approval_data)
        log_activity(f"APPROVAL: User '{session['username']}' added note to '{filename}'.")
        
        patient_name = approval_data[filename]['parsed_info'].get('patient_name', 'N/A')
        subject = f"New Note Added for Request: {patient_name}"
        context = {
            "patient_name": patient_name,
            "user_full_name": session['full_name'],
            "note_text": note_text
        }
        _notify_subscribers(filename, subject, 'new_note', context)

        flash("Note added and subscribers notified.", "success")
    else: flash("Could not find the specified request.", "danger")
    return redirect(url_for('approval_tracker'))

@app.route("/approval/update_status", methods=["POST"])
@permission_required('manage_approvals')
def update_approval_status():
    filename, new_status = request.form.get('filename'), request.form.get('new_status')
    
    # This line is updated to validate the shorter "Info Needed" status
    if not filename or not new_status.startswith(('Pending Review', 'Info Needed', 'Approved', 'Scan Failed', 'Completed', 'Rejected')):
        flash("Invalid status selected.", "warning"); return redirect(url_for('approval_tracker'))

    approval_data = load_approval_data()
    if filename in approval_data:
        approval_data[filename]['status'] = new_status
        now_iso = datetime.datetime.now().isoformat()
        approval_data[filename]['last_activity_on'] = now_iso
        approval_data[filename]['last_activity_by'] = session['full_name']
        
        status_note = {'text': f"Status changed to: {new_status}", 'by': session['full_name'], 'on': now_iso}
        approval_data[filename]['notes'].append(status_note)

        save_approval_data(approval_data)
        log_activity(f"APPROVAL: User '{session['username']}' set status of '{filename}' to '{new_status}'.")
        
        patient_name = approval_data[filename]['parsed_info'].get('patient_name', 'N/A')
        subject = f"Status Update for Request: {patient_name}"
        context = {
            "patient_name": patient_name,
            "status": new_status,
            "user_full_name": session['full_name']
        }
        _notify_subscribers(filename, subject, 'status_update', context)

        flash(f"Status updated to '{new_status}' and subscribers notified.", "success")
    else: flash("Could not find the specified request.", "danger")
    return redirect(url_for('approval_tracker'))

@app.route("/approval/delete", methods=["POST"])
@permission_required('manage_approvals')
def delete_approval_entry():
    filename = request.form.get('filename')
    approval_data = load_approval_data()
    if filename in approval_data:
        approval_data[filename]['visible'] = False
        save_approval_data(approval_data)
        log_activity(f"APPROVAL: User '{session['username']}' hid entry '{filename}'.")
        flash(f"Request for '{filename}' has been hidden.", "info")
    else: flash("Could not find the specified request.", "danger")
    return redirect(url_for('approval_tracker'))


import base64

def build_email_html(template_name, context):
    """
    Builds the final HTML for an email by embedding the logo from a local path
    and rendering the content template.
    """
    logo_html = ''  # Default to an empty string
    try:
        # Use CONFIG.get() to safely access the LOGO_PATH
        logo_path = CONFIG.get('LOGO_PATH')
        if logo_path and os.path.exists(logo_path):
            with open(logo_path, 'rb') as f:
                b64_data = base64.b64encode(f.read()).decode('ascii')
                logo_source = f"data:image/png;base64,{b64_data}"
                # If the logo is found, create the full img tag with styling
                logo_html = f'<img src="{logo_source}" alt="Logo" style="max-height: 40px; vertical-align: middle; margin-right: 15px;"> '
        else:
            # Log if the LOGO_PATH key is missing from config.json or path is invalid
            log_activity("EMAIL LOGO INFO: 'LOGO_PATH' not found in config.json or path is invalid. Logo will be omitted.")

    except Exception as e:
        # Catch any other exceptions during logo processing
        log_activity(f"EMAIL LOGO ERROR: Could not embed logo. Error: {e}. Logo will be omitted.")

    content_tmpl = CONFIG["EMAIL_TEMPLATES"].get(template_name, "")
    email_content = content_tmpl.format(**context)
    
    base_html = CONFIG["EMAIL_TEMPLATES"]["base_template"]
    
    # This is the line to change. Ensure it uses app_logo_html.
    final_html = base_html.format(
    email_content=email_content,
    app_logo_url=logo_html,
    app_logo_html=logo_html   # satisfy the other placeholder
)
    
    return final_html

def format_activity_log_for_email(notes_list, creation_time_iso):
    """Formats the activity log of a request into a clean HTML table for emails."""
    html = """
    <table class="activity-table" width="100%" cellpadding="8" cellspacing="0" style="border-collapse: collapse; width: 100%; margin-top: 15px;">
        <tr style="background-color: #e9ecef;">
            <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">Date & Time</th>
            <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">Event</th>
            <th style="border: 1px solid #ddd; padding: 8px; text-align: left;">User</th>
        </tr>
    """
    events = []
    try:
        events.append({
            'timestamp': datetime.datetime.fromisoformat(creation_time_iso),
            'text': 'Request Card Uploaded',
            'by': 'System'
        })
    except (ValueError, TypeError):
        pass

    for note in notes_list:
        try:
            events.append({
                'timestamp': datetime.datetime.fromisoformat(note['on']),
                'text': note.get('text', 'N/A'),
                'by': note.get('by', 'N/A')
            })
        except (ValueError, TypeError):
            continue

    events.sort(key=lambda x: x['timestamp'], reverse=True)

    for event in events:
        time_str = event['timestamp'].strftime('%b %d, %Y at %I:%M %p')
        html += f"""
        <tr>
            <td style="border: 1px solid #ddd; padding: 8px;">{time_str}</td>
            <td style="border: 1px solid #ddd; padding: 8px;">{event['text']}</td>
            <td style="border: 1px solid #ddd; padding: 8px;">{event['by']}</td>
        </tr>
        """
    html += "</table>"
    return html

def _notify_subscribers(filename, subject, template_name, context):
    """Helper function to email all subscribers of an approval item using templates."""
    approval_data = load_approval_data()
    request_item = approval_data.get(filename)
    if not request_item: return

    subscriber_usernames = request_item.get('subscribers', [])
    if not subscriber_usernames: return
    
    all_users = load_users()
    recipients = []
    for username in subscriber_usernames:
        user_data = all_users.get(username)
        if user_data and user_data.get('notify_on_updates') and user_data.get('email'):
            recipients.append(user_data['email'])
    
    if recipients:
        try:
            html_body = build_email_html(template_name, context)
            log_activity(f"NOTIFY: Sending '{template_name}' update for '{filename}' to {len(recipients)} subscribers.")
            # Use a thread to avoid blocking the main application
            email_thread = threading.Thread(target=_send_email_worker, args=(subject, html_body, recipients, CONFIG))
            email_thread.start()
        except Exception as e:
            log_activity(f"NOTIFY_ERROR: Failed to send subscriber email for '{filename}'. Error: {e}")

@app.route("/approval/toggle_approval_subscription", methods=["POST"])
@permission_required('manage_own_subscriptions')
def toggle_approval_subscription():
    filename = request.form.get('filename')
    username = session['username']
    approval_data = load_approval_data()

    if filename in approval_data:
        if 'subscribers' not in approval_data[filename]:
            approval_data[filename]['subscribers'] = []

        if username in approval_data[filename]['subscribers']:
            approval_data[filename]['subscribers'].remove(username)
            flash("You will no longer receive notifications for this request.", "info")
        else:
            approval_data[filename]['subscribers'].append(username)
            flash("You will now receive email notifications for this request.", "success")
            
            # Send one-time summary email to new subscriber
            all_users = load_users()
            user_data = all_users.get(username)
            if user_data and user_data.get('email'):
                request_item = approval_data[filename]
                patient_name = request_item.get('parsed_info', {}).get('patient_name', 'N/A')
                subject = f"Subscription Confirmed for Request: {patient_name}"
                
                context = {
                    "user_full_name": user_data.get('full_name', username),
                    "patient_name": patient_name,
                    "patient_id": request_item.get('parsed_info', {}).get('patient_id_part', 'N/A'),
                    "status": request_item.get('status', 'N/A'),
                    "activity_history_table": format_activity_log_for_email(request_item.get('notes', []), request_item.get('creation_time'))
                }
                
                html_body = build_email_html('new_subscriber', context)
                email_thread = threading.Thread(target=_send_email_worker, args=(subject, html_body, [user_data['email']], CONFIG))
                email_thread.start()

        approval_data[filename]['last_activity_on'] = datetime.datetime.now().isoformat()
        save_approval_data(approval_data)
    else:
        flash("Could not find the specified request.", "danger")
    return redirect(url_for('approval_tracker'))

@app.route("/approval/subscribe_user", methods=["POST"])
@permission_required('manage_all_subscriptions')
def subscribe_user():
    filename = request.form.get('filename')
    username_to_subscribe = request.form.get('username_to_subscribe')
    if not username_to_subscribe:
        flash("You must select a user to add.", "warning")
        return redirect(url_for('approval_tracker'))

    approval_data = load_approval_data()
    if filename in approval_data:
        if 'subscribers' not in approval_data[filename]:
            approval_data[filename]['subscribers'] = []
        
        if username_to_subscribe not in approval_data[filename]['subscribers']:
            approval_data[filename]['subscribers'].append(username_to_subscribe)
            approval_data[filename]['last_activity_on'] = datetime.datetime.now().isoformat()
            save_approval_data(approval_data)
            flash(f"User subscribed successfully.", "success")

            # Send one-time summary email to new subscriber
            all_users = load_users()
            user_data = all_users.get(username_to_subscribe)
            if user_data and user_data.get('email'):
                request_item = approval_data[filename]
                patient_name = request_item.get('parsed_info', {}).get('patient_name', 'N/A')
                subject = f"You've been subscribed to a request for: {patient_name}"

                context = {
                    "user_full_name": user_data.get('full_name', username_to_subscribe),
                    "patient_name": patient_name,
                    "patient_id": request_item.get('parsed_info', {}).get('patient_id_part', 'N/A'),
                    "status": request_item.get('status', 'N/A'),
                    "activity_history_table": format_activity_log_for_email(request_item.get('notes', []), request_item.get('creation_time'))
                }
                
                html_body = build_email_html('new_subscriber', context)
                email_thread = threading.Thread(target=_send_email_worker, args=(subject, html_body, [user_data['email']], CONFIG))
                email_thread.start()
        else:
            flash(f"User is already subscribed.", "info")
    else:
        flash("Could not find the specified request.", "danger")
    return redirect(url_for('approval_tracker'))


@app.route("/approval/unsubscribe_user", methods=["POST"])
@permission_required('manage_all_subscriptions')
def unsubscribe_user():
    filename = request.form.get('filename')
    username_to_unsubscribe = request.form.get('username_to_unsubscribe')
    
    approval_data = load_approval_data()
    if filename in approval_data and 'subscribers' in approval_data[filename]:
        if username_to_unsubscribe in approval_data[filename]['subscribers']:
            approval_data[filename]['subscribers'].remove(username_to_unsubscribe)
            approval_data[filename]['last_activity_on'] = datetime.datetime.now().isoformat()
            save_approval_data(approval_data)
            flash("User unsubscribed.", "info")
    else:
        flash("Could not find the specified request or user.", "danger")
    return redirect(url_for('approval_tracker'))


# --- MWL & PRE-REGISTRATION ROUTES ---
@app.route("/approval/preregister", methods=["POST"])
@permission_required('preregister_patient')
def preregister_patient():
    form = request.form
    filename = form.get('filename') # Get the filename from the form
    scheduled_datetime_str = form.get('scheduled_datetime')

    # Stricter validation for required fields based on DB schema
    if not all([form.get('patient_name'), form.get('patient_id'), scheduled_datetime_str, filename]):
        return jsonify({'success': False, 'message': 'Missing required form data.'})

    dob_input = form.get('dob', '').strip()
    dob_yyyymmdd = None
    if dob_input:
        try:
            dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d/%m/%Y").strftime("%Y%m%d")
        except ValueError:
            try:
                dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d%m%Y").strftime("%Y%m%d")
            except ValueError:
                return jsonify({'success': False, 'message': 'Invalid Date of Birth format. Please use DD/MM/YYYY or DDMMYYYY.'})

    sched = scheduled_datetime_str.strip()
    if re.fullmatch(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", sched):
        sched += ":00"
    
    query = """
        INSERT INTO expected_patients (patient_name, patient_id, dob_yyyymmdd, sex, study_description, modality, referred_from, scheduled_datetime, requesting_physician)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """
    params = (
        form['patient_name'], form['patient_id'], dob_yyyymmdd, form.get('sex'),
        form.get('study_description'), form.get('modality'), form.get('referred_from'),
        sched, form.get('requesting_physician') or form.get('referred_from')
    )
    
    result = mwl_db_execute(query, params, commit=True)
    
    if result:
        log_mwl_activity(f"User '{session['username']}' pre-registered patient '{form['patient_name']}'")
        
        # --- NEW: Add note to the approval request's activity log ---
        try:
            approval_data = load_approval_data()
            if filename in approval_data:
                # Format the date for the note
                dt_obj = datetime.datetime.strptime(sched, "%Y-%m-%d %H:%M:%S")
                formatted_dt = dt_obj.strftime('%b %d, %Y @ %I:%M %p')
                
                note_text = f"Appointment set for: {formatted_dt}"
                new_note = {'text': note_text, 'by': session['full_name'], 'on': datetime.datetime.now().isoformat()}
                
                approval_data[filename].setdefault('notes', []).append(new_note)
                approval_data[filename]['last_activity_on'] = datetime.datetime.now().isoformat()
                approval_data[filename]['last_activity_by'] = session['full_name']
                
                save_approval_data(approval_data)
                log_activity(f"APPROVAL_NOTE: Added pre-registration note to '{filename}'.")
        except Exception as e:
            log_activity(f"APPROVAL_NOTE_ERROR: Failed to add pre-reg note for '{filename}'. Error: {e}")
        # --- END OF NEW CODE ---

        return jsonify({'success': True, 'message': f"Patient {form['patient_name']} added to the expected list."})
    else:
        return jsonify({'success': False, 'message': 'A database error occurred. The patient might already be on the expected list.'})

@app.route("/api/expected_patients")
@permission_required('manage_mwl_entries')
def api_expected_patients():
    today_start = datetime.datetime.now().strftime("%Y-%m-%d 00:00:00")
    today_end = datetime.datetime.now().strftime("%Y-%m-%d 23:59:59")
    
    query = """
        SELECT id, patient_name, patient_id, scheduled_datetime, modality
        FROM expected_patients
        WHERE status = 'pending' AND scheduled_datetime BETWEEN ? AND ?
        ORDER BY scheduled_datetime ASC
    """
    patients = mwl_db_execute(query, (today_start, today_end), fetchall=True)
    
    if patients:
        patient_list = []
        for p in patients:
            p_dict = dict(p)
            # Format time for display
            try:
                p_dict['scheduled_time'] = datetime.datetime.strptime(p_dict['scheduled_datetime'], "%Y-%m-%d %H:%M:%S").strftime("%H:%M")
            except (ValueError, TypeError):
                 p_dict['scheduled_time'] = 'N/A'
            patient_list.append(p_dict)
        return jsonify(patient_list)
    return jsonify([])

@app.route("/api/get_expected_patient/<int:patient_id>")
@permission_required('manage_mwl_entries')
def api_get_expected_patient(patient_id):
    record = mwl_db_execute("SELECT * FROM expected_patients WHERE id = ?", (patient_id,), fetchone=True)
    if record:
        record_dict = dict(record)
        # Format DOB for form
        try:
            record_dict['dob'] = datetime.datetime.strptime(record_dict['dob_yyyymmdd'], "%Y%m%d").strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            record_dict['dob'] = ''
        return jsonify(record_dict)
    return jsonify({'error': 'Patient not found'})

@app.route('/api/expected_patient/cancel/<int:patient_id>', methods=['DELETE'])
@permission_required('manage_mwl_entries')
def api_cancel_expected_patient(patient_id):
    # First, get the patient name for logging before deleting
    record = mwl_db_execute("SELECT patient_name FROM expected_patients WHERE id = ?", (patient_id,), fetchone=True)
    if not record:
        return jsonify({'success': False, 'error': 'Patient not found on the expected list.'})

    patient_name = record['patient_name']
    
    result = mwl_db_execute("DELETE FROM expected_patients WHERE id = ?", (patient_id,), commit=True)
    
    if result is not None:
        log_mwl_activity(f"User '{session['username']}' cancelled expected patient: {patient_name} (ID: {patient_id})")
        return jsonify({'success': True, 'message': f'Registration for {patient_name} has been cancelled.'})
    else:
        log_mwl_activity(f"DB_ERROR: Failed to cancel expected patient with ID: {patient_id}")
        return jsonify({'success': False, 'error': 'Failed to cancel the registration due to a database error.'})

@app.route("/mwl")
@permission_required('view_mwl_server')
def mwl_server():
    """Renders the MWL server status page."""
    return render_template_string(MWL_TEMPLATE, config=CONFIG)

@app.route("/api/mwl_worklist")
@permission_required('view_mwl_server')
def api_mwl_worklist():
    """Returns all current worklist entries from the MWL database."""
    worklist = mwl_db_execute("SELECT * FROM patient_records ORDER BY created_at DESC", fetchall=True)
    if worklist:
        return jsonify([dict(row) for row in worklist])
    return jsonify([])

@app.route("/api/mwl_activity_log")
@permission_required('view_mwl_server')
def api_mwl_activity_log():
    """Returns the content of the MWL activity log file."""
    try:
        if os.path.exists(MWL_ACTIVITY_LOG_FILE):
            with open(MWL_ACTIVITY_LOG_FILE, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                return "".join(lines[-100:])
        return "Log file not created yet."
    except Exception as e:
        return f"Error reading log file: {e}"

@app.route("/mwl/add_patient", methods=["POST"])
@permission_required('manage_mwl_entries')
def add_mwl_patient():
    form_data = request.form
    is_check_only = form_data.get('check_only') == 'true'

    required_fields = ['patient_name', 'patient_id', 'accession_number', 'dob', 'sex', 'study_description', 'referred_from', 'requesting_physician']
    if not all(form_data.get(field) for field in required_fields):
        return jsonify({'success': False, 'error': 'All fields are required.'})

    dob_input = form_data['dob'].strip()
    dob_yyyymmdd = None
    try:
        dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d/%m/%Y").strftime("%Y%m%d")
    except ValueError:
        if re.fullmatch(r"\d{8}", dob_input):
            try:
                dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d%m%Y").strftime("%Y%m%d")
            except ValueError:
                dob_yyyymmdd = None
    if not dob_yyyymmdd:
        return jsonify({'success': False, 'error': "Date of Birth must be either DD/MM/YYYY or DDMMYYYY."})

    modality = get_modality_from_accession(form_data['accession_number'])
    if not modality:
        return jsonify({'success': False, 'error': 'Invalid Accession Number format. Cannot determine modality.'})
    
    if form_data.get('bypass_duplicate_check') != 'true':
        duplicate_warning = check_duplicate_mwl_record(form_data['patient_name'], form_data['patient_id'], form_data['accession_number'])
        if duplicate_warning:
            return jsonify({'success': False, 'duplicate_warning': duplicate_warning})

    if is_check_only:
        return jsonify({'success': True})

    now = datetime.datetime.now()
    patient_data_for_db = (
        form_data['patient_name'], form_data['patient_id'], form_data['accession_number'].upper(),
        dob_yyyymmdd, form_data['sex'], now.strftime("%Y%m%d"), now.strftime("%H%M%S"),
        form_data['study_description'], form_data['referred_from'], modality,
        form_data['requesting_physician'], form_data['accession_number'].upper(), 
        CONFIG.get("DEFAULT_SCHEDULED_STATION_AE", "ANY_MODALITY")
    )
    
    query = '''
        INSERT INTO patient_records
        (patient_name, patient_id, accession_number, dob_yyyymmdd, sex,
         study_date, study_time, study_description, referred_from, modality,
         requesting_physician, requested_procedure_id, scheduled_station_ae_title)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    '''
    
    result = mwl_db_execute(query, patient_data_for_db, commit=True)
    if result is not None:
        log_mwl_activity(f"User '{session['username']}' added new patient: {form_data['patient_name']} (Acc: {form_data['accession_number']})")
        
        # If this was from an expected patient, mark them as registered
        expected_id = form_data.get('expected_id')
        if expected_id:
            mwl_db_execute("UPDATE expected_patients SET status = 'registered' WHERE id = ?", (expected_id,), commit=True)
            log_mwl_activity(f"Marked expected patient ID {expected_id} as registered.")

        patient_data_for_docx = {key: form_data[key] for key in form_data}
        patient_data_for_docx.update({
            "modality": modality, "dob_yyyymmdd": dob_yyyymmdd, "study_date": now.strftime("%Y%m%d")
        })
        docx_error = generate_and_save_mwl_docx(patient_data_for_docx)
        
        if docx_error:
            return jsonify({'success': True, 'message': f"Patient added to worklist. DOCX Report: {docx_error}"})
        else:
            return jsonify({'success': True, 'message': "Patient added and report generated successfully!"})

    else:
        return jsonify({'success': False, 'error': 'Database error occurred. The Accession Number may already exist.'})

@app.route('/api/mwl/delete/<int:record_id>', methods=['DELETE'])
@permission_required('manage_mwl_entries')
def api_mwl_delete(record_id):
    record = mwl_db_execute("SELECT patient_name, accession_number FROM patient_records WHERE id = ?", (record_id,), fetchone=True)
    if not record:
        return jsonify({'success': False, 'error': 'Record not found.'})
    
    result = mwl_db_execute("DELETE FROM patient_records WHERE id = ?", (record_id,), commit=True)
    if result is not None:
        log_mwl_activity(f"User '{session['username']}' deleted patient: {record['patient_name']} (Acc: {record['accession_number']})")
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'error': 'Failed to delete record from database.'})

@app.route('/api/mwl/record/<int:record_id>')
@permission_required('manage_mwl_entries')
def api_mwl_get_record(record_id):
    record = mwl_db_execute("SELECT * FROM patient_records WHERE id=?", (record_id,), fetchone=True)
    if not record:
        return jsonify({'error': 'Record not found'})
    
    record_dict = dict(record)
    try:
        record_dict['dob'] = datetime.datetime.strptime(record_dict['dob_yyyymmdd'], "%Y%m%d").strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        record_dict['dob'] = ''
        
    return jsonify(record_dict)

@app.route('/api/mwl/update/<int:record_id>', methods=['POST'])
@permission_required('manage_mwl_entries')
def api_mwl_update(record_id):
    form_data = request.form
    
    dob_input = form_data['dob'].strip()
    dob_yyyymmdd = None
    try:
        dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d/%m/%Y").strftime("%Y%m%d")
    except ValueError:
        if re.fullmatch(r"\d{8}", dob_input):
            try:
                dob_yyyymmdd = datetime.datetime.strptime(dob_input, "%d%m%Y").strftime("%Y%m%d")
            except ValueError:
                dob_yyyymmdd = None
    if not dob_yyyymmdd:
        return jsonify({'success': False, 'error': 'Invalid Date of Birth format. Use DD/MM/YYYY or DDMMYYYY.'})
        
    query = """
        UPDATE patient_records SET
        patient_name=?, patient_id=?, dob_yyyymmdd=?, sex=?, study_description=?,
        referred_from=?, requesting_physician=?
        WHERE id=?
    """
    params = (
        form_data['patient_name'], form_data['patient_id'], dob_yyyymmdd, form_data['sex'],
        form_data['study_description'], form_data['referred_from'], form_data['requesting_physician'],
        record_id
    )
    
    result = mwl_db_execute(query, params, commit=True)
    if result is not None:
        log_mwl_activity(f"User '{session['username']}' updated record for Acc: {form_data['accession_number']}")
        return jsonify({'success': True, 'message': 'Patient updated successfully!'})
    else:
        return jsonify({'success': False, 'error': 'Database error during update.'})

# --- ASYNCHRONOUS DOWNLOAD IMPLEMENTATION ---

def create_study_zip_thread(accession, task_id):
    """
    This function runs in a background thread to create the zip file.
    **NEW**: It now copies the database to a temp file to avoid file lock issues.
    """
    log_activity(f"[THREAD_TASK {task_id}] Starting for Accession: {accession}")
    temp_db_path = None  # Initialize here to ensure it's available in 'finally'
    try:
        db_path = os.path.join(CONFIG.get('DICOM_ROOT', ''), '.Database', 'PacsDBv2.db')
        if not os.path.exists(db_path):
            log_activity(f"[THREAD_TASK {task_id}] ERROR: Live database not found at path.")
            with tasks_lock:
                DOWNLOAD_TASKS[task_id] = {'state': 'FAILURE', 'status': 'Server config error: DB not found.'}
            return

        # --- NEW: Copy database to a temporary location to bypass file locks ---
        temp_dir = os.path.join(SCRIPT_DIR, 'temp_reports') # Using a known temporary directory
        os.makedirs(temp_dir, exist_ok=True)
        temp_db_path = os.path.join(temp_dir, f"{task_id}.db")
        log_activity(f"[THREAD_TASK {task_id}] Copying live DB to temporary path: {temp_db_path}")
        shutil.copy2(db_path, temp_db_path)
        # --------------------------------------------------------------------

        log_activity(f"[THREAD_TASK {task_id}] Connecting to temporary database...")
        conn = sqlite3.connect(temp_db_path, timeout=5)
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT I.IMAGE_PATH_NAME
            FROM IMAGE I JOIN SERIES S ON I.SERIES_PRKEY = S.SERIES_PRKEY
            JOIN STUDY ST ON S.STUDY_PRKEY = ST.STUDY_PRKEY
            WHERE ST.ACCESSION_NUMBER = ?
        """, (accession,))
        rows = cursor.fetchall()
        conn.close()
        log_activity(f"[THREAD_TASK {task_id}] Found {len(rows)} image paths in temporary database.")

        if not rows:
            with tasks_lock:
                DOWNLOAD_TASKS[task_id] = {'state': 'FAILURE', 'status': 'No images found for this accession.'}
            return

        files_to_zip = []
        dicom_root = CONFIG.get('DICOM_ROOT')
        for (img_rel_path,) in rows:
            img_rel_path_os = os.path.join(*img_rel_path.split('/'))
            img_abs_path = os.path.join(dicom_root, img_rel_path_os)
            
            if os.path.isfile(img_abs_path):
                arcname = img_rel_path.replace('\\', '/')
                files_to_zip.append((arcname, img_abs_path))
        
        log_activity(f"[THREAD_TASK {task_id}] Found {len(files_to_zip)} existing files to zip.")

        if not files_to_zip:
             with tasks_lock:
                DOWNLOAD_TASKS[task_id] = {'state': 'FAILURE', 'status': 'Image files listed in DB but not found on disk.'}
             return

        with tasks_lock:
            DOWNLOAD_TASKS[task_id]['state'] = 'PROGRESS'
            DOWNLOAD_TASKS[task_id]['status'] = f'Zipping {len(files_to_zip)} files...'

        safe_accession = re.sub(r'[^a-zA-Z0-9_-]', '', accession)
        output_filename = f"{safe_accession}_{task_id[:8]}.zip"
        output_path = os.path.join(DOWNLOAD_DIR, output_filename)
        log_activity(f"[THREAD_TASK {task_id}] Creating zip file at: {output_path}")

        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for arcname, file_path in files_to_zip:
                zf.write(file_path, arcname=arcname)
    
        with tasks_lock:
            DOWNLOAD_TASKS[task_id] = {'state': 'SUCCESS', 'status': 'Complete!', 'filename': output_filename}
        log_activity(f"[THREAD_TASK {task_id}] Successfully created zip file.")

    except Exception as e:
        error_message = f"An unexpected error occurred: {e}"
        log_activity(f"[THREAD_TASK {task_id}] FATAL ERROR: {error_message}")
        with tasks_lock:
            DOWNLOAD_TASKS[task_id] = {'state': 'FAILURE', 'status': error_message}
    finally:
        # --- NEW: Cleanup the temporary database file ---
        if temp_db_path and os.path.exists(temp_db_path):
            try:
                os.remove(temp_db_path)
                log_activity(f"[THREAD_TASK {task_id}] Cleaned up temporary database file.")
            except Exception as e:
                log_activity(f"[THREAD_TASK {task_id}] ERROR: Could not clean up temp db file {temp_db_path}: {e}")
        # ---------------------------------------------

@app.route('/download/start/<accession>')
@permission_required('download_images')
def download_start(accession):
    """
    Starts the background zip creation thread and returns a task ID.
    """
    task_id = str(uuid.uuid4())
    
    with tasks_lock:
        DOWNLOAD_TASKS[task_id] = {'state': 'PENDING', 'status': 'Task is waiting to start...'}
    
    thread = threading.Thread(target=create_study_zip_thread, args=(accession, task_id))
    thread.daemon = True
    thread.start()
    
    log_activity(f"DOWNLOAD_TASK: User '{session['username']}' started zip thread {task_id} for Acc# {accession}.")
    return jsonify({'task_id': task_id}), 202

@app.route('/download/status/<task_id>')
@permission_required('download_images')
def download_status(task_id):
    """
    Checks the status of the background task from our global dictionary.
    """
    with tasks_lock:
        task = DOWNLOAD_TASKS.get(task_id, {})
    
    response = {
        'state': task.get('state', 'FAILURE'),
        'status': task.get('status', 'Unknown task ID. The job may have been lost.'),
        'filename': task.get('filename')
    }
    return jsonify(response)
    
@app.route('/download/get/<filename>')
@permission_required('download_images')
def get_download(filename):
    """
    Serves the completed (static) zip file.
    """
    if '..' in filename or filename.startswith('/'):
        abort(400, "Invalid filename.")
    
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        abort(404, "File not found. It may have expired or failed to create.")
    
    log_activity(f"DOWNLOAD_SERVE: User '{session['username']}' is downloading {filename}.")
    return send_file(file_path, as_attachment=True)

# REPLACE your current function with this one
# =================================================================
def process_report_upload_in_background(accession, save_path, user_full_name, patient_name):
    """
    Handles all post-upload tasks in a background thread to avoid blocking.
    """
    log_activity(f"BACKGROUND_TASK: Thread started for Acc#{accession}.")
    try:
        # Update metadata file
        try:
            with report_metadata_lock:
                report_metadata[accession] = save_path
                save_report_metadata()
            log_activity(f"BACKGROUND_TASK: Successfully updated metadata for Acc#{accession}")
        except Exception as e:
            log_activity(f"BACKGROUND_TASK_ERROR: Could not update metadata.json: {e}")
            return # Stop if this fails

        # Update live data cache
        try:
            with pacs_cache_lock:
                study_found_in_cache = False
                for study in pacs_data_cache.get('data', []):
                    if study.get('accession') == accession:
                        study['report_path'] = save_path
                        study_found_in_cache = True
                        break
            if study_found_in_cache:
                log_activity(f"BACKGROUND_TASK: Successfully updated live cache for Acc#{accession}")
            else:
                log_activity(f"BACKGROUND_TASK_WARN: Did not find Acc#{accession} in the live cache to update.")

        except Exception as e:
            log_activity(f"BACKGROUND_TASK_ERROR: Failed to update live cache for Acc#{accession}. Error: {e}")

        # Send notifications
        notify_report_action("UPLOADED", user_full_name, accession, patient_name, save_path)
    
    finally:
        log_activity(f"BACKGROUND_TASK: Thread finished for Acc#{accession}.")


# --- NEW REPORT MANAGEMENT ROUTES ---
# PASTE THIS CORRECTED BLOCK IN ITS PLACE
# =================================================================

@app.route('/upload_report', methods=['POST'])
@permission_required('upload_reports')
def upload_report():
    if 'report_file' not in request.files:
        flash('No file part in the request.', 'danger')
        return redirect(url_for('dashboard'))

    file = request.files['report_file']
    accession = request.form.get('accession', '').strip()
    patient_name = request.form.get('patient_name', 'N/A').strip()
    patient_id = request.form.get('patient_id', 'NoID').strip()

    if not file or file.filename == '' or not accession:
        flash('Missing PDF file or accession number.', 'danger')
        return redirect(url_for('dashboard'))

    if not file.filename.lower().endswith('.pdf'):
        flash('Invalid file type. Only PDF is allowed.', 'danger')
        return redirect(url_for('dashboard'))

    # Create the folder path
    safe_patient_folder_name = secure_filename(f"{patient_name}_{patient_id}")
    patient_folder = os.path.join(REPORTS_DIR, safe_patient_folder_name)
    accession_folder = os.path.join(patient_folder, secure_filename(accession))
    os.makedirs(accession_folder, exist_ok=True)
    
    # Create the file path and save the file
    original_filename = secure_filename(file.filename)
    timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
    new_filename = f"{accession}_{timestamp}_{original_filename}"
    save_path = os.path.join(accession_folder, new_filename)
    
    try:
        file.save(save_path)
        log_activity(f"REPORT_UPLOAD: Saved report to {save_path}")
    except Exception as e:
        log_activity(f"REPORT_UPLOAD_ERROR: Failed to save report: {e}")
        flash('Failed to save the report. Please try again.', 'danger')
        return redirect(url_for('dashboard'))

    # Start all slow tasks in the background and respond to the user immediately.
    background_thread = threading.Thread(
        target=process_report_upload_in_background,
        args=(accession, save_path, session.get('full_name'), patient_name),
        daemon=True
    )
    background_thread.start()

    flash('Report upload started. It will appear on the dashboard shortly.', 'success')
    return redirect(url_for('dashboard'))


@app.route('/delete_report', methods=['POST'])
@permission_required('delete_reports')
def delete_report():
    accession = request.form.get('accession')
    report_path = request.form.get('report_path')
    patient_name = request.form.get('patient_name', 'N/A')

    if not all([accession, report_path]):
        flash('Missing information to delete the report.', 'danger')
        return redirect(url_for('dashboard'))

    # Security check: Ensure the file path is within the allowed directory
    if not os.path.abspath(report_path).startswith(os.path.abspath(REPORTS_DIR)):
        log_activity(f"SECURITY_FAIL: User '{session['full_name']}' attempted to delete report outside of scope: {report_path}")
        flash('Invalid report path.', 'danger')
        abort(403)

    if os.path.exists(report_path):
        filename = os.path.basename(report_path)
        deleted_path = os.path.join(DELETED_REPORTS_DIR, f"{datetime.datetime.now():%Y%m%d%H%M%S}_{filename}")
        
        try:
            shutil.move(report_path, deleted_path)

            # Update metadata
            with report_metadata_lock:
                if accession in report_metadata:
                    del report_metadata[accession]
                    save_report_metadata()

            log_activity(f"REPORT: User '{session['full_name']}' deleted report '{filename}' for Acc# {accession}. Moved to deleted folder.")
            flash('Report has been deleted and archived.', 'info')

            # Trigger archive notification
            notification_thread = threading.Thread(
                target=notify_report_action,
                args=("DELETED", session['full_name'], accession, patient_name, deleted_path)
            )
            notification_thread.start()

        except Exception as e:
            log_activity(f"REPORT_DELETE_ERROR: Failed to move report for Acc# {accession}. Error: {e}")
            flash('An error occurred while deleting the report.', 'danger')
    else:
        flash('Report file not found. It may have already been deleted.', 'warning')
        # Also clean up metadata if file is missing
        with report_metadata_lock:
            if accession in report_metadata:
                del report_metadata[accession]
                save_report_metadata()

    return redirect(url_for('dashboard'))


def notify_report_action(action, user_full_name, accession, patient_name, file_path):
    """
    Sends two kinds of emails:
      1) ALWAYS: to ARCHIVING_EMAILS with the PDF attached using 'report_archive' template.
      2) IF action == 'UPLOADED': also to any approval-tracker subscribers
         using 'report_uploaded_subscriber' template (no attachment).
    """
    if not CONFIG.get('EMAIL_ENABLED'):
        log_activity("EMAIL_DISABLED: Skipping report notification.")
        return

    # 1) Archival notification (This part is correct)
    arch_emails = CONFIG.get('ARCHIVING_EMAILS', [])
    if arch_emails:
        subject = f"Report {action.title()} for {patient_name} (Acc#: {accession})"
        context = {
            "action": action.title(),
            "action_past_tense": "Uploaded" if action=="UPLOADED" else "Deleted",
            "patient_name": patient_name,
            "accession_number": accession,
            "user_full_name": user_full_name,
            "timestamp": datetime.datetime.now().strftime('%Y%m-%d %H:%M:%S')
        }
        try:
            html = build_email_html('report_archive', context)
            _send_email_worker(subject, html, arch_emails, CONFIG, attachments=[file_path])
            log_activity(f"NOTIFY_ARCHIVE: Sent archival email for Acc#{accession} to {len(arch_emails)} addresses.")
        except Exception as e:
            log_activity(f"NOTIFY_ARCHIVE_ERROR: {e}")

    # 2) Subscriber notification (only upon UPLOAD) - CORRECTED LOGIC
    if action == "UPLOADED":
        approval_data = load_approval_data()
        matching_filename = None

        # Search for the approval request by matching the accession number from the PACS data
        # with the patient ID in the approval request's parsed info. This assumes the
        # PACS accession number is used as the patient ID on the request card.
        for fn, req in approval_data.items():
            pacs_study_info = next((study for study in pacs_data_cache.get('data', []) if study.get('accession') == accession), None)
            if pacs_study_info and req.get('parsed_info', {}).get('patient_id') == pacs_study_info.get('patient_id'):
                matching_filename = fn
                break
        
        if matching_filename:
            subject = f"Report Uploaded for {patient_name}"
            context = {
                "patient_name": patient_name,
                "accession_number": accession,
                "user_full_name": user_full_name,
            }
            _notify_subscribers(matching_filename, subject, 'report_uploaded_subscriber', context)
        else:
            log_activity(f"NOTIFY_SUBSCRIBER_WARN: Could not find matching approval request for patient '{patient_name}' to notify subscribers.")

# --- NEW RESOURCES PAGE ROUTES ---

@app.route('/resources')
@permission_required('view_resources_page')
def resources():
    try:
        # List files, filtering out any subdirectories or hidden files
        files = [f for f in os.listdir(RESOURCES_DIR) if os.path.isfile(os.path.join(RESOURCES_DIR, f))]
        files.sort()
    except OSError:
        files = []
        flash("Could not read the resources directory.", "danger")
    return render_template_string(RESOURCES_TEMPLATE, files=files)

@app.route('/upload_resource', methods=['POST'])
@permission_required('manage_resources')
def upload_resource():
    if 'resource_file' not in request.files:
        flash('No file part in the request.', 'danger')
        return redirect(url_for('resources'))
    
    file = request.files['resource_file']
    if file.filename == '':
        flash('No selected file.', 'warning')
        return redirect(url_for('resources'))

    if file:
        filename = secure_filename(file.filename)
        save_path = os.path.join(RESOURCES_DIR, filename)
        file.save(save_path)
        log_activity(f"RESOURCES: User '{session['full_name']}' uploaded resource file '{filename}'.")
        flash(f"File '{filename}' uploaded successfully.", "success")

    return redirect(url_for('resources'))

@app.route('/report')
@permission_required('view_dashboard')
def serve_report():
    """
    Securely serves a report PDF from the uploaded reports directory.

    Adds a fallback step: if the exact file isn't found, try swapping
    spaces ↔ underscores in the basename to locate the file.
    """
    raw = request.args.get("path", "")
    if not raw:
        abort(400, "No report path specified.")

    # URL-decode
    file_path = unquote(raw)

    # Absolute paths
    reports_dir_abs = os.path.abspath(REPORTS_DIR)
    requested_abs = os.path.abspath(file_path)

    # 1) Security check: path must live under REPORTS_DIR
    if not requested_abs.startswith(reports_dir_abs):
        log_activity(f"SECURITY_FAIL: User '{session.get('username')}' tried to access forbidden path: {file_path}")
        abort(403)

    # 2) If exact file exists, serve immediately
    if os.path.isfile(requested_abs):
        return send_file(requested_abs, mimetype='application/pdf')

    # 3) Fallback #1: swap spaces → underscores in basename
    parent = os.path.dirname(requested_abs)
    base = os.path.basename(requested_abs)

    alt1 = base.replace(' ', '_')
    alt1_path = os.path.join(parent, alt1)
    if os.path.isfile(alt1_path):
        log_activity(f"REPORT_SERVE_FALLBACK: Serving underscore‐mapped file for {base}")
        return send_file(alt1_path, mimetype='application/pdf')

    # 4) Fallback #2: swap underscores → spaces
    alt2 = base.replace('_', ' ')
    alt2_path = os.path.join(parent, alt2)
    if os.path.isfile(alt2_path):
        log_activity(f"REPORT_SERVE_FALLBACK: Serving space‐mapped file for {base}")
        return send_file(alt2_path, mimetype='application/pdf')

    # 5) Still not found
    log_activity(f"REPORT_SERVE_FAIL: File not found at any tested path: {requested_abs}")
    abort(404)

@app.route('/download_resource/<path:filename>')
@permission_required('view_resources_page')
def download_resource(filename):
    # Use send_from_directory for security
    return send_file(os.path.join(RESOURCES_DIR, filename), as_attachment=True)

@app.route('/delete_resource/<path:filename>', methods=['POST'])
@permission_required('manage_resources')
def delete_resource(filename):
    try:
        # Security check to prevent path traversal
        safe_filename = secure_filename(filename)
        if safe_filename != filename:
            abort(400, "Invalid filename.")

        file_path = os.path.join(RESOURCES_DIR, safe_filename)
        if os.path.exists(file_path):
            os.remove(file_path)
            log_activity(f"RESOURCES: User '{session['full_name']}' deleted resource file '{safe_filename}'.")
            flash(f"File '{safe_filename}' has been deleted.", "info")
        else:
            flash("File not found.", "warning")
    except Exception as e:
        log_activity(f"RESOURCES_ERROR: Could not delete file '{filename}'. Error: {e}")
        flash("An error occurred while deleting the file.", "danger")
        
    return redirect(url_for('resources'))

# --- NEW QUIZ ROUTES ---

@app.route('/quiz/data')
@permission_required('view_resources_page')
def quiz_data_endpoint():
    # Load leaderboard and sort by score
    leaderboard = sorted(load_leaderboard(), key=lambda x: x.get('score', 0), reverse=True)
    
    # Check user's last attempt
    attempts = load_attempts()
    username = session['username']
    user_can_play = True
    cooldown_message = ""
    last_attempt_ts = attempts.get(username)

    if last_attempt_ts:
        last_attempt_dt = datetime.datetime.fromtimestamp(last_attempt_ts)
        cooldown_ends = last_attempt_dt + datetime.timedelta(hours=24)
        now = datetime.datetime.now()
        if now < cooldown_ends:
            user_can_play = False
            time_left = cooldown_ends - now
            hours, remainder = divmod(time_left.seconds, 3600)
            minutes, _ = divmod(remainder, 60)
            cooldown_message = f"You can play again in {hours} hours and {minutes} minutes."
            
    # Load resource files list
    try:
        files = [f for f in os.listdir(RESOURCES_DIR) if os.path.isfile(os.path.join(RESOURCES_DIR, f))]
        files.sort()
    except OSError:
        files = []

    return jsonify({
        "leaderboard": leaderboard[:6], # Top 6
        "user_can_play": user_can_play,
        "cooldown_message": cooldown_message,
        "files": files
    })


@app.route('/quiz/start', methods=['POST'])
@permission_required('view_resources_page')
def quiz_start():
    if not quiz_data:
        return jsonify({"message": "The quiz is currently unavailable. The question bank could not be loaded."}), 503

    attempts = load_attempts()
    username = session['username']
    last_attempt_ts = attempts.get(username)

    if last_attempt_ts:
        if (time.time() - last_attempt_ts) < 86400: # 24 hours
            return jsonify({"message": "You have already attempted the quiz in the last 24 hours."}), 429

    # Select 15 random questions
    random.shuffle(quiz_data)
    num_questions = min(15, len(quiz_data))
    selected_questions = quiz_data[:num_questions]
    
    # Store questions and start time in server-side memory
    quiz_sessions[username] = {
        'questions': selected_questions,
        'start_time': time.time()
    }
    
    # Return only questions and options to the client
    client_questions = [{
        "question": q["question"],
        "options": q["options"],
        "answer": q["answer"] # Sending answer for immediate feedback on client
    } for q in selected_questions]

    return jsonify({"questions": client_questions})

@app.route('/quiz/submit', methods=['POST'])
@permission_required('view_resources_page')
def quiz_submit():
    submission = request.json
    username = session['username']
    quiz_info = quiz_sessions.get(username)
    
    if not quiz_info:
        return jsonify({"message": "Invalid or expired quiz session. Please start over."}), 400

    server_questions = quiz_info.get('questions')
    start_time = quiz_info.get('start_time')

    if not all([submission, server_questions, start_time]):
        return jsonify({"message": "Invalid quiz session data. Please start over."}), 400

    # --- Server-side score calculation ---
    final_score = 0
    total_time_taken = time.time() - start_time
    time_bonus_per_q = max(0, 30 - (total_time_taken / len(server_questions)))
    submitted_answers = submission.get('answers', [])
    
    for i, s_answer in enumerate(submitted_answers):
        if i < len(server_questions):
            correct_answer = server_questions[i]['answer']
            if s_answer.get('is_correct') and s_answer.get('answer') == correct_answer:
                final_score += (1000 + int(time_bonus_per_q * 50))

    # --- NEW: Save score to leaderboard, only keeping the highest score ---
    leaderboard = load_leaderboard()
    user_entry = next((item for item in leaderboard if item.get("username") == username), None)

    if user_entry:
        # User exists, update only if new score is higher
        if final_score > user_entry.get('score', 0):
            user_entry['score'] = final_score
            user_entry['timestamp'] = datetime.datetime.now().isoformat()
            log_activity(f"QUIZ: User '{username}' achieved a new high score: {final_score}.")
        else:
            log_activity(f"QUIZ: User '{username}' scored {final_score}, which did not beat their high score of {user_entry.get('score', 0)}.")
    else:
        # New user for the leaderboard
        leaderboard.append({
            "username": username,
            "full_name": session['full_name'],
            "score": final_score,
            "timestamp": datetime.datetime.now().isoformat()
        })
        log_activity(f"QUIZ: User '{username}' set their first score: {final_score}.")

    save_leaderboard(leaderboard)
    
    # Record this attempt to enforce the 24-hour cooldown
    attempts = load_attempts()
    attempts[username] = time.time()
    save_attempts(attempts)

    # Clear quiz from server-side memory
    quiz_sessions.pop(username, None)

    return jsonify({"message": "Score submitted successfully!", "final_score": final_score})

@app.route('/generate_report')
@permission_required('upload_reports')
def generate_report():
    """
    Looks up a study by Accession Number in the live PACS data cache,
    generates a blank DOCX report from it, and serves it to the user.
    """
    accession = request.args.get('accession', '').strip()

    if not accession:
        flash("No Accession Number provided to generate report.", "danger")
        return redirect(url_for('dashboard'))

    # Fetch the study record from the live PACS data cache
    pacs_study_record = None
    with pacs_cache_lock:
        # Use next() for a more efficient search
        pacs_study_record = next((study for study in pacs_data_cache.get('data', []) if study.get('accession') == accession), None)

    if not pacs_study_record:
        flash(f"Could not find a matching study for Accession '{accession}' in the recent PACS query. The study may be too old or the cache is updating.", "danger")
        return redirect(url_for('dashboard'))

    # The data from the PACS cache is already in the correct dictionary format
    # for the docx generator.
    study_data = pacs_study_record

    # Call the generator, saving to a temporary directory
    out_path, error = generate_docx_report_internal(study_data, output_dir=TEMP_REPORTS_DIR)

    if error:
        flash(f"Could not create report: {error}", "danger")
        return redirect(url_for('dashboard'))

    # Clean up the temporary file after the request is finished
    @after_this_request
    def cleanup(response):
        try:
            if os.path.exists(out_path):
                os.remove(out_path)
        except Exception as e:
            log_activity(f"Error cleaning up temp report file {out_path}: {e}")
        return response

    # Serve the file for download
    return send_file(out_path,
                     as_attachment=True,
                     download_name=os.path.basename(out_path),
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
# --- MAIN EXECUTION ---
if __name__ == "__main__":
    load_config()
    init_mwl_db()
    load_report_metadata()
    load_leaderboard()
    load_attempts()
    load_quiz_data()
    
    pacs_thread = threading.Thread(target=pacs_poller, daemon=True, name="PACS_Poller")
    approval_thread = threading.Thread(target=approval_poller, daemon=True, name="Approval_Poller")
    mwl_thread = threading.Thread(target=mwl_server_worker, daemon=True, name="MWL_Server") 
    archive_thread = threading.Thread(target=mwl_archive_worker, daemon=True, name="MWL_Archiver")
    radiologist_thread = threading.Thread(target=radiologist_update_poller, daemon=True, name="Radiologist_Poller")
    # ADDED: The new thread for cleaning up zip files
    zip_cleanup_thread = threading.Thread(target=cleanup_zip_files_worker, daemon=True, name="Zip_Cleanup_Worker")
    
    pacs_thread.start()
    approval_thread.start()
    mwl_thread.start()
    archive_thread.start()
    radiologist_thread.start()
    zip_cleanup_thread.start() # ADDED: Start the new thread
    
    app.run(host='0.0.0.0', port=5000, debug=False)
